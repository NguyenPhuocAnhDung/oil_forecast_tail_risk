import docx
import xml.etree.ElementTree as ET

def verify_doc(doc_path):
    print(f"\n=======================================================")
    print(f"VERIFYING: {doc_path}")
    print(f"=======================================================")
    doc = docx.Document(doc_path)
    xml_str = doc._element.xml
    root = ET.fromstring(xml_str)
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    }
    
    # 1. Inspect bold runs
    bold_headings = 0
    bold_non_headings = []
    
    for i, p in enumerate(doc.paragraphs):
        is_heading = (p.style.name.startswith('Heading') or 
                      p.paragraph_format.space_before == docx.shared.Pt(12) or 
                      p.paragraph_format.space_before == docx.shared.Pt(9) or 
                      p.paragraph_format.space_before == docx.shared.Pt(6) or 
                      "TÓM TẮT" in p.text or "ABSTRACT" in p.text or i == 0)
        for r in p.runs:
            if r.bold:
                if is_heading:
                    bold_headings += 1
                else:
                    bold_non_headings.append(f"P{i}: {p.text[:60]} -> Run text: {r.text}")
                    
    # Inspect tables for bold
    bold_table_cells = []
    for t_idx, tbl in enumerate(doc.tables):
        # ignore formula tables
        if len(tbl.columns) == 2 and len(tbl.rows) == 1:
            continue
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.bold:
                            bold_table_cells.append(f"Table {t_idx} Row {r_idx} Col {c_idx}: {cell.text[:30]}")
                            
    print(f"1. Typography & Bold Check:")
    print(f"   - Bold headings count: {bold_headings}")
    print(f"   - Bold non-headings count: {len(bold_non_headings)}")
    print(f"   - Bold table cells count: {len(bold_table_cells)}")
        
    # 2. Check XML for OMML equations & empty base boxes
    empty_e_count = xml_str.count("<m:e><m:r><m:t></m:t></m:r></m:e>")
    print(f"2. Math & OMML Check:")
    print(f"   - Empty OMML base elements: {empty_e_count}")
    
    # Check unformatted multi-quantile raw strings
    unformatted_q = "(q in {0.1, 0.5, 0.9})" in xml_str or "(q in {0.1,0.5,0.9})" in xml_str
    print(f"   - Raw unformatted '(q in {{0.1, 0.5, 0.9}})' present: {unformatted_q}")
    
    # 3. Check equation links & bookmarks
    eq_links = []
    tbl_links = []
    fig_links = []
    ref_links = []
    
    for hyperlink in root.findall('.//w:hyperlink', ns):
        anchor = hyperlink.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor', '')
        text = "".join(t.text for t in hyperlink.findall('.//w:t', ns) if t.text)
        if anchor.startswith('eq_'):
            eq_links.append((anchor, text))
        elif anchor.startswith('tbl_'):
            tbl_links.append((anchor, text))
        elif anchor.startswith('fig_'):
            fig_links.append((anchor, text))
        elif anchor.startswith('ref_'):
            ref_links.append((anchor, text))
            
    # Check bookmarks
    bookmarks = [b.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name', '') 
                 for b in root.findall('.//w:bookmarkStart', ns)]
    eq_bms = [b for b in bookmarks if b.startswith('eq_')]
    tbl_bms = [b for b in bookmarks if b.startswith('tbl_')]
    fig_bms = [b for b in bookmarks if b.startswith('fig_')]
    ref_bms = [b for b in bookmarks if b.startswith('ref_')]
    
    print(f"3. Bookmarks & Hyperlinks Cross-Referencing:")
    print(f"   - Equation Bookmarks: {len(eq_bms)} | In-text Equation Links: {len(eq_links)}")
    print(f"   - Table Bookmarks: {len(tbl_bms)} (Expected: tbl_1..tbl_10) | Found: {tbl_bms}")
    print(f"   - In-text Table Links: {len(tbl_links)} | Sample: {tbl_links[:6]}")
    print(f"   - Figure Bookmarks: {len(fig_bms)} (Expected: fig_1..fig_7) | Found: {fig_bms}")
    print(f"   - In-text Figure Links: {len(fig_links)} | Sample: {fig_links[:6]}")
    print(f"   - Reference Bookmarks: {len(ref_bms)} | In-text Citation Links: {len(ref_links)}")
    
    # 4. Check Table Background Shading (Pure White FFFFFF)
    non_white_shading = []
    shd_elements = root.findall('.//w:tcPr/w:shd', ns)
    for shd in shd_elements:
        fill_val = shd.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '')
        if fill_val.upper() != 'FFFFFF':
            non_white_shading.append(fill_val)
            
    print(f"4. Table Shading Check:")
    print(f"   - Total cell shading elements: {len(shd_elements)}")
    print(f"   - Non-white shading elements: {len(non_white_shading)}")
    if non_white_shading:
        print(f"     Non-white fills detected: {set(non_white_shading)}")
        
    print(f"5. Structural Integrity:")
    newlines_in_p = [i for i, p in enumerate(doc.paragraphs) if '\n' in p.text]
    print(f"   - Paragraphs with illegal raw '\\n': {len(newlines_in_p)}")

verify_doc("GUMNETHET_FAIR_v1_TIENG_VIET.docx")
verify_doc("GUMNETHET_FAIR_v1.docx")
