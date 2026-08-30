import docx
import zipfile
import xml.etree.ElementTree as ET
import re

def verify_template(docx_path):
    print("=" * 65)
    print(f"VERIFYING: {docx_path}")
    print("=" * 65)
    
    doc = docx.Document(docx_path)
    
    # 1. Content & Consistency Check
    full_text = " ".join([p.text for p in doc.paragraphs])
    
    print("\n1. Sample Size Consistency:")
    n_4512 = full_text.count("4.512")
    n_4517 = full_text.count("4.517")
    n_4486 = full_text.count("4.486")
    print(f"   - Occurrences of '4.512': {n_4512}")
    print(f"   - Occurrences of '4.517': {n_4517} (Should be 0)")
    print(f"   - Occurrences of '4.486': {n_4486}")
    
    print("\n2. Vietnam Petroleum Context & Removal of Retail/BOG:")
    has_dung_quat = "Dung Quất" in full_text
    has_nghi_son = "Nghi Sơn" in full_text
    has_singapore = "Singapore" in full_text
    has_mops = "MOPS" in full_text or "Platts" in full_text
    has_petrolimex = "Petrolimex" in full_text
    has_pvoil = "PVOIL" in full_text
    has_bog = "Quỹ bình ổn" in full_text or "BOG" in full_text
    has_step_fn = "step-function" in full_text
    
    print(f"   - Dung Quất mentioned: {has_dung_quat}")
    print(f"   - Nghi Sơn mentioned: {has_nghi_son}")
    print(f"   - Singapore Platts mentioned: {has_singapore and has_mops}")
    print(f"   - Petrolimex & PVOIL mentioned: {has_petrolimex and has_pvoil}")
    print(f"   - 'Quỹ bình ổn' / BOG present: {has_bog} (Should be False)")
    print(f"   - 'step-function' present: {has_step_fn} (Should be False)")

    print("\n3. Typography & Bold Check:")
    bold_headings = 0
    bold_non_headings = 0
    bold_table_cells = 0
    
    for p in doc.paragraphs:
        is_heading = p.style.name in ['Heading 1', 'Heading 2', 'Heading 3', 'H1x', 'H2x', 'PaperTitle']
        has_bold = any(r.bold for r in p.runs if r.text.strip())
        if has_bold:
            if is_heading:
                bold_headings += 1
            else:
                bold_non_headings += 1
                
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if any(run.bold for run in p.runs if run.text.strip()):
                        bold_table_cells += 1
                        
    print(f"   - Bold headings count: {bold_headings}")
    print(f"   - Bold non-headings count: {bold_non_headings}")
    print(f"   - Bold table cells count: {bold_table_cells} (Should be 0)")

    print("\n4. Table Properties & Shading:")
    print(f"   - Total tables: {len(doc.tables)}")
    with zipfile.ZipFile(docx_path, 'r') as z:
        tree = ET.fromstring(z.read('word/document.xml'))
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        shds = tree.findall('.//w:shd', ns)
        non_white = [s.get(f'{{{ns["w"]}}}fill') for s in shds if s.get(f'{{{ns["w"]}}}fill') not in ['FFFFFF', 'ffffff', None]]
        print(f"   - Total cell shading elements: {len(shds)}")
        print(f"   - Non-white shading elements: {len(non_white)} (Should be 0)")

    print("\n5. Bookmarks & Cross-Referencing:")
    with zipfile.ZipFile(docx_path, 'r') as z:
        tree = ET.fromstring(z.read('word/document.xml'))
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        bms = [b.get(f'{{{ns["w"]}}}name') for b in tree.findall('.//w:bookmarkStart', ns)]
        tbl_bms = [b for b in bms if b and b.startswith('tbl_')]
        fig_bms = [b for b in bms if b and b.startswith('fig_')]
        print(f"   - Table Bookmarks: {len(tbl_bms)} | Found: {tbl_bms}")
        print(f"   - Figure Bookmarks: {len(fig_bms)} | Found: {fig_bms}")

    print("\n6. Embedded Media:")
    with zipfile.ZipFile(docx_path, 'r') as z:
        media_files = [f for f in z.namelist() if f.startswith('word/media/')]
        print(f"   - Media files count: {len(media_files)}")
        for mf in media_files:
            print(f"     * {mf} ({z.getinfo(mf).file_size} bytes)")

    print("\n" + "=" * 65)
    print("VERIFICATION COMPLETED!")
    print("=" * 65)

if __name__ == '__main__':
    verify_template('GUMNETHet_FAIRv3_template.docx')
