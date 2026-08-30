import docx
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc_v3 = docx.Document('GUMNETHet_FAIRv3.docx')

print(f"Total Paragraphs in GUMNETHet_FAIRv3.docx: {len(doc_v3.paragraphs)}")
print(f"Total Tables in GUMNETHet_FAIRv3.docx: {len(doc_v3.tables)}")

for i, p in enumerate(doc_v3.paragraphs):
    # check if paragraph has drawings or math
    drawings = len(p._p.xpath('.//w:drawing'))
    maths = len(p._p.xpath('.//m:oMath'))
    print(f"P{i:02d} [{p.style.name:12s}] (draw={drawings}, math={maths}): len={len(p.text):4d} | {p.text[:65]}")
