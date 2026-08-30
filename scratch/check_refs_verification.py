import docx
import re
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document('GUMNETHet_FAIRv4_final_IEEE.docx')

print("=== CHECKING CITATIONS IN TEXT ===")
full_text = "\n".join([p.text for p in doc.paragraphs])
cited_refs = set()

# Find all [X] or [X], [Y] in text
matches = re.findall(r'\[(\d+)\]', full_text)
for m in matches:
    cited_refs.add(int(m))

print(f"Total distinct reference numbers found in document: {len(cited_refs)}")
print(f"Numbers: {sorted(list(cited_refs))}")

missing = [i for i in range(1, 29) if i not in cited_refs]
print(f"Missing references (1-28): {missing}")

print("\n=== VERIFYING REFERENCE PARAGRAPHS ===")
ref_paras = [p.text.strip() for p in doc.paragraphs if p.style.name == 'references']
print(f"Total reference items in bibliography: {len(ref_paras)}")
for i, r in enumerate(ref_paras):
    print(f"  {r[:80]}...")
