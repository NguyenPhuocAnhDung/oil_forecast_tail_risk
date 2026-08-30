import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

src_doc = docx.Document('GUMNETHet_FAIRv3.docx')
print(f"Total tables in FAIRv3: {len(src_doc.tables)}")

for idx, t in enumerate(src_doc.tables):
    rows = len(t.rows)
    cols = len(t.columns)
    first_cell = t.cell(0,0).text.strip()[:30]
    tblW = t._tbl.tblPr.xpath('./w:tblW')
    w_val = tblW[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w') if tblW else 'None'
    w_type = tblW[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') if tblW else 'None'
    print(f"Table {idx:02d} ({rows}x{cols}) w={w_val} ({w_type}): first_cell={repr(first_cell)}")
