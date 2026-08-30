import docx
import zipfile
import xml.etree.ElementTree as ET
import re

docx_path = r'D:\1.Cong_Viec\NguyencuuKhoaHoc\oil_forecast_tail_risk\GUMNETHet_FAIRv6_final.docx'

doc = docx.Document(docx_path)
print(f"Total paragraphs in doc: {len(doc.paragraphs)}")
print(f"Total tables in doc: {len(doc.tables)}")

# Print all paragraphs around equations and comments
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if any(k in text for k in ['(1)', '(2)', '(3)', '(4)', '(5)', '(6)', '(7)', '(8)', '(9)', '(10)', 'Fig', 'Figure', 'github', 'GitHub', 'Abstract', 'Volatilit']):
        print(f"P[{i}]: {text[:120]}")
