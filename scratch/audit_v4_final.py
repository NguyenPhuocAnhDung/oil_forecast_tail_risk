import docx
import sys
import zipfile
import xml.etree.ElementTree as ET

sys.stdout.reconfigure(encoding='utf-8')

f = 'GUMNETHet_FAIRv4_final.docx'
doc = docx.Document(f)
print('='*75)
print(f'AUDIT OF IEEE FINAL MANUSCRIPT: {f}')
print('='*75)
print(f'Paragraph count: {len(doc.paragraphs)}')
print(f'Table count: {len(doc.tables)}')
print(f'Section count: {len(doc.sections)}')

print('\n--- SECTIONS CONFIGURATION ---')
for i, s in enumerate(doc.sections):
    sectPr = s._sectPr
    cols = sectPr.xpath('./w:cols')
    col_str = 'no cols'
    if cols:
        col_str = f'num={cols[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num")}, space={cols[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space")}'
    print(f'Section {i}: start_type={s.start_type}, top={s.top_margin.pt:.1f}pt, bot={s.bottom_margin.pt:.1f}pt, left={s.left_margin.pt:.1f}pt, right={s.right_margin.pt:.1f}pt | cols: {col_str}')

print('\n--- KEY PARAGRAPHS & HEADINGS ---')
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip().replace('\n', ' ')
    if any(txt.startswith(h) for h in ['Robust', 'Phuoc', 'Danh', 'Van', 'Abstract', 'Keywords', 'I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'REFERENCES', 'A.', 'B.', 'C.', '1)', '2)', '3)', '4)', '5)', '6)', 'Fig.', 'TABLE']):
        print(f'P{i:02d} [{p.style.name:15s}]: {txt[:70]}')

print('\n--- ALL TABLES ---')
for i, t in enumerate(doc.tables):
    txts = [c.text.strip().replace('\n', ' ') for r in t.rows for c in r.cells]
    print(f'Table {i:02d} ({len(t.rows)}x{len(t.columns)}): {txts[:4]}')

with zipfile.ZipFile(f, 'r') as z:
    media = [m for m in z.namelist() if m.startswith('word/media/')]
    print(f'\n--- MEDIA IMAGES ({len(media)}) ---')
    for m in media:
        print(f'  {m} ({len(z.read(m))} bytes)')
        
    tree = ET.fromstring(z.read('word/document.xml'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', 'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    math_count = len(tree.findall('.//m:oMath', ns))
    bms = [b.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name') for b in tree.findall('.//w:bookmarkStart', ns)]
    tbl_bms = [b for b in bms if b and b.startswith('tbl_')]
    fig_bms = [b for b in bms if b and b.startswith('fig_')]
    ref_bms = [b for b in bms if b and b.startswith('ref_')]
    print(f'\n--- STATS ---')
    print(f'OMML Math Formulas: {math_count}')
    print(f'Table Bookmarks ({len(tbl_bms)}): {tbl_bms}')
    print(f'Figure Bookmarks ({len(fig_bms)}): {fig_bms}')
    print(f'Reference Bookmarks ({len(ref_bms)}): count={len(ref_bms)}')
