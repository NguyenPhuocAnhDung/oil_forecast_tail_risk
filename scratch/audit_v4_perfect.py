import docx
import sys
import zipfile
import xml.etree.ElementTree as ET

sys.stdout.reconfigure(encoding='utf-8')

f = 'GUMNETHet_FAIRv4_final.docx'
doc = docx.Document(f)
print('='*75)
print(f'FINAL AUDIT OF MANUSCRIPT: {f}')
print('='*75)
print(f'Total Paragraphs: {len(doc.paragraphs)}')
print(f'Total Tables: {len(doc.tables)}')
print(f'Total Sections: {len(doc.sections)}')

s = doc.sections[0]
print(f'Page Margins: Top={s.top_margin.pt:.1f}pt, Bottom={s.bottom_margin.pt:.1f}pt, Left={s.left_margin.pt:.1f}pt, Right={s.right_margin.pt:.1f}pt')

print('\n--- FIRST 15 PARAGRAPHS ---')
for i in range(15):
    p = doc.paragraphs[i]
    print(f'P{i:02d} [{p.style.name:12s}]: {p.text[:70]}')

print('\n--- ALL 13 TABLES ---')
for i, t in enumerate(doc.tables):
    txts = [c.text.strip().replace('\n', ' ') for r in t.rows for c in r.cells]
    print(f'Table {i:02d} ({len(t.rows)}x{len(t.columns)}): {txts[:4]}')

with zipfile.ZipFile(f, 'r') as z:
    media = [m for m in z.namelist() if m.startswith('word/media/')]
    print(f'\n--- MEDIA IMAGES ({len(media)}) ---')
    for m in media:
        print(f'  {m} ({len(z.read(m))} bytes)')
        
    tree = ET.fromstring(z.read('word/document.xml'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', 'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    math_count = len(tree.findall('.//m:oMath', ns))
    bms = [b.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name') for b in tree.findall('.//w:bookmarkStart', ns)]
    tbl_bms = [b for b in bms if b and b.startswith('tbl_')]
    fig_bms = [b for b in bms if b and b.startswith('fig_')]
    ref_bms = [b for b in bms if b and b.startswith('ref_')]
    print(f'\n--- STATS ---')
    print(f'OMML Math Formulas: {math_count}')
    print(f'Table Bookmarks ({len(tbl_bms)}): {tbl_bms}')
    print(f'Figure Bookmarks ({len(fig_bms)}): {fig_bms}')
    print(f'Reference Bookmarks ({len(ref_bms)}): count={len(ref_bms)}')
