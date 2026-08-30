import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document('conference-template-a4_transitional.docx')
print(f'Total Paragraphs: {len(doc.paragraphs)}')
print(f'Total Tables: {len(doc.tables)}')
print(f'Total Sections: {len(doc.sections)}')

print('\n--- ALL PARAGRAPHS IN TEMPLATE ---')
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip().replace('\n', ' ')
    if txt or p.style.name != 'Normal':
        print(f'P{i:02d} [{p.style.name:15s}]: {txt[:70]}')

if doc.tables:
    print('\n--- ALL TABLES IN TEMPLATE ---')
    for t_idx, t in enumerate(doc.tables):
        print(f'Table {t_idx} ({len(t.rows)}x{len(t.columns)}):')
        for r_idx, r in enumerate(t.rows):
            cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
            print(f'  Row {r_idx}: {cells}')
