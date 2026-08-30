import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document('GUMNETHet_FAIRv4_final_IEEE.docx')

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print(f"Total sections: {len(doc.sections)}")

for i, sec in enumerate(doc.sections):
    print(f"Section {i}: start_type={sec.start_type}, page_width={sec.page_width.inches}in, cols={sec._sectPr.xpath('./w:cols/@w:num')}")

for idx, tbl in enumerate(doc.tables):
    tblW = tbl._tbl.xpath('./w:tblPr/w:tblW/@w:w')
    cols = tbl._tbl.xpath('./w:tblGrid/w:gridCol/@w:w')
    print(f"Table {idx}: rows={len(tbl.rows)}, cols={len(tbl.columns)}, tblW={tblW}, gridCols={cols[:4]}...")
