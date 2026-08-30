import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

f = 'GUMNETHet_FAIRv4_final.docx'
doc = docx.Document(f)

print(f"=== FULL AUDIT OF SANITIZED MANUSCRIPT: {f} ===")
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")

for idx, t in enumerate(doc.tables):
    tblW = t._tbl.tblPr.xpath('./w:tblW')
    w_val = tblW[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w') if tblW else 'None'
    borders = len(t._tbl.tblPr.xpath('./w:tblBorders/*[@w:val!="none"]'))
    cell_widths = [c.width for c in t.rows[0].cells]
    print(f"Table {idx:02d} ({len(t.rows)}x{len(t.columns)}): tblW={w_val}, visible_borders={borders}, cell_widths={cell_widths}")
