import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

src_doc = docx.Document('GUMNETHet_FAIRv3.docx')

print(f"Total paragraphs in FAIRv3: {len(src_doc.paragraphs)}")

for i, p in enumerate(src_doc.paragraphs):
    txt = p.text.strip()
    if txt and not txt.startswith('['):
        print(f"\n--- FAIRv3 P{i:02d} [{p.style.name}] ---")
        print(txt)
