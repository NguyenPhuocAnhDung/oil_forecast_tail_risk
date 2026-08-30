import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_vi = docx.Document('GUMNETHet_FAIRv4_final_VI.docx')
doc_en = docx.Document('GUMNETHet_FAIRv4_final_IEEE.docx')

print("EXTRACTING BILINGUAL SENTENCE MAPPINGS...")
for i, (p_v, p_e) in enumerate(zip(doc_vi.paragraphs, doc_en.paragraphs)):
    tv = p_v.text.strip()
    te = p_e.text.strip()
    st = p_v.style.name
    if tv and not tv.startswith('[') and st not in ['references', 'Author']:
        print(f"\n=== SECTION / PARAGRAPH {i} [{st}] ===")
        print(f"[VI]: {tv}")
        print(f"[EN]: {te}")
