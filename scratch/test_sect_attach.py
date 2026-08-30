import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document('conference-template-a4_clean_styles.docx')
body_elem = doc._body._body
for p in list(body_elem.xpath('./w:p')):
    body_elem.remove(p)

p1 = doc.add_paragraph("Paragraph in Section 1 (2 cols)")
s1_xml = (
    f'<w:sectPr {nsdecls("w")}>'
    f'<w:pgSz w:w="11906" w:h="16838"/>'
    f'<w:pgMar w:top="720" w:right="720" w:bottom="1080" w:left="720"/>'
    f'<w:cols w:num="2" w:space="360"/>'
    f'<w:type w:val="continuous"/>'
    f'</w:sectPr>'
)
p1._p.get_or_add_pPr().append(parse_xml(s1_xml))

p2 = doc.add_paragraph("Paragraph in Section 2 (1 col)")
s2_xml = (
    f'<w:sectPr {nsdecls("w")}>'
    f'<w:pgSz w:w="11906" w:h="16838"/>'
    f'<w:pgMar w:top="720" w:right="720" w:bottom="1080" w:left="720"/>'
    f'<w:cols w:num="1" w:space="720"/>'
    f'<w:type w:val="continuous"/>'
    f'</w:sectPr>'
)
p2._p.get_or_add_pPr().append(parse_xml(s2_xml))

p3 = doc.add_paragraph("Paragraph in Section 3 (2 cols)")

doc.save("scratch/test_sect_attach.docx")
print("Saved test_sect_attach.docx successfully! Sections:", len(doc.sections))
for i, s in enumerate(doc.sections):
    print(f"  Sec {i}: cols={s._sectPr.xpath('./w:cols/@w:num')}")
