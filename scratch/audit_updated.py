import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

f = 'GUMNETHet_FAIRv4_final_updated.docx'
doc = docx.Document(f)

print(f"=== AUDIT OF {f} ===")
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")
print(f"Total Sections: {len(doc.sections)}")

print("\n--- AUTHOR TABLE (Table 0) ---")
t0 = doc.tables[0]
print(f"Table 0 shape: {len(t0.rows)} rows x {len(t0.columns)} cols")
for c_idx, cell in enumerate(t0.rows[0].cells):
    print(f"  Col {c_idx+1} text:\n{cell.text}\n")

print("--- FIRST 10 PARAGRAPHS ---")
for i in range(10):
    p = doc.paragraphs[i]
    print(f"P{i:02d} [{p.style.name:15s}]: len={len(p.text):4d} | {p.text[:65]}")
