import docx
import sys
import re

sys.stdout.reconfigure(encoding='utf-8')

def cross_audit():
    doc_v3 = docx.Document('GUMNETHet_FAIRv3.docx')
    doc_vi = docx.Document('GUMNETHet_FAIRv4_final_VI.docx')
    doc_en = docx.Document('GUMNETHet_FAIRv4_final_IEEE.docx')

    print("==========================================================================")
    print("CROSS-VERIFICATION AUDIT: FAIRv3 vs FAIRv4_VI vs FAIRv4_EN (IEEE Template)")
    print("==========================================================================")
    
    print(f"Original FAIRv3: {len(doc_v3.paragraphs)} paragraphs, {len(doc_v3.tables)} tables")
    print(f"Final IEEE VI:   {len(doc_vi.paragraphs)} paragraphs, {len(doc_vi.tables)} tables")
    print(f"Final IEEE EN:   {len(doc_en.paragraphs)} paragraphs, {len(doc_en.tables)} tables")

    # 1. Verify Authors
    print("\n--- 1. AUTHOR BLOCK VERIFICATION ---")
    tbl_vi = doc_vi.tables[0]
    tbl_en = doc_en.tables[0]
    print(f"VI Authors: {[c.text.splitlines()[0] for c in tbl_vi.rows[0].cells]}")
    print(f"EN Authors: {[c.text.splitlines()[0] for c in tbl_en.rows[0].cells]}")
    print(f"VI Emails:  {[c.text.splitlines()[-1] for c in tbl_vi.rows[0].cells]}")
    print(f"EN Emails:  {[c.text.splitlines()[-1] for c in tbl_en.rows[0].cells]}")

    # 2. Verify Headings alignment
    print("\n--- 2. STRUCTURE & HEADINGS ALIGNMENT ---")
    h_vi = [(i, p.style.name, p.text.strip()) for i, p in enumerate(doc_vi.paragraphs) if 'Heading' in p.style.name or 'title' in p.style.name or 'Abstract' in p.style.name or 'Keywords' in p.style.name]
    h_en = [(i, p.style.name, p.text.strip()) for i, p in enumerate(doc_en.paragraphs) if 'Heading' in p.style.name or 'title' in p.style.name or 'Abstract' in p.style.name or 'Keywords' in p.style.name]

    print(f"Total structured headings/sections in VI: {len(h_vi)}")
    print(f"Total structured headings/sections in EN: {len(h_en)}")
    
    for (idx_v, st_v, txt_v), (idx_e, st_e, txt_e) in zip(h_vi, h_en):
        print(f"  [{st_v:12s}] VI: {txt_v[:40]:40s} <===> EN: {txt_e[:40]}")

    # 3. Verify Tables & Figures Alignment
    print("\n--- 3. TABLES AND EQUATIONS VERIFICATION ---")
    print(f"Table count match: VI ({len(doc_vi.tables)}) == EN ({len(doc_en.tables)})")
    for t_idx in range(len(doc_vi.tables)):
        t_v = doc_vi.tables[t_idx]
        t_e = doc_en.tables[t_idx]
        print(f"  Table #{t_idx:02d}: shape=({len(t_v.rows)}x{len(t_v.columns)}), VI text sample={repr(t_v.cell(0,0).text[:25])} <===> EN text sample={repr(t_e.cell(0,0).text[:25])}")

    # 4. Verify References
    print("\n--- 4. REFERENCES COUNT & NUMBERING VERIFICATION ---")
    refs_vi = [p.text.strip() for p in doc_vi.paragraphs if p.style.name == 'references']
    refs_en = [p.text.strip() for p in doc_en.paragraphs if p.style.name == 'references']
    print(f"VI References count: {len(refs_vi)}")
    print(f"EN References count: {len(refs_en)}")
    print(f"First VI Ref: {refs_vi[0][:60]}")
    print(f"First EN Ref: {refs_en[0][:60]}")
    print(f"Last VI Ref:  {refs_vi[-1][:60]}")
    print(f"Last EN Ref:  {refs_en[-1][:60]}")

if __name__ == '__main__':
    cross_audit()
