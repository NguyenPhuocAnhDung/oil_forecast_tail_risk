import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_vi = docx.Document('GUMNETHet_FAIRv4_final_VI.docx')
doc_en = docx.Document('GUMNETHet_FAIRv4_final_IEEE.docx')

print("==========================================================================")
print("PARAGRAPH-BY-PARAGRAPH COMPARISON (VI vs EN)")
print("==========================================================================")

for i in range(len(doc_vi.paragraphs)):
    p_vi = doc_vi.paragraphs[i]
    p_en = doc_en.paragraphs[i]
    st = p_vi.style.name
    
    # Check text lengths and previews
    txt_v = p_vi.text.strip()
    txt_e = p_en.text.strip()
    
    if len(txt_v) > 0 and not txt_v.startswith('[') and not any(k in st for k in ['references', 'table', 'figure', 'Author']):
        print(f"\n--- [P{i:02d}] Style: {st} ---")
        print(f"  [VI ({len(txt_v)} chars)]: {txt_v[:140]}...")
        print(f"  [EN ({len(txt_e)} chars)]: {txt_e[:140]}...")
