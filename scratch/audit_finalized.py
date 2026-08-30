import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

def audit_file(f):
    print(f"\n==================================================")
    print(f"AUDITING: {f}")
    print(f"==================================================")
    doc = docx.Document(f)
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    print(f"Total Sections: {len(doc.sections)}")
    
    print("\n--- SAMPLE HEADINGS & CAPTIONS ---")
    for i, p in enumerate(doc.paragraphs):
        st = p.style.name if p.style else ''
        if any(k in st for k in ['Heading', 'title', 'head', 'caption', 'Abstract', 'Keywords', 'references']):
            print(f"P{i:02d} [{st:15s}]: {p.text[:75]}")

audit_file('GUMNETHet_FAIRv4_final.docx')
audit_file('GUMNETHet_FAIRv4_final_VI.docx')
