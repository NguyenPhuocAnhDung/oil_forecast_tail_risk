import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os
import re
import pandas as pd
import numpy as np

# Load seed42 metrics for verification
df_metrics = pd.read_csv('seed42_metrics.csv')

def create_document():
    # Use template as base to preserve exact styles, section margins, header/footer
    doc = docx.Document('GUMNetHet_FAIR_ban_thao_template.docx')
    
    # Helper to set 100% pure white cell formatting with standard 3-line borders
    def style_table(table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = table._tbl.tblPr
        
        # Remove old table borders if any
        for old_b in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old_b)
            
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                            f'<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                            f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
                            f'<w:insideV w:val="none"/>'
                            f'<w:left w:val="none"/>'
                            f'<w:right w:val="none"/>'
                            f'</w:tblBorders>')
        tblPr.append(borders)
        
        for row_idx, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            if row_idx == 0:
                trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
                
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                
                # Remove ALL existing shading elements
                for old_shd in tcPr.findall(qn('w:shd')):
                    tcPr.remove(old_shd)
                    
                # 100% Pure white cell shading
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="FFFFFF"/>')
                tcPr.append(shd)
                
                # Margins
                mar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                                f'<w:top w:w="80" w:type="dxa"/>'
                                f'<w:bottom w:w="80" w:type="dxa"/>'
                                f'<w:left w:w="100" w:type="dxa"/>'
                                f'<w:right w:w="100" w:type="dxa"/>'
                                f'</w:tcMar>')
                tcPr.append(mar)
                
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.05
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(8.5)
                        r.bold = False

    # Apply styling to all tables in doc
    for t in doc.tables:
        style_table(t)

    # Helper function to add table caption with bookmark
    def make_table_caption(p, text, tbl_id):
        p.style = 'Caption'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.text = ""
        bm_id = f"300{tbl_id}"
        bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="tbl_{tbl_id}"/>')
        bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>')
        p._p.append(bm_start)
        p._p.append(bm_end)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8.5)
        run.bold = False

    # Helper function to add figure caption with bookmark
    def make_figure_caption(p, text, fig_id):
        p.style = 'Caption'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        p.text = ""
        bm_id = f"400{fig_id}"
        bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="fig_{fig_id}"/>')
        bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>')
        p._p.append(bm_start)
        p._p.append(bm_end)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8.0)
        run.bold = False

    # 1. Update Section 1: GIỚI THIỆU (Paragraphs 8, 9, 10) to incorporate Comment 1
    intro_p8 = (
        "Thị trường xăng dầu Việt Nam có đặc thù cấu trúc nguồn cung mang tính chiến lược: khoảng 70% tổng sản lượng "
        "tiêu thụ nội địa được cung ứng bởi hai nhà máy lọc hóa dầu trong nước (Dung Quất và Nghi Sơn), trong khi 30% "
        "nhu cầu còn lại bắt buộc phải nhập khẩu trực tiếp từ các thị trường quốc tế. Trong đó, thị trường Singapore là "
        "địa bàn nhập khẩu trọng yếu nhất, với giá giao dịch thành phẩm Mean of Platts Singapore (MOPS)—tiêu biểu là Mogas 95 "
        "(MG95) và Gasoil 0.001%S (DO 0.001%)—đóng vai trò là hệ quy chiếu định giá cơ sở cho mọi hợp đồng thương mại, "
        "tính toán giá vốn và quản lý chuỗi cung ứng. Tuy nhiên, giá Platts biến động cực kỳ phức tạp do phản ứng đồng thời "
        "với các cú sốc cung–cầu toàn cầu, biến động tỷ giá và rủi ro địa chính trị (GPR) [1]–[3]. Các giai đoạn sụp đổ giá dầu "
        "2014–2016, chiến tranh giá OPEC+ năm 2020, xung đột Nga–Ukraine 2022 và căng thẳng Biển Đỏ 2023–2024 cho thấy chuỗi giá "
        "thành phẩm thường xuyên xuất hiện bước nhảy phi tuyến, dịch chuyển chế độ và phân cụm biến động mạnh. Do đó, nhu cầu "
        "dự báo chính xác giá xăng dầu Platts trong ngắn hạn (H1–H7) và trung hạn (H10–H60) là đòi hỏi cấp thiết phục vụ trực "
        "tiếp cho việc ra quyết định kinh doanh, tối ưu hóa kế hoạch mua hàng, quản trị tồn kho và phòng hộ rủi ro (hedging) "
        "của các doanh nghiệp đầu mối xăng dầu lớn như Petrolimex và PVOIL."
    )
    
    intro_p9 = (
        "Các kiến trúc chuỗi thời gian hiện đại như PatchTST [8], iTransformer [9], TimesNet [10], DLinear [11], Mamba [15] và "
        "Chronos [16] đã cải thiện đáng kể hiệu năng trên các benchmark tổng quát, nhưng phần lớn đều xử lý toàn bộ các biến đầu vào "
        "trong một không gian biểu diễn tương đối đồng nhất. Với dữ liệu năng lượng, ba nhóm tín hiệu có bản chất cơ bản khác nhau: "
        "động lượng giá tần số cao, trạng thái vĩ mô biến đổi chậm và phản ứng phi tuyến nhạy cú sốc biên độ lớn đòi hỏi các inductive "
        "bias chuyên biệt. Các mạng kết hợp chuyên gia (MoE) truyền thống [20]–[22] thường đưa cùng một tập đặc trưng tới mọi expert, "
        "làm suy giảm mức độ chuyên môn hóa."
    )
    
    intro_p10 = (
        "GUMNetHet giải quyết triệt để điểm nghẽn này bằng kỹ thuật phân vùng đặc trưng (feature partitioning): nhóm giá và benchmark "
        "được xử lý bởi CNN-1D đa tỷ lệ; nhóm vĩ mô và chỉ số GPR bởi GRU-Attention; nhóm tỷ lệ crack-spread và độ biến động bởi Wavelet-KAN. "
        "Ba biểu diễn được hợp nhất linh hoạt thông qua một bộ định tuyến (gating router) phụ thuộc vào horizon dự báo và ngữ cảnh thị trường. "
        "Đóng góp chính của bài báo gồm: (i) Kiến trúc MoE dị thể với phân vùng đặc trưng theo bản chất kinh tế và miền tần số; (ii) Bộ định tuyến "
        "nhận biết horizon (horizon-aware routing); (iii) Đầu ra đa phân vị (multi-quantile head) kết hợp residual scaling để kiểm soát độ trôi phương sai; "
        "và (iv) Đánh giá thực nghiệm mở rộng walk-forward trên N=4.512 ngày giao dịch với bảng kết quả đầy đủ các chỉ số MAE, RMSE, MAPE, R² và DA%."
    )

    # Set intro text
    doc.paragraphs[8].text = intro_p8
    doc.paragraphs[9].text = intro_p9
    doc.paragraphs[10].text = intro_p10

    # Ensure Abstract has exact N=4.512
    abstract_text = (
        "Tóm tắt—Dự báo giá xăng dầu thành phẩm trở nên khó khăn khi chuỗi giá chịu đồng thời biến động ngắn hạn, dịch chuyển chế độ vĩ mô "
        "và các cú sốc địa chính trị có tính phi tuyến, đuôi dày. Bài báo đề xuất GUMNetHet (Heterogeneous Gated Unified Mixture Network), "
        "một mạng kết hợp chuyên gia không đồng nhất cho dự báo xác suất đa chu kỳ. Mô hình phân vùng đặc trưng thành ba nhóm và giao cho "
        "ba chuyên gia chuyên biệt: CNN-1D đa tỷ lệ cho động lượng giá, GRU-Attention cho chế độ vĩ mô, và Wavelet-KAN cho quan hệ phi tuyến "
        "nhạy cú sốc. Bộ định tuyến nhận biết horizon kết hợp biểu diễn chuyên gia, nhúng horizon và thống kê ngữ cảnh để phân bổ trọng số động; "
        "đầu ra đa phân vị q ∈ {0.1, 0.5, 0.9} được tối ưu bằng pinball loss và điều chuẩn cân bằng tải. Trên dữ liệu đa nguồn 11/2008–04/2026 "
        "(N=4.512 quan sát) với expanding walk-forward, GUMNetHet đạt MAE thấp nhất trong nhóm baseline được báo cáo ở cả bảy horizon H1–H60 cho "
        "MG95 và DO 0.001%. Tại H60, MAE giảm 30,1% cho xăng và 22,9% cho dầu so với baseline tốt nhất tương ứng. Directional accuracy đạt "
        "90,95–95,56% ở H1–H7 của xăng và 76,65–84,92% ở H1–H7 của dầu, nhưng giảm mạnh tại H10/H60, cho thấy cần phân biệt giữa độ chính xác "
        "mức giá và độ chính xác hướng ở horizon dài. Khoảng phân vị 80% đạt PICP=82,4% với PINAW=0,142. Với 0,34M tham số và độ trễ 1,42 ms/mẫu, "
        "GUMNetHet cho thấy sự cân bằng tốt giữa độ chính xác, bất định và chi phí tính toán."
    )
    doc.paragraphs[5].text = abstract_text

    # Ensure Data section has exact N=4.512
    data_p25 = (
        "Dữ liệu bao phủ 03/11/2008–30/04/2026 với N=4.512 quan sát ngày giao dịch. Hai target được báo cáo là MG95 và DO 0.001% (USD/thùng). "
        "Tập biến ngoại sinh gồm WTI, Brent DTD, chênh lệch liên sản phẩm (MG92, MG97, KERO, FO 180, Naphtha), chỉ số GPR [1], DXY, tỷ lệ crack spread "
        "và độ biến động thực tế. Lookback cố định L=30 ngày; horizon dự báo gồm h ∈ {1, 3, 5, 7, 10, 20, 60}. Giao thức expanding walk-forward chia tập ban đầu "
        "70% train, 10% validation và 20% test; sau mỗi bước cuộn, cửa sổ train được mở rộng để phản ánh điều kiện vận hành thực tế."
    )
    doc.paragraphs[25].text = data_p25

    # Refine Table captions and references to sequential numbering: Bảng 1, Bảng 2, Bảng 3, Bảng 4, Bảng 5
    make_table_caption(doc.paragraphs[29], "Bảng 1. Chẩn đoán thống kê rút gọn của các chuỗi chính.", "1")
    make_table_caption(doc.paragraphs[39], "Bảng 2. Kết quả chi tiết trên MG95 (Seed=42). MAE/RMSE tính theo USD/thùng; giá trị DA là phần trăm.", "2")
    make_table_caption(doc.paragraphs[40], "Bảng 3. Kết quả chi tiết trên DO 0.001% (Seed=42). Bảng nguồn thực nghiệm không báo cáo BiMamba cho mục tiêu dầu.", "3")
    make_table_caption(doc.paragraphs[47], "Bảng 4. Ablation rút gọn của GUMNetHet (Seed=42).", "4")
    make_table_caption(doc.paragraphs[52], "Bảng 5. Chi phí tính toán trên GPU Tesla T4.", "5")

    # Update in-text mention of Bảng 4 (was typed as Bảng 5 in P46)
    doc.paragraphs[46].text = (
        "Khoảng [q0.1, q0.9] đạt PICP=82,4% so với danh định 80% và PINAW=0,142. Fan chart cho thấy biên bất định mở rộng khi biến động tăng. "
        "Ablation trong Bảng 4 cho thấy thay Wav-KAN bằng MLP gây suy giảm lớn nhất trong các biến thể expert; router đồng nhất cũng làm MAE tăng "
        "đáng kể, củng cố vai trò của chuyên môn hóa và định tuyến thích ứng."
    )

    # Refine Figure captions
    make_figure_caption(doc.paragraphs[20], "Hình 1. Kiến trúc GUMNetHet: phân vùng đặc trưng, ba expert dị thể, horizon-aware router và multi-quantile head.", "1")
    make_figure_caption(doc.paragraphs[38], "Hình 2. Đường cong MAE và R² qua bảy horizon cho MG95 và DO 0.001% (Seed=42).", "2")
    make_figure_caption(doc.paragraphs[44], "Hình 3. Directional accuracy (DA%) của GUMNetHet và các baseline qua H1–H60 cho xăng và dầu.", "3")
    make_figure_caption(doc.paragraphs[50], "Hình 4. Trên: fan chart đa phân vị dưới biến động mạnh. Dưới: trọng số router trong chế độ GPR thấp và GPR cao.", "4")

    # Apply font consistency to all paragraphs
    for p in doc.paragraphs:
        if p.style.name == 'Normal':
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9.5)
        elif p.style.name == 'Abstract':
            p.paragraph_format.line_spacing = 1.10
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9.0)
        elif p.style.name == 'Refs':
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(8.0)

    # Save to GUMNETHet_FAIRv3_template.docx
    output_path = 'GUMNETHet_FAIRv3_template.docx'
    doc.save(output_path)
    print(f"Successfully generated {output_path}!")

if __name__ == '__main__':
    create_document()
