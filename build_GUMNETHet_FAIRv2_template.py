import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os
import re
import pandas as pd
import numpy as np

# Load seed42 metrics for verification
df_metrics = pd.read_csv('seed42_metrics.csv')

OMML_EQUATIONS = {
    '1': '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:sSub><m:e><m:r><m:t>R</m:t></m:r></m:e><m:sub><m:r><m:t>t→t+h, c</m:t></m:r></m:sub></m:sSub><m:r><m:t> = ln(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t+h, c</m:t></m:r></m:sub></m:sSub><m:r><m:t>) − ln(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t, c</m:t></m:r></m:sub></m:sSub><m:r><m:t>)</m:t></m:r></m:oMath>',
    '2': '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:sSub><m:e><m:r><m:t>P̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h, c</m:t></m:r></m:sub></m:sSub><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t, c</m:t></m:r></m:sub></m:sSub><m:r><m:t> · exp(</m:t></m:r><m:sSub><m:e><m:r><m:t>R̂</m:t></m:r></m:e><m:sub><m:r><m:t>t→t+h, c</m:t></m:r></m:sub></m:sSub><m:r><m:t>)</m:t></m:r></m:oMath>',
    '3': '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:sSub><m:e><m:r><m:t>out</m:t></m:r></m:e><m:sub><m:r><m:t>k</m:t></m:r></m:sub></m:sSub><m:r><m:t> = ReLU(</m:t></m:r><m:sSub><m:e><m:r><m:t>Conv1D</m:t></m:r></m:e><m:sub><m:r><m:t>k</m:t></m:r></m:sub></m:sSub><m:r><m:t>(</m:t></m:r><m:sSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sup><m:r><m:t>CNN</m:t></m:r></m:sup></m:sSup><m:r><m:t>)),   k ∈ {3, 7, 15}</m:t></m:r></m:oMath>',
    '4': '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:sSub><m:e><m:r><m:t>h</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub></m:sSub><m:r><m:t> = GRU(</m:t></m:r><m:sSubSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub><m:sup><m:r><m:t>GRU</m:t></m:r></m:sup></m:sSubSup><m:r><m:t>, </m:t></m:r><m:sSub><m:e><m:r><m:t>h</m:t></m:r></m:e><m:sub><m:r><m:t>t-1</m:t></m:r></m:sub></m:sSub><m:r><m:t>),   </m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>GRU</m:t></m:r></m:sub></m:sSub><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>h</m:t></m:r></m:e><m:sub><m:r><m:t>L</m:t></m:r></m:sub></m:sSub><m:r><m:t> ∈ </m:t></m:r><m:sSup><m:e><m:r><m:t>ℝ</m:t></m:r></m:e><m:sup><m:r><m:t>d</m:t></m:r></m:sup></m:sSup></m:oMath>',
    '5': '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>z = </m:t></m:r><m:f><m:num><m:sSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sup><m:r><m:t>KAN</m:t></m:r></m:sup></m:sSup><m:r><m:t> − t</m:t></m:r></m:num><m:den><m:r><m:t>|s| + </m:t></m:r><m:sSup><m:e><m:r><m:t>10</m:t></m:r></m:e><m:sup><m:r><m:t>-4</m:t></m:r></m:sup></m:sSup></m:den></m:f><m:r><m:t>,   ψ(z) = (1 − </m:t></m:r><m:sSup><m:e><m:r><m:t>z</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup><m:r><m:t>) · </m:t></m:r><m:sSup><m:e><m:r><m:t>e</m:t></m:r></m:e><m:sup><m:r><m:t>−0.5z²</m:t></m:r></m:sup></m:sSup></m:oMath>',
    '6': '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:sSub><m:e><m:r><m:t>g</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> = MLP([</m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>CNN</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>GRU</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>KAN</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>Pos</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>ctx</m:t></m:r></m:sub></m:sSub><m:r><m:t>])</m:t></m:r></m:oMath>',
    '7': '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> = Softmax(</m:t></m:r><m:sSub><m:e><m:r><m:t>g</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t>) = [</m:t></m:r><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>1</m:t></m:r></m:sub></m:sSub><m:r><m:t>, </m:t></m:r><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>2</m:t></m:r></m:sub></m:sSub><m:r><m:t>, </m:t></m:r><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>3</m:t></m:r></m:sub></m:sSub><m:sSup><m:e><m:r><m:t>]</m:t></m:r></m:e><m:sup><m:r><m:t>T</m:t></m:r></m:sup></m:sSup><m:r><m:t>,   </m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>fused</m:t></m:r></m:sub></m:sSub><m:r><m:t> = </m:t></m:r><m:sSubSup><m:e><m:r><m:t>∑</m:t></m:r></m:e><m:sub><m:r><m:t>i=1</m:t></m:r></m:sub><m:sup><m:r><m:t>3</m:t></m:r></m:sup></m:sSubSup><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub></m:oMath>',
    '8': '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:sSubSup><m:e><m:r><m:t>ŷ</m:t></m:r></m:e><m:sub><m:r><m:t>t+h, c</m:t></m:r></m:sub><m:sup><m:r><m:t>(q)</m:t></m:r></m:sup></m:sSubSup><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>Head</m:t></m:r></m:e><m:sub><m:r><m:t>q</m:t></m:r></m:sub></m:sSub><m:r><m:t>(</m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>fused</m:t></m:r></m:sub></m:sSub><m:r><m:t>) + </m:t></m:r><m:sSub><m:e><m:r><m:t>γ</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> · </m:t></m:r><m:sSubSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>t, c</m:t></m:r></m:sub><m:sup><m:r><m:t>target</m:t></m:r></m:sup></m:sSubSup></m:oMath>',
    '9': '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:sSub><m:e><m:r><m:t>ℒ</m:t></m:r></m:e><m:sub><m:r><m:t>total</m:t></m:r></m:sub></m:sSub><m:r><m:t> = </m:t></m:r><m:f><m:num><m:r><m:t>1</m:t></m:r></m:num><m:den><m:r><m:t>C · H · |𝒬|</m:t></m:r></m:den></m:f><m:sSub><m:e><m:r><m:t>∑</m:t></m:r></m:e><m:sub><m:r><m:t>c,h,q</m:t></m:r></m:sub></m:sSub><m:r><m:t> max(q(y − </m:t></m:r><m:sSup><m:e><m:r><m:t>ŷ</m:t></m:r></m:e><m:sup><m:r><m:t>(q)</m:t></m:r></m:sup></m:sSup><m:r><m:t>), (q−1)(y − </m:t></m:r><m:sSup><m:e><m:r><m:t>ŷ</m:t></m:r></m:e><m:sup><m:r><m:t>(q)</m:t></m:r></m:sup></m:sSup><m:r><m:t>)) + α </m:t></m:r><m:sSubSup><m:e><m:r><m:t>∑</m:t></m:r></m:e><m:sub><m:r><m:t>i=1</m:t></m:r></m:sub><m:sup><m:r><m:t>3</m:t></m:r></m:sup></m:sSubSup><m:r><m:t>(</m:t></m:r><m:sSub><m:e><m:r><m:t>w̄</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub><m:r><m:t> − </m:t></m:r><m:f><m:num><m:r><m:t>1</m:t></m:r></m:num><m:den><m:r><m:t>3</m:t></m:r></m:den></m:f><m:sSup><m:e><m:r><m:t>)</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup></m:oMath>',
    '10': '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>DA = </m:t></m:r><m:f><m:num><m:r><m:t>1</m:t></m:r></m:num><m:den><m:r><m:t>N</m:t></m:r></m:den></m:f><m:sSubSup><m:e><m:r><m:t>∑</m:t></m:r></m:e><m:sub><m:r><m:t>t=1</m:t></m:r></m:sub><m:sup><m:r><m:t>N</m:t></m:r></m:sup></m:sSubSup><m:r><m:t> 𝕀[sign(</m:t></m:r><m:sSub><m:e><m:r><m:t>P̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub></m:sSub><m:r><m:t> − </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub></m:sSub><m:r><m:t>) = sign(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub></m:sSub><m:r><m:t> − </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub></m:sSub><m:r><m:t>)]</m:t></m:r></m:oMath>'
}

def make_equation_run(paragraph, eq_id, eq_omml_xml):
    # Inline OMML
    paragraph._p.append(parse_xml(eq_omml_xml))
    
    # Bookmark Start & End
    bm_id = f"200{eq_id}"
    bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="eq_{eq_id}"/>')
    bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>')
    paragraph._p.append(bm_start)
    paragraph._p.append(bm_end)
    
    # Display equation number in blue
    run = paragraph.add_run(f"  ({eq_id})")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# Helper function to add hyperlinked equation reference ((X))
def add_eq_link(paragraph, eq_id, prefix=""):
    target_bm = f"eq_{eq_id}"
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{target_bm}" w:history="1">'
                          f'<w:r>'
                          f'<w:rPr>'
                          f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
                          f'<w:sz w:val="19"/>'
                          f'<w:color w:val="1A56DB"/>'
                          f'<w:u w:val="none"/>'
                          f'</w:rPr>'
                          f'<w:t>{prefix}({eq_id})</w:t>'
                          f'</w:r>'
                          f'</w:hyperlink>')
    paragraph._p.append(hyperlink)

# Helper function to add hyperlinked citation [X]
def add_citation_link(paragraph, ref_id):
    r_id = f"ref_{ref_id}"
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                          f'<w:r>'
                          f'<w:rPr>'
                          f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
                          f'<w:sz w:val="19"/>'
                          f'<w:color w:val="1A56DB"/>'
                          f'<w:u w:val="none"/>'
                          f'</w:rPr>'
                          f'<w:t>[{ref_id}]</w:t>'
                          f'</w:r>'
                          f'</w:hyperlink>')
    paragraph._p.append(hyperlink)

# Helper function to add hyperlinked Table reference (Bảng X)
def add_tbl_link(paragraph, tbl_id, display_text):
    t_id = f"tbl_{tbl_id}"
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{t_id}" w:history="1">'
                          f'<w:r>'
                          f'<w:rPr>'
                          f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
                          f'<w:sz w:val="19"/>'
                          f'<w:color w:val="1A56DB"/>'
                          f'<w:u w:val="none"/>'
                          f'</w:rPr>'
                          f'<w:t>{display_text}</w:t>'
                          f'</w:r>'
                          f'</w:hyperlink>')
    paragraph._p.append(hyperlink)

# Helper function to add hyperlinked Figure reference (Hình X)
def add_fig_link(paragraph, fig_id, display_text):
    f_id = f"fig_{fig_id}"
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{f_id}" w:history="1">'
                          f'<w:r>'
                          f'<w:rPr>'
                          f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
                          f'<w:sz w:val="19"/>'
                          f'<w:color w:val="1A56DB"/>'
                          f'<w:u w:val="none"/>'
                          f'</w:rPr>'
                          f'<w:t>{display_text}</w:t>'
                          f'</w:r>'
                          f'</w:hyperlink>')
    paragraph._p.append(hyperlink)

def add_inline_math(paragraph, inner_xml):
    omml_str = f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{inner_xml}</m:oMath>'
    paragraph._p.append(parse_xml(omml_str))

def _emit_string_with_links_vn(paragraph, text_str):
    # Regex to split on Bảng X, Hình X, and Equation (X)
    parts = re.split(r'(Bảng \d+|Hình \d+|\(\d+\))', text_str)
    for part in parts:
        if not part:
            continue
        m_tbl = re.match(r'Bảng (\d+)', part)
        m_fig = re.match(r'Hình (\d+)', part)
        m_eq = re.match(r'\((\d+)\)', part)
        if m_tbl:
            add_tbl_link(paragraph, m_tbl.group(1), part)
        elif m_fig:
            add_fig_link(paragraph, m_fig.group(1), part)
        elif m_eq and int(m_eq.group(1)) <= 10:
            add_eq_link(paragraph, m_eq.group(1), prefix="")
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9.5)

def add_text_with_citations(paragraph, text_segments, justify=True):
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg in text_segments:
        if isinstance(seg, str):
            _emit_string_with_links_vn(paragraph, seg)
        elif isinstance(seg, int):
            add_citation_link(paragraph, seg)
        elif isinstance(seg, tuple) and len(seg) == 2 and seg[0] == 'eq':
            add_eq_link(paragraph, seg[1], prefix="")
        elif isinstance(seg, tuple) and len(seg) == 3 and seg[0] == 'eq':
            add_eq_link(paragraph, seg[1], prefix=seg[2])
        elif isinstance(seg, tuple) and len(seg) == 2 and seg[0] == 'tbl':
            add_tbl_link(paragraph, seg[1], f"Bảng {seg[1]}")
        elif isinstance(seg, tuple) and len(seg) == 2 and seg[0] == 'fig':
            add_fig_link(paragraph, seg[1], f"Hình {seg[1]}")
        elif isinstance(seg, tuple) and len(seg) == 2 and seg[0] == 'math':
            add_inline_math(paragraph, seg[1])
        elif isinstance(seg, list):
            for idx, c_id in enumerate(seg):
                add_citation_link(paragraph, c_id)
                if idx < len(seg) - 1:
                    r = paragraph.add_run(', ')
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(9.5)

def build_v2_conference_document(output_filename):
    # Use template as base to preserve exact styles, section margins, header/footer
    doc = docx.Document('GUMNetHet_FAIR_ban_thao_template.docx')
    
    # Helper to set 100% pure white cell formatting with standard 3-line borders
    def style_table(table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = table._tbl.tblPr
        
        # Remove old table borders if any
        for old_b in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old_b)
            
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                            f'<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                            f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
                            f'<w:insideV w:val="none"/>'
                            f'<w:left w:val="none"/>'
                            f'<w:right w:val="none"/>'
                            f'</w:tblBorders>')
        tblPr.append(borders)
        
        for row_idx, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            if row_idx == 0:
                trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
                
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                
                # Remove ALL existing shading elements
                for old_shd in tcPr.findall(qn('w:shd')):
                    tcPr.remove(old_shd)
                    
                # 100% Pure white cell shading
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="FFFFFF"/>')
                tcPr.append(shd)
                
                # Margins
                mar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                                f'<w:top w:w="80" w:type="dxa"/>'
                                f'<w:bottom w:w="80" w:type="dxa"/>'
                                f'<w:left w:w="100" w:type="dxa"/>'
                                f'<w:right w:w="100" w:type="dxa"/>'
                                f'</w:tcMar>')
                tcPr.append(mar)
                
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.05
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(8.5)
                        r.bold = False

    # Apply styling to all tables in doc
    for t in doc.tables:
        style_table(t)

    # Helper function to add table caption with bookmark
    def make_table_caption(p, text, tbl_id):
        p.style = 'Caption'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.text = ""
        bm_id = f"300{tbl_id}"
        bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="tbl_{tbl_id}"/>')
        bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>')
        p._p.append(bm_start)
        p._p.append(bm_end)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8.5)
        run.bold = False

    # Helper function to add figure caption with bookmark
    def make_figure_caption(p, text, fig_id):
        p.style = 'Caption'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        p.text = ""
        bm_id = f"400{fig_id}"
        bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="fig_{fig_id}"/>')
        bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>')
        p._p.append(bm_start)
        p._p.append(bm_end)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8.0)
        run.bold = False

    # 1. Update Section 1: GIỚI THIỆU (Paragraphs 8, 9, 10) to incorporate Reviewer Comments
    doc.paragraphs[8].text = ""
    add_text_with_citations(doc.paragraphs[8], [
        "Thị trường xăng dầu Việt Nam có đặc thù cấu trúc nguồn cung mang tính chiến lược: khoảng 70% tổng sản lượng "
        "tiêu thụ nội địa được cung ứng bởi hai nhà máy lọc hóa dầu trong nước (Dung Quất và Nghi Sơn), trong khi 30% "
        "nhu cầu còn lại bắt buộc phải nhập khẩu trực tiếp từ các thị trường quốc tế. Trong đó, thị trường Singapore là "
        "địa bàn nhập khẩu trọng yếu nhất, với giá giao dịch thành phẩm Mean of Platts Singapore (MOPS)—tiêu biểu là Mogas 95 "
        "(MG95) và Gasoil 0.001%S (DO 0.001%)—đóng vai trò là hệ quy chiếu định giá cơ sở cho mọi hợp đồng thương mại, "
        "tính toán giá vốn và quản lý chuỗi cung ứng. Tuy nhiên, giá Platts biến động cực kỳ phức tạp do phản ứng đồng thời "
        "với các cú sốc cung–cầu toàn cầu, biến động tỷ giá và rủi ro địa chính trị (GPR) ",
        [1, 2, 3],
        ". Các giai đoạn sụp đổ giá dầu 2014–2016, chiến tranh giá OPEC+ năm 2020, xung đột Nga–Ukraine 2022 và căng thẳng Biển Đỏ 2023–2024 cho thấy chuỗi giá "
        "thành phẩm thường xuyên xuất hiện bước nhảy phi tuyến, dịch chuyển chế độ và phân cụm biến động mạnh. Do đó, nhu cầu "
        "dự báo chính xác giá xăng dầu Platts trong ngắn hạn (H1–H7) và trung hạn (H10–H60) là đòi hỏi cấp thiết phục vụ trực "
        "tiếp cho việc ra quyết định kinh doanh, tối ưu hóa kế hoạch mua hàng, quản trị tồn kho và phòng hộ rủi ro (hedging) "
        "của các doanh nghiệp đầu mối xăng dầu lớn như Petrolimex và PVOIL."
    ])
    
    doc.paragraphs[9].text = ""
    add_text_with_citations(doc.paragraphs[9], [
        "Các kiến trúc chuỗi thời gian hiện đại như PatchTST ", 8,
        ", iTransformer ", 9,
        ", TimesNet ", 10,
        ", DLinear ", 11,
        ", Mamba ", 15,
        " và Chronos ", 16,
        " đã cải thiện đáng kể hiệu năng trên các benchmark tổng quát, nhưng phần lớn đều xử lý toàn bộ các biến đầu vào "
        "trong một không gian biểu diễn tương đối đồng nhất. Với dữ liệu năng lượng, ba nhóm tín hiệu có bản chất cơ bản khác nhau: "
        "động lượng giá tần số cao, trạng thái vĩ mô biến đổi chậm và phản ứng phi tuyến nhạy cú sốc biên độ lớn đòi hỏi các inductive "
        "bias chuyên biệt. Các mạng kết hợp chuyên gia (MoE) truyền thống ",
        [20, 21, 22],
        " thường đưa cùng một tập đặc trưng tới mọi expert, làm suy giảm mức độ chuyên môn hóa."
    ])
    
    doc.paragraphs[10].text = ""
    add_text_with_citations(doc.paragraphs[10], [
        "GUMNetHet giải quyết triệt để điểm nghẽn này bằng kỹ thuật phân vùng đặc trưng (feature partitioning): nhóm giá và benchmark "
        "được xử lý bởi CNN-1D đa tỷ lệ; nhóm vĩ mô và chỉ số GPR bởi GRU-Attention; nhóm tỷ lệ crack-spread và độ biến động bởi Wavelet-KAN. "
        "Ba biểu diễn được hợp nhất linh hoạt thông qua một bộ định tuyến (gating router) phụ thuộc vào horizon dự báo và ngữ cảnh thị trường. "
        "Đóng góp chính của bài báo gồm: (i) Kiến trúc MoE dị thể với phân vùng đặc trưng theo bản chất kinh tế và miền tần số; (ii) Bộ định tuyến "
        "nhận biết horizon (horizon-aware routing); (iii) Đầu ra đa phân vị (multi-quantile head) kết hợp residual scaling để kiểm soát độ trôi phương sai; "
        "và (iv) Đánh giá thực nghiệm mở rộng walk-forward trên N=4.512 ngày giao dịch với bảng kết quả đầy đủ các chỉ số MAE, RMSE, MAPE, R² và DA%."
    ])

    # Abstract with math notation
    doc.paragraphs[5].text = ""
    add_text_with_citations(doc.paragraphs[5], [
        "Tóm tắt—Dự báo giá xăng dầu thành phẩm trở nên khó khăn khi chuỗi giá chịu đồng thời biến động ngắn hạn, dịch chuyển chế độ vĩ mô "
        "và các cú sốc địa chính trị có tính phi tuyến, đuôi dày. Bài báo đề xuất GUMNetHet (Heterogeneous Gated Unified Mixture Network), "
        "một mạng kết hợp chuyên gia không đồng nhất cho dự báo xác suất đa chu kỳ. Mô hình phân vùng đặc trưng thành ba nhóm và giao cho "
        "ba chuyên gia chuyên biệt: CNN-1D đa tỷ lệ cho động lượng giá, GRU-Attention cho chế độ vĩ mô, và Wavelet-KAN cho quan hệ phi tuyến "
        "nhạy cú sốc. Bộ định tuyến nhận biết horizon kết hợp biểu diễn chuyên gia, nhúng horizon và thống kê ngữ cảnh để phân bổ trọng số động; "
        "đầu ra đa phân vị ",
        ('math', '<m:r><m:t>q ∈ {0.1, 0.5, 0.9}</m:t></m:r>'),
        " được tối ưu bằng pinball loss và điều chuẩn cân bằng tải. Trên dữ liệu đa nguồn 11/2008–04/2026 "
        "(N=4.512 quan sát) với expanding walk-forward, GUMNetHet đạt MAE thấp nhất trong nhóm baseline được báo cáo ở cả bảy horizon H1–H60 cho "
        "MG95 và DO 0.001%. Tại H60, MAE giảm 30,1% cho xăng và 22,9% cho dầu so với baseline tốt nhất tương ứng. Directional accuracy đạt "
        "90,95–95,56% ở H1–H7 của xăng và 76,65–84,92% ở H1–H7 của dầu, nhưng giảm mạnh tại H10/H60, cho thấy cần phân biệt giữa độ chính xác "
        "mức giá và độ chính xác hướng ở horizon dài. Khoảng phân vị 80% đạt PICP=82,4% với PINAW=0,142. Với 0,34M tham số và độ trễ 1,42 ms/mẫu, "
        "GUMNetHet cho thấy sự cân bằng tốt giữa độ chính xác, bất định và chi phí tính toán."
    ])

    # Section 3.1: Problem Formulation & Partitioning with Equations (1) and (2)
    doc.paragraphs[16].text = ""
    add_text_with_citations(doc.paragraphs[16], [
        "Với giá mục tiêu ", ('math', '<m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t,c</m:t></m:r></m:sub></m:sSub>'),
        " của sản phẩm ", ('math', '<m:r><m:t>c ∈ {MG95, DO 0.001%}</m:t></m:r>'),
        ", lookback ", ('math', '<m:r><m:t>L = 30</m:t></m:r>'),
        " và horizon ", ('math', '<m:r><m:t>h ∈ {1, 3, 5, 7, 10, 20, 60}</m:t></m:r>'),
        ", mục tiêu học là ánh xạ từ cửa sổ lịch sử đa biến ", ('math', '<m:sSub><m:e><m:r><m:t>X</m:t></m:r></m:e><m:sub><m:r><m:t>t-L+1:t</m:t></m:r></m:sub></m:sSub>'),
        " sang log-return tích lũy đa chu kỳ ", ('math', '<m:sSub><m:e><m:r><m:t>R</m:t></m:r></m:e><m:sub><m:r><m:t>t→t+h, c</m:t></m:r></m:sub></m:sSub>'),
        " theo công thức ", ('eq', '1'), ": "
    ])
    make_equation_run(doc.paragraphs[16], '1', OMML_EQUATIONS['1'])
    
    p_eq1_cont = doc.paragraphs[16]
    p_eq1_cont.add_run(". Mức giá dự báo tất định tương ứng được khôi phục trực tiếp theo ")
    add_eq_link(p_eq1_cont, '2', prefix="")
    p_eq1_cont.add_run(": ")
    make_equation_run(p_eq1_cont, '2', OMML_EQUATIONS['2'])
    p_eq1_cont.add_run(". Tập đặc trưng được phân chia chặt chẽ thành ba không gian con dị thể: ")
    p_eq1_cont.add_run("(i) Nhóm giá và benchmark (P_t, WTI, Brent) đưa vào CNN; (ii) Nhóm vĩ mô và chỉ số GPR đưa vào GRU; (iii) Nhóm tỷ lệ crack-spread và độ biến động đưa vào Wavelet-KAN.")

    # Section 3.2: 3 Experts and Router with Equations (3), (4), (5), (6), (7)
    doc.paragraphs[18].text = ""
    add_text_with_citations(doc.paragraphs[18], [
        "Expert giá sử dụng ba nhánh Conv1D đa tỷ lệ với kích thước kernel ",
        ('math', '<m:r><m:t>k ∈ {3, 7, 15}</m:t></m:r>'),
        " kết hợp attention pooling ", ('eq', '3'), ": "
    ])
    make_equation_run(doc.paragraphs[18], '3', OMML_EQUATIONS['3'])
    doc.paragraphs[18].add_run(". Expert vĩ mô dùng GRU 2 lớp với dropout 0.1 trích xuất vector trạng thái ")
    add_eq_link(doc.paragraphs[18], '4', prefix="")
    doc.paragraphs[18].add_run(": ")
    make_equation_run(doc.paragraphs[18], '4', OMML_EQUATIONS['4'])
    doc.paragraphs[18].add_run(". Expert cú sốc dùng Wavelet-KAN với hàm cơ sở Mexican Hat wavelet ")
    add_eq_link(doc.paragraphs[18], '5', prefix="")
    doc.paragraphs[18].add_run(": ")
    make_equation_run(doc.paragraphs[18], '5', OMML_EQUATIONS['5'])
    doc.paragraphs[18].add_run(". Router nhận biết horizon kết hợp cả 3 biểu diễn chuyên gia cùng vector nhúng horizon ")
    add_eq_link(doc.paragraphs[18], '6', prefix="")
    doc.paragraphs[18].add_run(" để sinh trọng số softmax ")
    add_eq_link(doc.paragraphs[18], '7', prefix="")
    doc.paragraphs[18].add_run(": ")
    make_equation_run(doc.paragraphs[18], '6', OMML_EQUATIONS['6'])
    doc.paragraphs[18].add_run(". Trọng số định tuyến tổng hợp biểu diễn: ")
    make_equation_run(doc.paragraphs[18], '7', OMML_EQUATIONS['7'])

    # Section 3.3: Quantile Head and Loss Function with Equations (8) and (9)
    doc.paragraphs[22].text = ""
    add_text_with_citations(doc.paragraphs[22], [
        "Head sinh ba phân vị xác suất ",
        ('math', '<m:r><m:t>q ∈ {0.1, 0.5, 0.9}</m:t></m:r>'),
        ". Kỹ thuật Residual scaling học hệ số neo ",
        ('math', '<m:r><m:t>γ_h ∈ (0, 1)</m:t></m:r>'),
        " để giới hạn độ trôi phương sai ở horizon dài theo ", ('eq', '8'), ": "
    ])
    make_equation_run(doc.paragraphs[22], '8', OMML_EQUATIONS['8'])
    doc.paragraphs[22].add_run(". Mô hình được tối ưu bằng hàm mất mát tổng hợp Pinball loss kết hợp điều chuẩn cân bằng tải router theo ")
    add_eq_link(doc.paragraphs[22], '9', prefix="")
    doc.paragraphs[22].add_run(": ")
    make_equation_run(doc.paragraphs[22], '9', OMML_EQUATIONS['9'])

    # Ensure Data section has exact N=4.512
    data_p25 = (
        "Dữ liệu bao phủ 03/11/2008–30/04/2026 với N=4.512 quan sát ngày giao dịch. Hai target được báo cáo là MG95 và DO 0.001% (USD/thùng). "
        "Tập biến ngoại sinh gồm WTI, Brent DTD, chênh lệch liên sản phẩm (MG92, MG97, KERO, FO 180, Naphtha), chỉ số GPR [1], DXY, tỷ lệ crack spread "
        "và độ biến động thực tế. Lookback cố định L=30 ngày; horizon dự báo gồm h ∈ {1, 3, 5, 7, 10, 20, 60}. Giao thức expanding walk-forward chia tập ban đầu "
        "70% train, 10% validation và 20% test; sau mỗi bước cuộn, cửa sổ train được mở rộng để phản ánh điều kiện vận hành thực tế."
    )
    doc.paragraphs[25].text = data_p25

    # Section 4.2: Diagnostics
    doc.paragraphs[28].text = ""
    add_text_with_citations(doc.paragraphs[28], [
        "Kiểm định ADF và KPSS trong ", ('tbl', '1'),
        " cho bằng chứng về tính dừng ở mức giá; sau chuyển log-return, kiểm định ADF bác bỏ giả thuyết nghiệm đơn vị (p < 0.001) trên tất cả các chuỗi. Độ nhọn (Kurtosis) của lợi suất rất cao (đạt 213,35 ở WTI và 17,51 ở MG95), khẳng định tính chất đuôi dày và các cú sốc cực đoan trong chuỗi giá năng lượng."
    ])

    # Section 4.3: Baseline and Metrics with Equation (10)
    doc.paragraphs[32].text = ""
    add_text_with_citations(doc.paragraphs[32], [
        "Các baseline đối chuẩn được báo cáo chi tiết gồm 6 mô hình đại diện tiêu biểu: PatchTST ", 8,
        ", iTransformer ", 9,
        ", TimesNet ", 10,
        ", DLinear ", 11,
        ", Chronos ", 16,
        " và BiMamba (ở bảng xăng của nguồn thực nghiệm). Các mô hình dùng chung lookback L=30, quy tắc trễ dữ liệu và giao thức "
        "walk-forward expanding. Các chỉ số đánh giá điểm gồm MAE, RMSE, MAPE (%) và Hệ số xác định R². Độ chính xác xu hướng được tính theo ",
        ('eq', '10'), ": "
    ])
    make_equation_run(doc.paragraphs[32], '10', OMML_EQUATIONS['10'])
    doc.paragraphs[32].add_run(". Đánh giá phân phối xác suất sử dụng hệ số bao phủ PICP và độ rộng dải chuẩn hóa PINAW.")

    # Section 5.1: Empirical multi-horizon performance
    doc.paragraphs[36].text = ""
    add_text_with_citations(doc.paragraphs[36], [
        ('fig', '2'), " cùng ", ('tbl', '2'), " và ", ('tbl', '3'),
        " cho thấy GUMNetHet duy trì MAE thấp hơn nhóm baseline được báo cáo ở cả hai sản phẩm, đặc biệt ở H20–H60. Ở H60, MAE của xăng là 4,847 so với 6,933 của baseline tốt nhất (giảm 30,1%); với dầu là 7,066 so với 9,167 (giảm 22,9%). Tuy nhiên, R² tại H60 giảm còn 0,155 (xăng) và −0,007 (dầu), vì vậy kết quả dài hạn nên được hiểu là ổn định sai số mức giá tốt hơn baseline, không phải dự báo quỹ đạo dài hạn hoàn hảo."
    ])

    # Section 5.2: Directional Accuracy
    doc.paragraphs[42].text = ""
    add_text_with_citations(doc.paragraphs[42], [
        "Độ chính xác hướng trong ", ('fig', '3'), " cho thấy một hành vi khác với sai số mức. Với xăng (", ('tbl', '2'),
        "), GUMNetHet đạt 91,46%, 91,37%, 90,95% và 95,56% ở H1, H3, H5, H7; với dầu (", ('tbl', '3'),
        ") tương ứng là 84,92%, 76,65%, 76,88% và 83,28%. Ở H20, DA vẫn cao (91,65% xăng; 71,11% dầu). Ngược lại, H10 và H60 giảm dưới 50%, lần lượt 42,24%/27,95% cho xăng và 32,29%/19,10% cho dầu. Do đó, GUMNetHet phù hợp hơn như mô hình dự báo mức và biên bất định ở horizon dài, thay vì công cụ sinh tín hiệu hướng."
    ])

    # Section 5.3: Quantile, Ablation, Router
    doc.paragraphs[46].text = ""
    add_text_with_citations(doc.paragraphs[46], [
        "Khoảng ", ('math', '<m:r><m:t>[q0.1, q0.9]</m:t></m:r>'),
        " đạt PICP=82,4% so với danh định 80% và PINAW=0,142. ", ('fig', '4'),
        " cho thấy biên bất định mở rộng khi biến động tăng. Ablation trong ", ('tbl', '4'),
        " cho thấy thay Wav-KAN bằng MLP gây suy giảm lớn nhất trong các biến thể expert; router đồng nhất cũng làm MAE tăng "
        "đáng kể, củng cố vai trò của chuyên môn hóa và định tuyến thích ứng."
    ])

    doc.paragraphs[48].text = ""
    add_text_with_citations(doc.paragraphs[48], [
        "Bản revised còn ghi nhận rằng loại bỏ residual scaling làm MAE tăng khoảng 8,5%/6,3% ở H20 và 14,1%/11,8% ở H60 cho xăng/dầu. Phân tích router trong ",
        ('fig', '4'),
        " cho thấy ở GPR thấp và horizon ngắn, CNN có trọng số trung bình khoảng 0,48; khi GPR vượt phân vị 90%, trọng số Wav-KAN tăng từ khoảng 0,29 lên 0,61 ở horizon trung–dài và CNN giảm xuống khoảng 0,21. Các kết quả này cho thấy router thực sự thay đổi chế độ thay vì chỉ trung bình hóa đầu ra."
    ])

    # Section 5.4: Computational Cost
    doc.paragraphs[53].text = ""
    add_text_with_citations(doc.paragraphs[53], [
        "Kết quả trong ", ('tbl', '5'),
        " cho thấy GUMNetHet không nhẹ bằng DLinear, nhưng nhỏ hơn đáng kể các Transformer đại diện và vẫn đạt độ trễ 1,42 ms/mẫu. Vì vậy lợi thế của mô hình là hiệu quả tương đối cao trong nhóm phi tuyến mạnh, phù hợp cho hệ thống giám sát và cập nhật dự báo cận thời gian thực."
    ])

    # Refine Table captions and references to sequential numbering: Bảng 1, Bảng 2, Bảng 3, Bảng 4, Bảng 5
    make_table_caption(doc.paragraphs[29], "Bảng 1. Chẩn đoán thống kê rút gọn của các chuỗi chính.", "1")
    make_table_caption(doc.paragraphs[39], "Bảng 2. Kết quả chi tiết trên MG95 (Seed=42). MAE/RMSE tính theo USD/thùng; giá trị DA là phần trăm.", "2")
    make_table_caption(doc.paragraphs[40], "Bảng 3. Kết quả chi tiết trên DO 0.001% (Seed=42). Bảng nguồn thực nghiệm không báo cáo BiMamba cho mục tiêu dầu.", "3")
    make_table_caption(doc.paragraphs[47], "Bảng 4. Ablation rút gọn của GUMNetHet (Seed=42).", "4")
    make_table_caption(doc.paragraphs[52], "Bảng 5. Chi phí tính toán trên GPU Tesla T4.", "5")

    # Refine Figure captions
    make_figure_caption(doc.paragraphs[20], "Hình 1. Kiến trúc GUMNetHet: phân vùng đặc trưng, ba expert dị thể, horizon-aware router và multi-quantile head.", "1")
    make_figure_caption(doc.paragraphs[38], "Hình 2. Đường cong MAE và R² qua bảy horizon cho MG95 và DO 0.001% (Seed=42).", "2")
    make_figure_caption(doc.paragraphs[44], "Hình 3. Directional accuracy (DA%) của GUMNetHet và các baseline qua H1–H60 cho xăng và dầu.", "3")
    make_figure_caption(doc.paragraphs[50], "Hình 4. Trên: fan chart đa phân vị dưới biến động mạnh. Dưới: trọng số router trong chế độ GPR thấp và GPR cao.", "4")

    # Apply font consistency to all paragraphs
    for p in doc.paragraphs:
        if p.style.name == 'Normal':
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9.5)
        elif p.style.name == 'Abstract':
            p.paragraph_format.line_spacing = 1.10
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9.0)
        elif p.style.name == 'Refs':
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(8.0)

    # Save document
    doc.save(output_filename)
    print(f"Successfully generated {output_filename} with full OMML equations and bookmarks!")

if __name__ == '__main__':
    # Generate all requested v2 filenames
    build_v2_conference_document('GUMNETHet_FAIRv2_template.docx')
    build_v2_conference_document('GUMNETHET_FAIR_v2_TIENG_VIET.docx')
    build_v2_conference_document('GUMNETHET_FAIR_v2.docx')
