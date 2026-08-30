import os
import json
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    root = 'results_v4/walkforward'
    seeds = [42, 123, 777, 2025, 9999]
    targets = ['XANG', 'DAU']
    horizons = [1, 3, 5, 7, 10, 20, 60]

    records = []
    
    # Scan all models
    if os.path.exists(root):
        for model in os.listdir(root):
            model_path = os.path.join(root, model)
            if os.path.isdir(model_path):
                for h in horizons:
                    for seed in seeds:
                        for target in targets:
                            run_dir = f'{target}_H{h}_seed{seed}'
                            run_path = os.path.join(model_path, run_dir)
                            results_json = os.path.join(run_path, 'results.json')
                            if os.path.exists(results_json):
                                try:
                                    with open(results_json) as f:
                                        d = json.load(f)
                                        m_dict = d.get('metrics', {})
                                        
                                        crps_val = m_dict.get('crps', m_dict.get('CRPS'))
                                        mql_val = m_dict.get('mql', m_dict.get('MQL'))
                                        
                                        records.append({
                                            'model': d.get('model', model),
                                            'target': d.get('target_type', target),
                                            'horizon': h,
                                            'seed': d.get('seed', seed),
                                            'MAE': m_dict.get('MAE'),
                                            'MSE': m_dict.get('MSE'),
                                            'RMSE': m_dict.get('RMSE'),
                                            'MAPE': m_dict.get('MAPE'),
                                            'SMAPE': m_dict.get('SMAPE'),
                                            'R2': m_dict.get('R2'),
                                            'CRPS': crps_val,
                                            'MQL': mql_val,
                                            'PICP': m_dict.get('PICP'),
                                            'PINAW': m_dict.get('PINAW')
                                        })
                                except Exception as e:
                                    pass

    df = pd.DataFrame(records)
    if len(df) == 0:
        print("No records found!")
        return

    # Initialize Word Document
    doc = Document()
    
    # Title
    doc.add_heading('BẢNG KẾT QUẢ THỰC NGHIỆM TRUNG BÌNH CỘNG 5 SEEDS', level=1)
    doc.add_paragraph('Trạng thái: Báo cáo đầy đủ cho tất cả các khung H (1, 3, 5, 7, 10, 20, 60).')
    doc.add_paragraph(f"Ngày tổng hợp: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    metrics_cols = ['MAE', 'MSE', 'RMSE', 'MAPE', 'SMAPE', 'R2', 'CRPS', 'MQL', 'PICP', 'PINAW']
    
    for h in horizons:
        doc.add_heading(f'KẾT QUẢ KHUNG H{h}', level=2)
        
        for target in targets:
            df_t = df[(df['target'] == target) & (df['horizon'] == h)].copy()
            if len(df_t) == 0:
                continue
                
            df_avg = df_t.groupby('model')[metrics_cols].mean().reset_index()
            sort_col = 'CRPS' if df_avg['CRPS'].notna().any() else 'MAE'
            df_avg = df_avg.sort_values(by=sort_col, ascending=True).reset_index(drop=True)
            
            best_vals = {
                'MAE': df_avg['MAE'].min(),
                'MSE': df_avg['MSE'].min() if df_avg['MSE'].notna().any() else None,
                'RMSE': df_avg['RMSE'].min(),
                'MAPE': df_avg['MAPE'].min(),
                'SMAPE': df_avg['SMAPE'].min() if df_avg['SMAPE'].notna().any() else None,
                'R2': df_avg['R2'].max(),
                'CRPS': df_avg['CRPS'].min() if df_avg['CRPS'].notna().any() else None,
                'MQL': df_avg['MQL'].min() if df_avg['MQL'].notna().any() else None,
                'PICP': df_avg['PICP'].max() if df_avg['PICP'].notna().any() else None,
                'PINAW': df_avg['PINAW'].min() if df_avg['PINAW'].notna().any() else None,
            }
            
            # Subheading (No icons)
            target_name = "XĂNG (GASOLINE)" if target == 'XANG' else "DẦU (DIESEL)"
            doc.add_heading(f"{target_name} - KHUNG H{h}", level=3)
            
            # Create Table
            headers = ['STT', 'Mô hình (Model)'] + metrics_cols
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Add headers
            hdr_cells = table.rows[0].cells
            for i, header_name in enumerate(headers):
                hdr_cells[i].text = header_name
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                
            for idx, row in df_avg.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx + 1)
                
                model = row['model']
                is_gumnet = 'GUMNet' in model
                
                p = row_cells[1].paragraphs[0]
                run = p.add_run(model)
                if is_gumnet:
                    run.bold = True
                
                for i, m in enumerate(metrics_cols):
                    cell_idx = i + 2
                    val = row[m]
                    if pd.isna(val):
                        row_cells[cell_idx].text = 'nan'
                        continue
                    
                    is_best = False
                    if m in ['R2', 'PICP']:
                        is_best = (val >= best_vals[m] - 1e-6)
                    else:
                        is_best = (val <= best_vals[m] + 1e-6)
                    
                    if m == 'MAPE':
                        val_str = f"{val:.2f}%"
                    else:
                        val_str = f"{val:.4f}"
                        
                    p2 = row_cells[cell_idx].paragraphs[0]
                    run2 = p2.add_run(val_str)
                    if is_best:
                        run2.bold = True
            
            doc.add_paragraph() # Spacing
            
    output_path = 'results_v4/COMPLETE_ALL_HORIZONS_TABLES_5SEEDS_AVERAGE.docx'
    doc.save(output_path)
    print(f"Saved {output_path}")

if __name__ == '__main__':
    main()
