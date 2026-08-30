import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os
import sys
import re
import zipfile
import io
from copy import deepcopy

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

from scripts.table_builder_helper import M, create_inline_omml, style_table

def convert_strict_to_transitional(in_path, out_path):
    mappings = [
        (b'http://purl.oclc.org/ooxml/wordprocessingml/main', b'http://schemas.openxmlformats.org/wordprocessingml/2006/main'),
        (b'http://purl.oclc.org/ooxml/drawingml/main', b'http://schemas.openxmlformats.org/drawingml/2006/main'),
        (b'http://purl.oclc.org/ooxml/drawingml/wordprocessingDrawing', b'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'),
        (b'http://purl.oclc.org/ooxml/officeDocument/relationships/', b'http://schemas.openxmlformats.org/officeDocument/2006/relationships/'),
        (b'http://purl.oclc.org/ooxml/package/relationships/', b'http://schemas.openxmlformats.org/package/2006/relationships/'),
        (b'http://purl.oclc.org/ooxml/schemaLibrary/main', b'http://schemas.openxmlformats.org/schemaLibrary/2006/main'),
        (b'http://purl.oclc.org/ooxml/officeDocument/customXml', b'http://schemas.openxmlformats.org/officeDocument/2006/customXml'),
        (b'http://purl.oclc.org/ooxml/officeDocument/extendedProperties', b'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'),
        (b'http://purl.oclc.org/ooxml/officeDocument/math', b'http://schemas.openxmlformats.org/officeDocument/2006/math')
    ]
    with zipfile.ZipFile(in_path, 'r') as z_in:
        with zipfile.ZipFile(out_path, 'w', zipfile.ZIP_DEFLATED) as z_out:
            for item in z_in.infolist():
                data = z_in.read(item.filename)
                for old_ns, new_ns in mappings:
                    data = data.replace(old_ns, new_ns)
                z_out.writestr(item, data)

def add_styled_para_ieee(p, segments, style_name='Body Text', align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3, line_spacing=1.05):
    p.text = ""
    try:
        p.style = style_name
    except Exception:
        pass
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    
    for seg in segments:
        if isinstance(seg, str):
            parts = re.split(r'(Table \d+|Fig\. \d+|\[\d+\])', seg)
            for part in parts:
                if not part:
                    continue
                m_tbl = re.match(r'Table (\d+)', part)
                m_fig = re.match(r'Fig\. (\d+)', part)
                m_ref = re.match(r'\[(\d+)\]', part)
                if m_tbl:
                    t_id = f"tbl_{m_tbl.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{t_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="20"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                elif m_fig:
                    f_id = f"fig_{m_fig.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{f_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="20"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                elif m_ref:
                    r_id = f"ref_{m_ref.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="20"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                else:
                    run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10.0 if style_name == 'Body Text' else 9.0)
                    run.font.bold = False
        elif isinstance(seg, tuple) and seg[0] == 'm':
            m_xml = M[seg[1]]
            p._p.append(parse_xml(create_inline_omml(m_xml)))
        elif isinstance(seg, tuple) and seg[0] == 'ref':
            r_id = f"ref_{seg[1]}"
            hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                           f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="20"/><w:color w:val="1A56DB"/></w:rPr>'
                           f'<w:t>[{seg[1]}]</w:t></w:r></w:hyperlink>')
            p._p.append(hl)

def set_ieee_heading1(p, text):
    p.text = ""
    try:
        p.style = 'Heading 1'
    except Exception:
        pass
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    pPr = p._p.get_or_add_pPr()
    existing_lvl = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
    if existing_lvl is None:
        pPr.append(parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="0"/>'))
    else:
        existing_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "0")
        
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.0)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_ieee_heading2(p, text):
    p.text = ""
    try:
        p.style = 'Heading 2'
    except Exception:
        pass
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    
    pPr = p._p.get_or_add_pPr()
    existing_lvl = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
    if existing_lvl is None:
        pPr.append(parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="1"/>'))
    else:
        existing_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "1")
        
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.0)
    run.font.italic = True
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_ieee_heading3(p, text):
    p.text = ""
    try:
        p.style = 'Heading 3'
    except Exception:
        pass
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    
    pPr = p._p.get_or_add_pPr()
    existing_lvl = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
    if existing_lvl is None:
        pPr.append(parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="2"/>'))
    else:
        existing_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "2")
        
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_ieee_caption_fig(p, tag_num, prefix, text):
    p.text = ""
    try:
        p.style = 'figure caption'
    except Exception:
        pass
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    r_pre = p.add_run(prefix)
    r_pre.font.name = "Times New Roman"
    r_pre.font.size = Pt(8.5)
    r_pre.font.bold = True
    
    r_txt = p.add_run(" " + text)
    r_txt.font.name = "Times New Roman"
    r_txt.font.size = Pt(8.5)
    r_txt.font.bold = False
    
    bm_name = f"fig_{tag_num}"
    bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{tag_num}" w:name="{bm_name}"/>')
    bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{tag_num}"/>')
    p._p.insert(0, bm_start)
    p._p.append(bm_end)

def set_ieee_caption_table(p, tag_num, prefix, text):
    p.text = ""
    try:
        p.style = 'table head'
    except Exception:
        pass
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    
    r_pre = p.add_run(prefix)
    r_pre.font.name = "Times New Roman"
    r_pre.font.size = Pt(8.5)
    r_pre.font.bold = True
    
    r_txt = p.add_run(" " + text)
    r_txt.font.name = "Times New Roman"
    r_txt.font.size = Pt(8.5)
    r_txt.font.bold = False
    
    bm_name = f"tbl_{tag_num}"
    bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{tag_num}" w:name="{bm_name}"/>')
    bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{tag_num}"/>')
    p._p.insert(0, bm_start)
    p._p.append(bm_end)

def build_v4_final():
    tmpl_strict = 'conference-template-a4 (1).docx'
    tmpl_trans = 'conference-template-a4_transitional.docx'
    convert_strict_to_transitional(tmpl_strict, tmpl_trans)
    
    doc = docx.Document(tmpl_trans)
    src_doc = docx.Document('GUMNETHet_FAIRv3_EN.docx')
    
    # 1. Set Title
    p_title = doc.paragraphs[0]
    p_title.text = "Robust Probabilistic Energy Forecasting under Geopolitical Shocks: An Adaptive Mixture of Local-Global Experts"
    p_title.style = 'paper title'
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_title.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(24.0)
        r.font.bold = True

    # 2. Clear subtitle note
    doc.paragraphs[1].text = ""
    doc.paragraphs[2].text = ""
    doc.paragraphs[3].text = ""

    # 3. Author block in 3 columns
    # P04: Author 1
    doc.paragraphs[4].text = "Phuoc Anh Dung Nguyen\nFaculty of Information Technology\nHo Chi Minh City University of Technology (HUTECH)\nHo Chi Minh City, Vietnam\nanhdung.research@gmail.com"
    doc.paragraphs[4].style = 'Author'
    doc.paragraphs[4].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # P05: Author 2 (*Corresponding)
    doc.paragraphs[5].text = "Danh Huong Bui*\nFaculty of Information Technology\nHo Chi Minh City University of Technology (HUTECH)\nHo Chi Minh City, Vietnam\nbd.huong@hutech.edu.vn"
    doc.paragraphs[5].style = 'Author'
    doc.paragraphs[5].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # P06: Author 3
    doc.paragraphs[6].text = "Van Quy Hoang\nFaculty of Information Technology\nThuyloi University (TLU)\nHanoi, Vietnam\nquyhv@tlu.edu.vn"
    doc.paragraphs[6].style = 'Author'
    doc.paragraphs[6].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # P07: clear
    doc.paragraphs[7].text = ""

    # Abstract and Keywords
    p_abs = doc.paragraphs[10]
    p_abs.text = ""
    p_abs.style = 'Abstract'
    p_abs.paragraph_format.line_spacing = 1.05
    p_abs.paragraph_format.space_before = Pt(4)
    p_abs.paragraph_format.space_after = Pt(3)
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_abs_pre = p_abs.add_run("Abstract—")
    r_abs_pre.font.name = "Times New Roman"
    r_abs_pre.font.size = Pt(9.0)
    r_abs_pre.font.bold = True
    r_abs_pre.font.italic = True
    
    r_abs_txt = p_abs.add_run(
        "This paper proposes GUMNetHet (Heterogeneous Gated Unified Mixture Network) for multi-horizon probabilistic "
        "forecasting of refined oil prices under geopolitical volatility. The framework partitions input features into three "
        "distinct subsets and processes them via specialized domain experts: multi-scale 1D-CNN for price momentum, "
        "GRU-Attention for macroeconomic regimes, and Wavelet-KAN for non-linear shock-sensitive dynamics. These representations "
        "are dynamically fused by a horizon-aware and market-context gating router; multi-quantile outputs q ∈ {0.1, 0.5, 0.9} "
        "are optimized using pinball loss combined with residual scaling. On multi-source data spanning 11/2008–04/2026 (N = 4,512 "
        "observations) under an expanding walk-forward protocol during severe market turmoil (annualized test set volatility reaching "
        "73.04–96.29%, 1.9–2.9× historical training levels), GUMNetHet achieves the lowest MAE across all reported baselines over all "
        "seven horizons H1–H60 for both MG95 gasoline and DO 0.001% diesel. At H60, MAE is reduced by 30.1% and 22.9% relative to the "
        "strongest baseline, respectively. Directional accuracy remains high at H1–H7 but declines noticeably at H10 and H60, "
        "highlighting the structural dichotomy between price-level precision and long-term directional forecasting. The 80% prediction "
        "interval attains PICP=82.4% with PINAW=0.142. Ablation studies and router analysis confirm that expert specialization and "
        "adaptive routing contribute substantially to model performance."
    )
    r_abs_txt.font.name = "Times New Roman"
    r_abs_txt.font.size = Pt(9.0)
    r_abs_txt.font.bold = True

    p_kw = doc.paragraphs[11]
    p_kw.text = ""
    p_kw.style = 'Keywords'
    p_kw.paragraph_format.line_spacing = 1.05
    p_kw.paragraph_format.space_before = Pt(2)
    p_kw.paragraph_format.space_after = Pt(6)
    p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_kw_pre = p_kw.add_run("Keywords—")
    r_kw_pre.font.name = "Times New Roman"
    r_kw_pre.font.size = Pt(9.0)
    r_kw_pre.font.bold = True
    r_kw_pre.font.italic = True
    
    r_kw_txt = p_kw.add_run("energy price forecasting, mixture-of-experts, Wavelet-KAN, geopolitical risk, quantile forecasting, expanding walk-forward.")
    r_kw_txt.font.name = "Times New Roman"
    r_kw_txt.font.size = Pt(9.0)

    # Delete all dummy paragraphs after P11
    body_elem = doc._body._body
    p_elems = list(body_elem.xpath('./w:p'))
    for p_elem in p_elems[12:]:
        sectPr = p_elem.xpath('./w:pPr/w:sectPr')
        if sectPr:
            body_elem.append(sectPr[0])
        body_elem.remove(p_elem)

    # Also remove any tables from template
    tbl_elems = list(body_elem.xpath('./w:tbl'))
    for t_elem in tbl_elems:
        body_elem.remove(t_elem)

    # Helper function to append paragraph
    def append_p():
        return doc.add_paragraph()

    # --- SECTION I: INTRODUCTION ---
    p = append_p()
    set_ieee_heading1(p, "I. INTRODUCTION")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "The Vietnamese refined oil market exhibits distinctive structural supply characteristics, being heavily dependent "
        "on domestic refinery output and international refined imports pegged to Platts Singapore benchmarks. The price adjustment "
        "mechanism has evolved progressively toward shorter revision cycles (from 30 days, to 15 days, 10 days, and currently 7 days), "
        "escalating the operational demand for highly accurate multi-horizon price projections to safeguard domestic supply security "
        "and optimize commercial hedging."
    ])
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Modern time series architectures such as PatchTST [8], iTransformer [9], TimesNet [10], and Mamba [15] demonstrate remarkable "
        "capabilities in representation learning. Nevertheless, they frequently confront significant challenges in energy markets "
        "characterized by structural breaks, fat tails, and regime shifts triggered by geopolitical risk events (such as the Russia–Ukraine "
        "conflict or Red Sea disruptions). Furthermore, homogeneous neural networks often encounter difficulty in simultaneously "
        "handling short-term price momentum, long-term macroeconomic trends, and abrupt non-linear shocks."
    ])
    
    p = append_p()
    add_styled_para_ieee(p, [
        "The proposed GUMNetHet model addresses this bottleneck via heterogeneous feature partitioning: price and benchmark features "
        "are processed by a multi-scale 1D-CNN; macroeconomic and GPR indices by GRU-Attention; and crack-spread ratios alongside "
        "realized volatility by Wavelet-KAN. These three representations are flexibly aggregated through a gating router conditioned "
        "on the forecasting horizon and global market context. The principal contributions of this paper are fourfold: "
        "(i) A heterogeneous MoE architecture with feature partitioning grounded in economic principles and frequency domains; "
        "(ii) A horizon-aware routing mechanism (horizon-aware routing); "
        "(iii) A multi-quantile prediction head (multi-quantile head) combined with residual scaling to mitigate variance drift; and "
        "(iv) Extensive expanding walk-forward empirical evaluations on ", ('m', 'N_val'), " trading days directly within "
        "high geopolitical volatility regimes (test set volatility 1.9–2.9× training levels) with comprehensive reporting across "
        "MAE, RMSE, MAPE, ", ('m', 'R2'), ", and DA% metrics."
    ])

    # --- SECTION II: RELATED WORK ---
    p = append_p()
    set_ieee_heading1(p, "II. RELATED WORK")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Kilian [2] and Baumeister & Kilian [3] established foundational insights into decomposing structural shocks in global crude oil "
        "markets. Caldara & Iacoviello [1] introduced the Geopolitical Risk (GPR) index, establishing a standardized quantitative metric "
        "for measuring external political tensions."
    ])
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Regarding MoE, Jacobs et al. [20] introduced adaptive mixtures of local experts, while Shazeer et al. [21] and Fedus et al. [22] "
        "scaled MoE to large-scale deep neural networks. In representation learning, Kolmogorov–Arnold Networks (KAN) [23] and Wav-KAN [24] "
        "offer compelling non-linear approximation capabilities through learnable spline and continuous wavelet basis functions. "
        "Foundation time series models (Chronos [16], MOIRAI [17], UniTS [18], Time-MoE [26]) provide powerful generalized representations; "
        "however, energy forecasting under extreme geopolitical shocks still requires architectural inductive biases tailored specifically "
        "to market mechanics."
    ])

    # --- SECTION III: PROPOSED METHODOLOGY: GUMNetHet ---
    p = append_p()
    set_ieee_heading1(p, "III. PROPOSED METHODOLOGY: GUMNETHET")
    
    p = append_p()
    set_ieee_heading2(p, "A. Problem Formulation")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Let ", ('m', 'X_price'), ", ", ('m', 'X_macro'), ", ", ('m', 'X_shock'), " denote past input feature matrices with lookback length ",
        ('m', 'L_30'), " trading days. The forecasting objective is to directly estimate the cumulative log-return vector ", ('m', 'r_thc'),
        " across horizons ", ('m', 'h_set'), " for product basket ", ('m', 'c_set'), ", defined in (1):"
    ])
    
    # Insert Table 0 (Eq 1)
    doc._body._body.append(deepcopy(src_doc.tables[0]._tbl))
    
    p = append_p()
    add_styled_para_ieee(p, [
        "The future forecasted price level ", ('m', 'P_hat_thc'), " is subsequently reconstructed exactly via the deterministic inverse mapping "
        "in (2). This direct cumulative return formulation transforms the non-stationary price series into stationary log-returns ",
        ('m', 'p_val'), " (empirically validated in Section IV-B), while completely eliminating autoregressive error compounding across multi-step horizons."
    ])
    
    # Insert Table 1 (Eq 2)
    doc._body._body.append(deepcopy(src_doc.tables[1]._tbl))

    p = append_p()
    set_ieee_heading2(p, "B. Overall System Architecture")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Fig. 1 illustrates the neural architecture of GUMNetHet. Unlike conventional MoE networks that feed all features into every expert—which "
        "often degrades specialization—GUMNetHet partitions features across frequency domains and economic characteristics: high-frequency "
        "price dynamics to multi-scale 1D-CNN, slowly-evolving macroeconomic regimes to GRU-Attention, and fat-tailed crack-spread/shock indicators "
        "to Wavelet-KAN. This mechanism enforces strong domain inductive bias, effectively preventing expert degeneration."
    ])
    
    # Figure 1 Image and Caption
    p_img1 = append_p()
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img1.paragraph_format.space_before = Pt(4)
    p_img1.paragraph_format.space_after = Pt(2)
    p_img1.add_run().add_picture('scratch/images/image1.png', width=Inches(3.4))
    
    p_cap1 = append_p()
    set_ieee_caption_fig(p_cap1, 1, "Fig. 1.", "Architectural overview of GUMNetHet: feature partitioning, three heterogeneous experts, horizon-aware router, and multi-quantile head.")

    p = append_p()
    set_ieee_heading3(p, "1) Price Momentum Expert: Multi-Scale 1D-CNN:")
    p = append_p()
    add_styled_para_ieee(p, [
        "The price expert employs three parallel 1D convolutional layers with kernel sizes ", ('m', 'k_set'),
        " to extract multi-resolution temporal features, integrated with layer normalization and temporal attention according to (3) and (4):"
    ])
    doc._body._body.append(deepcopy(src_doc.tables[2]._tbl)) # Eq 3 & 4

    p = append_p()
    set_ieee_heading3(p, "2) Macroeconomic Regime Expert: GRU-Attention:")
    p = append_p()
    add_styled_para_ieee(p, [
        "The macroeconomic expert captures low-frequency trend signals via a 2-layer stacked GRU network with dropout = 0.1, "
        "extracting the final hidden state representation according to (5):"
    ])
    doc._body._body.append(deepcopy(src_doc.tables[3]._tbl)) # Eq 5

    p = append_p()
    set_ieee_heading3(p, "3) Non-Linear Shock-Absorption Expert: Wavelet-KAN:")
    p = append_p()
    add_styled_para_ieee(p, [
        "To capture severe non-linear structural breaks induced by extreme geopolitical risk events, the third expert implements a "
        "Kolmogorov–Arnold Network integrated with Mexican Hat wavelets ", ('m', 'mexican_hat'),
        " parameterized by learnable translation and dilation coefficients according to (6) and (7):"
    ])
    doc._body._body.append(deepcopy(src_doc.tables[4]._tbl)) # Eq 6 & 7

    p = append_p()
    set_ieee_heading3(p, "4) Horizon-Aware Dynamic Gating Router:")
    p = append_p()
    add_styled_para_ieee(p, [
        "To dynamically modulate expert contributions across diverse forecasting horizons and varying market regimes, the gating router "
        "accepts concatenated expert representations, horizon positional embeddings, and global summary context statistics ",
        ('m', 'x_ctx'), " according to (8), (9), and (10):"
    ])
    doc._body._body.append(deepcopy(src_doc.tables[5]._tbl)) # Eq 8, 9, 10

    p = append_p()
    set_ieee_heading3(p, "5) Residual Scaling Bounding & Multi-Quantile Prediction Head:")
    p = append_p()
    add_styled_para_ieee(p, [
        "To prevent variance explosion at long horizons (such as H60), GUMNetHet incorporates a learnable residual scaling vector ",
        ('m', 'gamma_h'), " initialized at 0.1, outputting multi-quantile return predictions for ", ('m', 'q_set'), " according to (11):"
    ])
    doc._body._body.append(deepcopy(src_doc.tables[6]._tbl)) # Eq 11

    p = append_p()
    set_ieee_heading3(p, "6) Dual-Loss Optimization:")
    p = append_p()
    add_styled_para_ieee(p, [
        "The overall objective function combines the multi-quantile pinball loss with a load-balancing regularization penalty to mitigate "
        "routing collapse according to (12), (13), and (14):"
    ])
    doc._body._body.append(deepcopy(src_doc.tables[7]._tbl)) # Eq 12, 13, 14

    # --- SECTION IV: EXPERIMENTAL SETUP ---
    p = append_p()
    set_ieee_heading1(p, "IV. EXPERIMENTAL SETUP")
    
    p = append_p()
    set_ieee_heading2(p, "A. Data and Walk-Forward Protocol")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "The dataset covers 03/11/2008–30/04/2026 with ", ('m', 'N_val'), " trading-day observations. The two reported primary targets are MG95 "
        "and DO 0.001% per Platts Singapore; exogenous covariates include inter-product Platts benchmarks, WTI, Brent, GPR ", ('ref', '1'),
        ", DXY, crude oil production, crack-spread ratios, realized volatility, and calendar features."
    ])
    
    p = append_p()
    add_styled_para_ieee(p, [
        "To prevent look-ahead bias, daily variables are indexed at ", ('m', 't_minus_1'), "; GPR is lagged by 30 calendar days; crude oil production "
        "is lagged by 7 days; and rolling features are computed strictly after applying lags. All scalers are fit exclusively on the training partition "
        "at each walk-forward step. Lookback length is ", ('m', 'L_30'), "; train/validation ratio is 85/15 per expansion. Expanding walk-forward test "
        "windows increase progressively with horizon: 100 trading days for H1–H5 (11/12/2025–30/04/2026), 150 days for H7 (02/10/2025–30/04/2026), "
        "200 days for H10 (24/07/2025–30/04/2026), 300 days for H20 (10/03/2025–30/04/2026), and 600 days for H60 (10/01/2024–30/04/2026). "
        "All test windows fall entirely within a period of clustered extreme energy shocks (Red Sea crisis, Middle East escalations, "
        "Russia–Ukraine war, and OPEC+ quota shifts). Specifically, annualized realized return volatility during the short-term test window "
        "(100 days) surges to 73.04% for MG95 (1.90× historical training volatility of 38.45%) and 96.29% for DO 0.001% (2.90× training volatility "
        "of 33.16%). The Geopolitical Risk (GPR) index averages 225.66 during the test period (nearly doubling the historical mean of 114.60) and "
        "peaks at 500.81 (90th percentile at 376.48). Price swings in the H60 test window record extreme peak-to-trough variations: 70.34 to 170.52 USD/bbl "
        "(+142.4%) for MG95 and 75.90 to 292.82 USD/bbl (+285.8%) for DO 0.001%. Comprehensive train/test splits and volatility characteristics "
        "across horizons are summarized in Table 1. This configuration serves as a rigorous stress-testing benchmark, ensuring models are evaluated "
        "on their capacity to adapt under fat-tailed regime shifts rather than static market conditions."
    ])

    # Table 1 Caption & Table
    p_cap_t1 = append_p()
    set_ieee_caption_table(p_cap_t1, 1, "TABLE I.", "EXPANDING WALK-FORWARD TRAIN/TEST SPLIT CONFIGURATIONS AND EMPIRICAL VOLATILITY CHARACTERISTICS")
    doc._body._body.append(deepcopy(src_doc.tables[8]._tbl))

    p = append_p()
    set_ieee_heading2(p, "B. Econometric Diagnostics and Training Configuration")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "ADF and KPSS unit root tests in Table 2 provide evidence of non-stationarity in raw price levels; upon log-return transformation, "
        "ADF tests reject the unit-root null hypothesis (", ('m', 'p_val'), ") across all series. Return kurtosis is exceptionally high "
        "(213.35 for WTI and 17.51 for MG95), substantiating fat tails and extreme jumps in energy price series."
    ])
    
    p_cap_t2 = append_p()
    set_ieee_caption_table(p_cap_t2, 2, "TABLE II.", "CONDENSED ECONOMETRIC STATISTICAL DIAGNOSTICS FOR PRIMARY ENERGY SERIES")
    doc._body._body.append(deepcopy(src_doc.tables[9]._tbl))
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Implementation utilizes PyTorch 2.1 / Python 3.10 with AdamW optimizer (lr = 10⁻³, weight_decay = 10⁻⁴), "
        "CosineAnnealingLR scheduler (T_max = 60, η_min = 10⁻⁵), batch size = 64, early stopping patience = 15, and 60 training epochs."
    ])

    p = append_p()
    set_ieee_heading2(p, "C. Baseline Methods and Evaluation Metrics")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Evaluated baselines comprise six representative state-of-the-art benchmarks: PatchTST [8], iTransformer [9], TimesNet [10], "
        "DLinear [11], TFT [12], and N-HiTS [14]. Evaluation metrics include MAE, RMSE, MAPE (%), ", ('m', 'R2'), ", and Directional Accuracy (DA, %). "
        "Probabilistic forecasts are quantified using Prediction Interval Coverage Probability (PICP) and Prediction Interval Normalized "
        "Average Width (PINAW) on ", ('m', 'q_interval'), "."
    ])

    # --- SECTION V: RESULTS AND DISCUSSION ---
    p = append_p()
    set_ieee_heading1(p, "V. RESULTS AND DISCUSSION")
    
    p = append_p()
    set_ieee_heading2(p, "A. Multi-Horizon Point Forecasting Performance")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Fig. 2 alongside Table 3 and Table 4 demonstrates that GUMNetHet consistently maintains lower MAE than all reported baselines "
        "across both products over all horizons, with pronounced advantages at H20–H60 where the test set undergoes massive price swings "
        "of up to +142.4% (MG95) and +285.8% (DO). At H60, gasoline MAE is 4.847 compared to 6.933 for the best baseline (a 30.1% reduction); "
        "for diesel, MAE is 7.066 versus 9.167 (a 22.9% reduction). However, ", ('m', 'R2'), " at H60 moderates to 0.155 (gasoline) and −0.007 (diesel), "
        "indicating that long-term performance represents superior price-level stabilization under high volatility rather than flawless trajectory forecasting."
    ])

    # Figure 2
    p_img2 = append_p()
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.add_run().add_picture('scratch/images/image2.png', width=Inches(3.4))
    p_cap2 = append_p()
    set_ieee_caption_fig(p_cap2, 2, "Fig. 2.", "MAE and R² degradation curves across seven forecasting horizons for MG95 and DO 0.001% (Seed=42).")

    # Table 3 (MG95)
    p_cap_t3 = append_p()
    set_ieee_caption_table(p_cap_t3, 3, "TABLE III.", "DETAILED EMPIRICAL FORECASTING PERFORMANCE ON MG95 GASOLINE (SEED=42)")
    doc._body._body.append(deepcopy(src_doc.tables[10]._tbl))

    # Table 4 (DO)
    p_cap_t4 = append_p()
    set_ieee_caption_table(p_cap_t4, 4, "TABLE IV.", "DETAILED EMPIRICAL FORECASTING PERFORMANCE ON DO 0.001% DIESEL (SEED=42)")
    doc._body._body.append(deepcopy(src_doc.tables[11]._tbl))

    p = append_p()
    set_ieee_heading2(p, "B. Directional Accuracy Analysis")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Directional accuracy in Fig. 3 reveals distinct behavior between short and long horizons. For gasoline (Table 3), GUMNetHet attains "
        "high DA between 90.95% and 95.56% at H1–H7; for diesel (Table 4), DA spans 76.65% to 84.92%. At H20, DA remains substantial "
        "(91.65% for gasoline; 71.11% for diesel). Conversely, at H10 and H60, DA drops below 50% (42.24%/27.95% for gasoline and 32.29%/19.10% for diesel). "
        "This phenomenon aligns with financial econometrics principles: over extended horizons (H60 spanning nearly three months), oil prices "
        "approximate a near random walk with exponentially compounding uncertainty, rendering directional signals noisy. Rather than taking "
        "extreme directional bets that risk catastrophic error explosion, GUMNetHet's residual scaling mechanism shrinks predictions toward "
        "robust central bounds to minimize absolute error (MAE reduced by 30.1% at H60). Consequently, GUMNetHet operates as an effective "
        "directional trading signal at short horizons (H1–H7), while transitioning into a price-level risk and uncertainty quantification "
        "tool at long horizons (H10–H60)."
    ])

    # Figure 3
    p_img3 = append_p()
    p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img3.add_run().add_picture('scratch/images/image3.png', width=Inches(3.4))
    p_cap3 = append_p()
    set_ieee_caption_fig(p_cap3, 3, "Fig. 3.", "Directional accuracy (DA%) of GUMNetHet and baselines across H1–H60 for gasoline and diesel.")

    p = append_p()
    set_ieee_heading2(p, "C. Probabilistic Forecasting, Ablation, and Router Behavior")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "The ", ('m', 'q_interval'), " prediction interval achieves an empirical coverage probability of PICP=82.4% "
        "(exceeding nominal 80%) with a normalized average width of PINAW=0.142. This validates that the model avoids under-coverage "
        "while maintaining sharpness, offering practical utility for oil importers (e.g., Petrolimex, PVOIL) in forward pricing, inventory buffer "
        "sizing, and hedging optimization. Fig. 4 illustrates that prediction intervals dynamically widen during heightened market volatility. "
        "Ablation results in Table 5 indicate that replacing Wav-KAN with MLP causes the steepest degradation among expert variants; a uniform "
        "router similarly inflates MAE, corroborating the essential roles of expert specialization and adaptive routing."
    ])

    # Table 5 (Ablation)
    p_cap_t5 = append_p()
    set_ieee_caption_table(p_cap_t5, 5, "TABLE V.", "CONDENSED ABLATION STUDY OF GUMNETHET VARIANTS (SEED=42)")
    doc._body._body.append(deepcopy(src_doc.tables[12]._tbl))

    p = append_p()
    add_styled_para_ieee(p, [
        "Empirical results reveal that removing residual scaling increases MAE by approximately 8.5%/6.3% at H20 and 14.1%/11.8% at H60 "
        "for gasoline/diesel. Router analysis in Fig. 4 indicates that under low GPR and short horizons, the CNN expert receives an average "
        "weight of ~0.48; when GPR exceeds the 90th percentile, the Wav-KAN weight surges from ~0.29 to 0.61 at medium–long horizons while "
        "the CNN weight contracts to ~0.21. These findings demonstrate that the router executes true regime-switching rather than trivial output averaging."
    ])

    # Figure 4
    p_img4 = append_p()
    p_img4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img4.add_run().add_picture('scratch/images/image4.png', width=Inches(3.4))
    p_cap4 = append_p()
    set_ieee_caption_fig(p_cap4, 4, "Fig. 4.", "Top: Multi-quantile fan chart under high volatility. Bottom: Router gating weights across low-GPR and high-GPR regimes.")

    # --- SECTION VI: CONCLUSION ---
    p = append_p()
    set_ieee_heading1(p, "VI. CONCLUSION")
    
    p = append_p()
    add_styled_para_ieee(p, [
        "This paper proposed GUMNetHet, a heterogeneous mixture-of-experts architecture for multi-horizon probabilistic refined oil price "
        "forecasting. The model synergizes multi-scale 1D-CNN, GRU-Attention, and Wavelet-KAN through a horizon-aware router, combined with "
        "multi-quantile outputs and residual scaling to bolster stability over extended horizons."
    ])
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Expanding walk-forward experiments on MG95 and DO 0.001% demonstrate that GUMNetHet achieves the lowest MAE across all evaluated "
        "baselines throughout H1–H60, yielding MAE reductions at H60 of 30.1% for MG95 and 22.9% for DO 0.001%. Ablation experiments and "
        "router weight analyses corroborate the contributions of feature partitioning, Wavelet-KAN, and adaptive routing. However, the "
        "attenuation of ", ('m', 'R2'), " and directional accuracy at long horizons underscores that model superiority lies in price-level error containment "
        "and uncertainty quantification, rather than long-range directional forecasting."
    ])
    
    p = append_p()
    add_styled_para_ieee(p, [
        "Future work will focus on multi-seed benchmarking, expanding probabilistic baseline comparisons, formal statistical significance "
        "testing for multi-step forecasts, and integrating real-time multimodal geopolitical news representations."
    ])

    # --- REFERENCES ---
    p = append_p()
    set_ieee_heading1(p, "REFERENCES")
    
    # 28 References from src_doc
    for p_ref in src_doc.paragraphs[68:]:
        if p_ref.text.strip():
            p = append_p()
            p.text = p_ref.text.strip()
            try:
                p.style = 'references'
            except Exception:
                pass
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Extract reference number [X] for bookmark
            m_r = re.match(r'\[(\d+)\]', p.text)
            if m_r:
                r_num = m_r.group(1)
                bm_name = f"ref_{r_num}"
                bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{100+int(r_num)}" w:name="{bm_name}"/>')
                bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{100+int(r_num)}"/>')
                p._p.insert(0, bm_start)
                p._p.append(bm_end)

    # Save to GUMNETHet_FAIRv4_final.docx
    target_out = 'GUMNETHet_FAIRv4_final.docx'
    doc.save(target_out)
    print(f"Successfully generated IEEE formatted final manuscript: {target_out}!")

if __name__ == '__main__':
    build_v4_final()
