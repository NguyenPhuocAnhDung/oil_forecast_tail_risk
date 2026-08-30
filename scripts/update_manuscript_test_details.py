import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

# Define helper to convert string or math token into Word XML
def create_inline_omml(xml_body):
    return f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{xml_body}</m:oMath>'

# Math snippets dictionary
M = {
    'P_tc': '<m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t,c</m:t></m:r></m:sub></m:sSub>',
    'c_set': '<m:r><m:t>c ∈ {MG95, DO 0.001%}</m:t></m:r>',
    'L_30': '<m:r><m:t>L = 30</m:t></m:r>',
    'h_set': '<m:r><m:t>h ∈ {1, 3, 5, 7, 10, 20, 60}</m:t></m:r>',
    'r_thc': '<m:sSub><m:e><m:r><m:t>r</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t> = ln(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t>) − ln(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t,c</m:t></m:r></m:sub></m:sSub><m:r><m:t>)</m:t></m:r>',
    'P_hat_thc': '<m:sSub><m:e><m:r><m:t>P̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t,c</m:t></m:r></m:sub></m:sSub><m:r><m:t> · exp(</m:t></m:r><m:sSub><m:e><m:r><m:t>r̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t>)</m:t></m:r>',
    'X_price': '<m:sSub><m:e><m:r><m:t>X</m:t></m:r></m:e><m:sub><m:r><m:t>price</m:t></m:r></m:sub></m:sSub>',
    'X_macro': '<m:sSub><m:e><m:r><m:t>X</m:t></m:r></m:e><m:sub><m:r><m:t>macro</m:t></m:r></m:sub></m:sSub>',
    'X_shock': '<m:sSub><m:e><m:r><m:t>X</m:t></m:r></m:e><m:sub><m:r><m:t>shock</m:t></m:r></m:sub></m:sSub>',
    'k_set': '<m:r><m:t>k ∈ {3, 7, 15}</m:t></m:r>',
    'router_in': '<m:r><m:t>[</m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>CNN</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>GRU</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>KAN</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ Emb(h) ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>ctx</m:t></m:r></m:sub></m:sSub><m:r><m:t>]</m:t></m:r>',
    'w_h': '<m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> = Softmax(MLP(·))</m:t></m:r>',
    'f_fused': '<m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>fused</m:t></m:r></m:sub></m:sSub><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>∑</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>h,i</m:t></m:r></m:sub></m:sSub><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>',
    'x_ctx': '<m:sSub><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>ctx</m:t></m:r></m:sub></m:sSub>',
    'q_set': '<m:r><m:t>q ∈ {0.1, 0.5, 0.9}</m:t></m:r>',
    'gamma_h_set': '<m:sSub><m:e><m:r><m:t>γ</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> ∈ (0, 1)</m:t></m:r>',
    'r_hat_q': '<m:sSubSup><m:e><m:r><m:t>r̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub><m:sup><m:r><m:t>(q)</m:t></m:r></m:sup></m:sSubSup><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>Head</m:t></m:r></m:e><m:sub><m:r><m:t>q</m:t></m:r></m:sub></m:sSub><m:r><m:t>(</m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>fused</m:t></m:r></m:sub></m:sSub><m:r><m:t>) + </m:t></m:r><m:sSub><m:e><m:r><m:t>γ</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> · </m:t></m:r><m:sSubSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub><m:sup><m:r><m:t>target</m:t></m:r></m:sup></m:sSubSup>',
    'L_tot': '<m:r><m:t>ℒ = </m:t></m:r><m:sSub><m:e><m:r><m:t>ℒ</m:t></m:r></m:e><m:sub><m:r><m:t>pinball</m:t></m:r></m:sub></m:sSub><m:r><m:t> + α</m:t></m:r><m:sSub><m:e><m:r><m:t>ℒ</m:t></m:r></m:e><m:sub><m:r><m:t>balance</m:t></m:r></m:sub></m:sSub>',
    'alpha_val': '<m:r><m:t>α = 0.01</m:t></m:r>',
    'L_bal': '<m:sSub><m:e><m:r><m:t>ℒ</m:t></m:r></m:e><m:sub><m:r><m:t>balance</m:t></m:r></m:sub></m:sSub>',
    'gamma_h': '<m:sSub><m:e><m:r><m:t>γ</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub>',
    't_minus_1': '<m:r><m:t>t − 1</m:t></m:r>',
    'N_val': '<m:r><m:t>N = 4.512</m:t></m:r>',
    'R2': '<m:sSup><m:e><m:r><m:t>R</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>',
    'DA_formula': '<m:r><m:t>DA = </m:t></m:r><m:f><m:num><m:r><m:t>1</m:t></m:r></m:num><m:den><m:r><m:t>N</m:t></m:r></m:den></m:f><m:sSubSup><m:e><m:r><m:t>∑</m:t></m:r></m:e><m:sub><m:r><m:t>t=1</m:t></m:r></m:sub><m:sup><m:r><m:t>N</m:t></m:r></m:sup></m:sSubSup><m:r><m:t> 𝕀[sign(</m:t></m:r><m:sSub><m:e><m:r><m:t>P̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub></m:sSub><m:r><m:t> − </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub></m:sSub><m:r><m:t>) = sign(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub></m:sSub><m:r><m:t> − </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub></m:sSub><m:r><m:t>)]</m:t></m:r>',
    'q_interval': '<m:r><m:t>[</m:t></m:r><m:sSub><m:e><m:r><m:t>q</m:t></m:r></m:e><m:sub><m:r><m:t>0.1</m:t></m:r></m:sub></m:sSub><m:r><m:t>, </m:t></m:r><m:sSub><m:e><m:r><m:t>q</m:t></m:r></m:e><m:sub><m:r><m:t>0.9</m:t></m:r></m:sub></m:sSub><m:r><m:t>]</m:t></m:r>',
    'p_val': '<m:r><m:t>p &lt; 0.001</m:t></m:r>',
    'mexican_hat': '<m:r><m:t>ψ(z) = (1 − </m:t></m:r><m:sSup><m:e><m:r><m:t>z</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup><m:r><m:t>) · </m:t></m:r><m:sSup><m:e><m:r><m:t>e</m:t></m:r></m:e><m:sup><m:r><m:t>−0.5z²</m:t></m:r></m:sup></m:sSup>'
}

def add_styled_paragraph(p, segments, line_spacing=1.15, space_after=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p.text = ""
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    
    for seg in segments:
        if isinstance(seg, str):
            # Parse text with Table/Figure/Reference links
            parts = re.split(r'(Bảng \d+|Hình \d+|\[\d+\])', seg)
            for part in parts:
                if not part:
                    continue
                m_tbl = re.match(r'Bảng (\d+)', part)
                m_fig = re.match(r'Hình (\d+)', part)
                m_ref = re.match(r'\[(\d+)\]', part)
                if m_tbl:
                    t_id = f"tbl_{m_tbl.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{t_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                elif m_fig:
                    f_id = f"fig_{m_fig.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{f_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                elif m_ref:
                    r_id = f"ref_{m_ref.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                else:
                    run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9.5)
        elif isinstance(seg, tuple) and seg[0] == 'm':
            m_xml = M[seg[1]]
            p._p.append(parse_xml(create_inline_omml(m_xml)))
        elif isinstance(seg, tuple) and seg[0] == 'ref':
            r_id = f"ref_{seg[1]}"
            hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                           f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                           f'<w:t>[{seg[1]}]</w:t></w:r></w:hyperlink>')
            p._p.append(hl)
        elif isinstance(seg, tuple) and seg[0] == 'refs':
            for idx, r_num in enumerate(seg[1]):
                r_id = f"ref_{r_num}"
                hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                               f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                               f'<w:t>[{r_num}]</w:t></w:r></w:hyperlink>')
                p._p.append(hl)
                if idx < len(seg[1]) - 1:
                    r_c = p.add_run(", ")
                    r_c.font.name = 'Times New Roman'
                    r_c.font.size = Pt(9.5)

def update_document(target_docx_path):
    print(f"Loading {target_docx_path}...")
    doc = docx.Document(target_docx_path)
    
    # 1. P[5] Abstract - Preserve wording, enrich volatile test set context
    add_styled_paragraph(doc.paragraphs[5], [
        "Tóm tắt—Bài báo đề xuất GUMNetHet (Heterogeneous Gated Unified Mixture Network) cho dự báo xác suất đa chu kỳ "
        "giá xăng dầu thành phẩm dưới biến động địa chính trị. Mô hình phân vùng đặc trưng thành ba nhóm và xử lý bằng "
        "các expert chuyên biệt: CNN-1D đa tỷ lệ cho động lượng giá, GRU-Attention cho trạng thái vĩ mô và Wavelet-KAN cho "
        "các quan hệ phi tuyến nhạy cú sốc. Các biểu diễn được kết hợp bởi bộ định tuyến nhận biết horizon và ngữ cảnh thị trường; "
        "đầu ra đa phân vị ", ('m', 'q_set'), " được tối ưu bằng pinball loss và residual scaling. Trên dữ liệu đa nguồn 11/2008–04/2026 (",
        ('m', 'N_val'), " quan sát) theo giao thức expanding walk-forward trong các giai đoạn thị trường biến động khốc liệt "
        "(độ biến động annualized volatility tập test đạt 73,04–96,29%, gấp 1,9–2,9 lần tập train lịch sử), GUMNetHet đạt MAE "
        "thấp nhất trong nhóm baseline được báo cáo ở cả bảy horizon H1–H60 cho MG95 và DO 0.001%. Tại H60, MAE giảm lần lượt 30,1% "
        "và 22,9% so với baseline tốt nhất. Độ chính xác xu hướng đạt mức cao ở H1–H7 nhưng giảm rõ ở H10 và H60, cho thấy sự khác biệt "
        "giữa độ chính xác mức giá và dự báo hướng ở horizon dài. Khoảng dự báo 80% đạt PICP=82,4% với PINAW=0,142. Kết quả ablation "
        "và phân tích router cho thấy chuyên môn hóa expert và định tuyến thích ứng đóng góp đáng kể vào hiệu năng của mô hình."
    ], line_spacing=1.10, space_after=4)

    # 2. P[10] Introduction Contributions - Preserve exact contributions, specify volatile test conditions
    add_styled_paragraph(doc.paragraphs[10], [
        "Mô hình đề xuất GUMNetHet giải quyết điểm nghẽn này bằng kỹ thuật phân vùng đặc trưng (feature partitioning): "
        "nhóm giá và benchmark được xử lý bởi CNN-1D đa tỷ lệ; nhóm vĩ mô và chỉ số GPR bởi GRU-Attention; nhóm tỷ lệ crack-spread "
        "và độ biến động bởi Wavelet-KAN. Ba biểu diễn được hợp nhất linh hoạt thông qua một bộ định tuyến (gating router) phụ thuộc "
        "vào horizon dự báo và ngữ cảnh thị trường. Đóng góp chính của bài báo gồm: (i) Kiến trúc MoE dị thể với phân vùng đặc trưng "
        "theo bản chất kinh tế và miền tần số; (ii) Bộ định tuyến nhận biết horizon (horizon-aware routing); (iii) Đầu ra đa phân vị "
        "(multi-quantile head) kết hợp residual scaling để kiểm soát độ trôi phương sai; và (iv) Đánh giá thực nghiệm mở rộng walk-forward "
        "trên ", ('m', 'N_val'), " ngày giao dịch trực tiếp trong các giai đoạn biến động địa chính trị cao (độ biến động tập test cao "
        "gấp 1,9–2,9 lần tập train) với bảng kết quả đầy đủ các chỉ số MAE, RMSE, MAPE, ", ('m', 'R2'), " và DA%."
    ])

    # 3. P[37] Data Section
    add_styled_paragraph(doc.paragraphs[37], [
        "Dữ liệu bao phủ 03/11/2008–30/04/2026 với ", ('m', 'N_val'), " quan sát ngày giao dịch. Hai target được báo cáo là MG95 và "
        "DO 0.001% theo Platts; các biến ngoại sinh gồm Platts liên sản phẩm, WTI, Brent, GPR ", ('ref', '1'),
        ", DXY, sản lượng dầu, crack-spread, realized volatility và biến lịch."
    ])

    # 4. P[38] Test set & Volatility details (100% empirical code grounding)
    add_styled_paragraph(doc.paragraphs[38], [
        "Để tránh look-ahead bias, các biến ngày được dùng ở ", ('m', 't_minus_1'), "; GPR được áp lag 30 ngày lịch; sản lượng dầu áp lag 7 ngày; "
        "các biến rolling được tính sau khi lag. Mọi scaler chỉ fit trên train tại từng bước walk-forward. Lookback ", ('m', 'L_30'),
        "; train/validation=85/15 ở mỗi lần mở rộng. Cửa sổ test expanding walk-forward được thiết kế tăng dần theo horizon: 100 ngày giao dịch "
        "ở H1–H5 (11/12/2025–30/04/2026), 150 ngày ở H7 (02/10/2025–30/04/2026), 200 ngày ở H10 (24/07/2025–30/04/2026), 300 ngày ở H20 "
        "(10/03/2025–30/04/2026) và 600 ngày ở H60 (10/01/2024–30/04/2026). Toàn bộ các cửa sổ test đều nằm trọn trong giai đoạn thị trường "
        "năng lượng chịu chuỗi cú sốc địa chính trị dồn dập và phân cụm biến động cực đoan (khủng hoảng Biển Đỏ, căng thẳng Trung Đông, "
        "xung đột Nga–Ukraine và biến động hạn ngạch OPEC+). Cụ thể, độ biến động thực tế hàng năm (annualized volatility) của lợi suất trong "
        "cửa sổ test ngắn hạn (100 ngày) tăng vọt lên 73,04% đối với MG95 (gấp 1,90 lần mức 38,45% của tập train lịch sử) và 96,29% đối với "
        "DO 0.001% (gấp 2,90 lần mức 33,16% của train). Chỉ số rủi ro địa chính trị GPR trong giai đoạn test đạt trung bình 225,66 "
        "(gần gấp đôi mức trung bình toàn kỳ 114,60) và đạt đỉnh 500,81 (phân vị 90% đạt 376,48). Biên độ dao động giá trong cửa sổ test H60 "
        "ghi nhận mức chênh lệch đỉnh–đáy cực lớn: từ 70,34 đến 170,52 USD/thùng (+142,4%) đối với MG95 và từ 75,90 đến 292,82 USD/thùng "
        "(+285,8%) đối với DO 0.001%. Thiết lập này đóng vai trò như một bài kiểm tra áp lực (stress-testing) thực thụ, bảo đảm các mô hình "
        "được đánh giá về khả năng thích ứng dưới các chế độ biến động đuôi dày (fat-tailed regimes) thay vì chỉ học trong điều kiện thị trường tĩnh."
    ])

    # 5. P[47] Horizon Results Discussion
    add_styled_paragraph(doc.paragraphs[47], [
        "Hình 2 cùng Bảng 2 và Bảng 3 cho thấy GUMNetHet duy trì MAE thấp hơn nhóm baseline được báo cáo ở cả hai sản phẩm trên toàn bộ các horizon, "
        "đặc biệt vượt trội ở H20–H60 trong bối cảnh tập test trải qua biên độ dao động giá lên tới +142,4% (MG95) và +285,8% (DO). Ở H60, MAE "
        "của xăng là 4,847 so với 6,933 của baseline tốt nhất (giảm 30,1%); với dầu là 7,066 so với 9,167 (giảm 22,9%). Tuy nhiên, ",
        ('m', 'R2'), " tại H60 giảm còn 0,155 (xăng) và −0,007 (dầu), vì vậy kết quả dài hạn nên được hiểu là ổn định sai số mức giá tốt hơn "
        "baseline trong môi trường biến động mạnh, không phải dự báo quỹ đạo dài hạn hoàn hảo."
    ])

    doc.save(target_docx_path)
    print(f"Successfully updated {target_docx_path}!")

if __name__ == '__main__':
    update_document('GUMNETHet_FAIRv3 - Copy.docx')
