import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os
import sys
import re
import zipfile
import io
from copy import deepcopy

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

from scripts.table_builder_helper import M, create_inline_omml, style_table

def build_manuscript(is_en=True, out_path="GUMNETHet_FAIRv4_final.docx"):
    print(f"Building {'English' if is_en else 'Vietnamese'} manuscript -> {out_path}...")
    
    # Use base document that has all XML drawings and tables perfectly configured
    base_path = 'GUMNETHet_FAIRv3 - Copy.backup.docx'
    doc = docx.Document(base_path)
    
    # Page setup - A4 with IEEE conference margins
    for s in doc.sections:
        s.page_width = Inches(8.27)  # A4 width
        s.page_height = Inches(11.69) # A4 height
        s.top_margin = Inches(0.75)   # 54pt - Prevents title clipping!
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.625) # 45pt
        s.right_margin = Inches(0.625)
    
    # Paragraph references
    p_title = doc.paragraphs[0]
    p_authors = doc.paragraphs[1]
    p_affil1 = doc.paragraphs[2]
    p_affil2 = doc.paragraphs[3]
    p_corr = doc.paragraphs[4]
    
    p_abstract = doc.paragraphs[5]
    p_keywords = doc.paragraphs[6]
    
    p_h1_intro = doc.paragraphs[7]
    p_intro1 = doc.paragraphs[8]
    p_intro2 = doc.paragraphs[9]
    p_intro_contrib = doc.paragraphs[10]
    
    p_h1_related = doc.paragraphs[11]
    p_related1 = doc.paragraphs[12]
    p_related2 = doc.paragraphs[13]
    
    p_h1_method = doc.paragraphs[14]
    p_h2_prob = doc.paragraphs[15]
    p_form1_txt = doc.paragraphs[16]
    p_form2_txt = doc.paragraphs[17]
    p_h2_arch = doc.paragraphs[18]
    p_arch_txt = doc.paragraphs[19]
    p_dr1 = doc.paragraphs[20] # DRAWING 1
    p_cap_fig1 = doc.paragraphs[21]
    
    p_h2_cnn = doc.paragraphs[22]
    p_cnn_txt = doc.paragraphs[23]
    p_h2_gru = doc.paragraphs[24]
    p_gru_txt = doc.paragraphs[25]
    p_h2_wkan = doc.paragraphs[26]
    p_wkan_txt = doc.paragraphs[27]
    p_h2_router = doc.paragraphs[28]
    p_router_txt = doc.paragraphs[29]
    p_h2_res = doc.paragraphs[30]
    p_res_txt = doc.paragraphs[31]
    p_h2_loss = doc.paragraphs[32]
    p_loss_txt = doc.paragraphs[33]
    
    p_h1_exp = doc.paragraphs[35]
    p_h2_data = doc.paragraphs[36]
    p_data_txt = doc.paragraphs[37]
    p_data_wf_txt = doc.paragraphs[38]
    
    p_h2_diag = doc.paragraphs[39]
    p_diag_txt = doc.paragraphs[40]
    p_cap_tbl2 = doc.paragraphs[41]
    p_impl_txt = doc.paragraphs[42]
    
    p_h2_base = doc.paragraphs[43]
    p_base_txt = doc.paragraphs[44]
    
    p_h1_res = doc.paragraphs[45]
    p_h2_point = doc.paragraphs[46]
    p_point_txt = doc.paragraphs[47]
    p_dr2 = doc.paragraphs[48] # DRAWING 2
    p_cap_fig2 = doc.paragraphs[49]
    p_cap_tbl3 = doc.paragraphs[50]
    p_cap_tbl4 = doc.paragraphs[51]
    
    p_h2_da = doc.paragraphs[52]
    p_da_txt = doc.paragraphs[53]
    p_dr3 = doc.paragraphs[54] # DRAWING 3
    p_cap_fig3 = doc.paragraphs[55]
    
    p_h2_prob_res = doc.paragraphs[56]
    p_prob_res_txt = doc.paragraphs[57]
    p_cap_tbl5 = doc.paragraphs[58]
    p_res_scale_txt = doc.paragraphs[59]
    p_dr4 = doc.paragraphs[60] # DRAWING 4
    p_cap_fig4 = doc.paragraphs[61]
    
    p_h1_concl = doc.paragraphs[62]
    p_concl1 = doc.paragraphs[63]
    p_concl2 = doc.paragraphs[64]
    p_concl3 = doc.paragraphs[65]
    p_h1_refs = doc.paragraphs[66]

    # Helper function for setting IEEE headings
    def set_h1(p, text):
        p.text = ""
        p.style = 'H1x'
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="0"/>'))
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.0)
        r.font.bold = True

    def set_h2(p, text):
        p.text = ""
        p.style = 'H2x'
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="1"/>'))
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.italic = True

    def set_caption(p, prefix, text, tag_num, is_table=True):
        p.text = ""
        p.style = 'Caption'
        p.paragraph_format.line_spacing = 1.05
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_table else WD_ALIGN_PARAGRAPH.LEFT
        
        r_pre = p.add_run(prefix)
        r_pre.font.name = "Times New Roman"
        r_pre.font.size = Pt(8.5)
        r_pre.font.bold = True
        
        r_txt = p.add_run(" " + text)
        r_txt.font.name = "Times New Roman"
        r_txt.font.size = Pt(8.5)
        r_txt.font.bold = False
        r_txt.font.italic = False
        
        bm_name = f"tbl_{tag_num}" if is_table else f"fig_{tag_num}"
        bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{tag_num}" w:name="{bm_name}"/>')
        bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{tag_num}"/>')
        p._p.insert(0, bm_start)
        p._p.append(bm_end)

    def set_body(p, segments):
        p.text = ""
        p.style = 'Normal'
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        
        pat = r'(Table \d+|Fig\. \d+|\[\d+\])' if is_en else r'(Bảng \d+|Hình \d+|\[\d+\])'
        
        for seg in segments:
            if isinstance(seg, str):
                parts = re.split(pat, seg)
                for part in parts:
                    if not part:
                        continue
                    m_tbl = re.match(r'(?:Table|Bảng) (\d+)', part)
                    m_fig = re.match(r'(?:Fig\.|Hình) (\d+)', part)
                    m_ref = re.match(r'\[(\d+)\]', part)
                    if m_tbl:
                        t_id = f"tbl_{m_tbl.group(1)}"
                        hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{t_id}" w:history="1">'
                                       f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                       f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                        p._p.append(hl)
                    elif m_fig:
                        f_id = f"fig_{m_fig.group(1)}"
                        hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{f_id}" w:history="1">'
                                       f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                       f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                        p._p.append(hl)
                    elif m_ref:
                        r_id = f"ref_{m_ref.group(1)}"
                        hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                                       f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                       f'<w:t>[{m_ref.group(1)}]</w:t></w:r></w:hyperlink>')
                        p._p.append(hl)
                    else:
                        run = p.add_run(part)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9.5)
                        run.font.bold = False
            elif isinstance(seg, tuple) and seg[0] == 'm':
                m_xml = M[seg[1]]
                p._p.append(parse_xml(create_inline_omml(m_xml)))
            elif isinstance(seg, tuple) and seg[0] == 'ref':
                r_id = f"ref_{seg[1]}"
                hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                               f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                               f'<w:t>[{seg[1]}]</w:t></w:r></w:hyperlink>')
                p._p.append(hl)

    def set_abstract(p, prefix, text_fragments):
        p.text = ""
        p.style = 'Abstract'
        p.paragraph_format.line_spacing = 1.10
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        r_pre = p.add_run(prefix)
        r_pre.font.name = "Times New Roman"
        r_pre.font.size = Pt(9.0)
        r_pre.font.bold = True
        r_pre.font.italic = True
        
        for item in text_fragments:
            if isinstance(item, str):
                r = p.add_run(item)
                r.font.name = "Times New Roman"
                r.font.size = Pt(9.0)
                r.font.bold = False
            elif isinstance(item, tuple) and item[0] == 'm':
                m_xml = M[item[1]]
                p._p.append(parse_xml(create_inline_omml(m_xml)))

    # 1. Title
    p_title.text = "ROBUST PROBABILISTIC ENERGY FORECASTING UNDER GEOPOLITICAL SHOCKS: AN ADAPTIVE MIXTURE OF LOCAL-GLOBAL EXPERTS"
    p_title.style = 'PaperTitle'
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.paragraph_format.line_spacing = 1.15
    for r in p_title.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(15.0)
        r.font.bold = True

    # 2. Authors
    if is_en:
        p_authors.text = "Phuoc Anh Dung Nguyen¹, Danh Huong Bui¹*, Van Quy Hoang²"
        p_affil1.text = "¹Faculty of Information Technology, Ho Chi Minh City University of Technology (HUTECH), Ho Chi Minh City, Vietnam"
        p_affil2.text = "²Faculty of Information Technology, Thuyloi University (TLU), Hanoi, Vietnam"
        p_corr.text = "Emails: anhdungnguyen955@gmail.com, bd.huong@hutech.edu.vn (*Corresponding author), hoangvanquy@tlu.edu.vn"
    else:
        p_authors.text = "Nguyễn Phước Anh Dũng¹, Bùi Danh Hường¹*, Hoàng Văn Quý²"
        p_affil1.text = "¹Khoa Công nghệ Thông tin, Trường Đại học Công nghệ TP.HCM (HUTECH), TP. Hồ Chí Minh, Việt Nam"
        p_affil2.text = "²Khoa Công nghệ Thông tin, Trường Đại học Thủy lợi (TLU), Hà Nội, Việt Nam"
        p_corr.text = "Email: anhdungnguyen955@gmail.com, bd.huong@hutech.edu.vn (*Tác giả liên hệ), hoangvanquy@tlu.edu.vn"

    p_authors.style = 'Authors'
    p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_authors.paragraph_format.space_before = Pt(2)
    p_authors.paragraph_format.space_after = Pt(2)
    for r in p_authors.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.0)
        r.font.bold = True

    for p_aff in [p_affil1, p_affil2, p_corr]:
        p_aff.style = 'Affil'
        p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_aff.paragraph_format.space_before = Pt(0)
        p_aff.paragraph_format.space_after = Pt(1)
        for r in p_aff.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(8.5)
            r.font.italic = True

    # 3. Abstract & Keywords
    if is_en:
        set_abstract(p_abstract, "Abstract—", [
            "This paper proposes GUMNetHet (Heterogeneous Gated Unified Mixture Network) for multi-horizon probabilistic "
            "forecasting of refined oil prices under geopolitical volatility. The framework partitions input features into three "
            "distinct subsets and processes them via specialized domain experts: multi-scale 1D-CNN for price momentum, "
            "GRU-Attention for macroeconomic regimes, and Wavelet-KAN for non-linear shock-sensitive dynamics. These representations "
            "are dynamically fused by a horizon-aware and market-context gating router; multi-quantile outputs ", ('m', 'q_set'),
            " are optimized using pinball loss combined with residual scaling. On multi-source data spanning 11/2008–04/2026 (",
            ('m', 'N_val'), " observations) under an expanding walk-forward protocol during severe market turmoil (annualized test set "
            "volatility reaching 73.04–96.29%, 1.9–2.9× historical training levels), GUMNetHet achieves the lowest MAE across all "
            "reported baselines over all seven horizons H1–H60 for both MG95 gasoline and DO 0.001% diesel. At H60, MAE is reduced "
            "by 30.1% and 22.9% relative to the strongest baseline, respectively. Directional accuracy remains high at H1–H7 but declines "
            "noticeably at H10 and H60, highlighting the structural dichotomy between price-level precision and long-term directional forecasting. "
            "The 80% prediction interval attains PICP=82.4% with PINAW=0.142. Ablation studies and router analysis confirm that expert "
            "specialization and adaptive routing contribute substantially to model performance."
        ])
        set_abstract(p_keywords, "Keywords—", [
            "energy price forecasting; mixture-of-experts; Wavelet-KAN; geopolitical risk; quantile forecasting; expanding walk-forward."
        ])
    else:
        set_abstract(p_abstract, "Tóm tắt—", [
            "Bài báo đề xuất GUMNetHet (Heterogeneous Gated Unified Mixture Network) cho dự báo xác suất đa chu kỳ "
            "giá xăng dầu thành phẩm dưới biến động địa chính trị. Mô hình phân vùng đặc trưng thành ba nhóm và xử lý bằng "
            "các expert chuyên biệt: CNN-1D đa tỷ lệ cho động lượng giá, GRU-Attention cho trạng thái vĩ mô và Wavelet-KAN cho "
            "các quan hệ phi tuyến nhạy cú sốc. Các biểu diễn được kết hợp bởi bộ định tuyến nhận biết horizon và ngữ cảnh thị trường; "
            "đầu ra đa phân vị ", ('m', 'q_set'), " được tối ưu bằng pinball loss và residual scaling. Trên dữ liệu đa nguồn 11/2008–04/2026 (",
            ('m', 'N_val'), " quan sát) theo giao thức expanding walk-forward trong các giai đoạn thị trường biến động khốc liệt "
            "(độ biến động annualized volatility tập test đạt 73,04–96,29%, gấp 1,9–2,9 lần tập train lịch sử), GUMNetHet đạt MAE "
            "thấp nhất trong nhóm baseline được báo cáo ở cả bảy horizon H1–H60 cho MG95 và DO 0.001%. Tại H60, MAE giảm lần lượt 30,1% "
            "và 22,9% so với baseline tốt nhất. Độ chính xác xu hướng đạt mức cao ở H1–H7 nhưng giảm rõ ở H10 và H60, cho thấy sự khác biệt "
            "giữa độ chính xác mức giá và dự báo hướng ở horizon dài. Khoảng dự báo 80% đạt PICP=82,4% với PINAW=0,142. Kết quả ablation "
            "và phân tích router cho thấy chuyên môn hóa expert và định tuyến thích ứng đóng góp đáng kể vào hiệu năng của mô hình."
        ])
        set_abstract(p_keywords, "Từ khóa—", [
            "dự báo giá năng lượng; mixture-of-experts; Wavelet-KAN; rủi ro địa chính trị; dự báo phân vị; expanding walk-forward."
        ])

    # 4. Section 1: Introduction
    set_h1(p_h1_intro, "1. INTRODUCTION" if is_en else "1. GIỚI THIỆU")
    if is_en:
        set_body(p_intro1, [
            "The Vietnamese refined oil market exhibits distinctive structural supply characteristics, being heavily dependent "
            "on domestic refinery output and international refined imports pegged to Platts Singapore benchmarks. The price adjustment "
            "mechanism has evolved progressively toward shorter revision cycles (from 30 days, to 15 days, 10 days, and currently 7 days), "
            "escalating the operational demand for highly accurate multi-horizon price projections to safeguard domestic supply security "
            "and optimize commercial hedging."
        ])
        set_body(p_intro2, [
            "Modern time series architectures such as PatchTST [8], iTransformer [9], TimesNet [10], and Mamba [15] demonstrate remarkable "
            "capabilities in representation learning. Nevertheless, they frequently confront significant challenges in energy markets "
            "characterized by structural breaks, fat tails, and regime shifts triggered by geopolitical risk events (such as the Russia–Ukraine "
            "conflict or Red Sea disruptions). Furthermore, homogeneous neural networks often encounter difficulty in simultaneously "
            "handling short-term price momentum, long-term macroeconomic trends, and abrupt non-linear shocks."
        ])
        set_body(p_intro_contrib, [
            "The proposed GUMNetHet model addresses this bottleneck via heterogeneous feature partitioning: price and benchmark features "
            "are processed by a multi-scale 1D-CNN; macroeconomic and GPR indices by GRU-Attention; and crack-spread ratios alongside "
            "realized volatility by Wavelet-KAN. These three representations are flexibly aggregated through a gating router conditioned "
            "on the forecasting horizon and global market context. The principal contributions of this paper are fourfold: "
            "(i) A heterogeneous MoE architecture with feature partitioning grounded in economic principles and frequency domains; "
            "(ii) A horizon-aware routing mechanism (horizon-aware routing); "
            "(iii) A multi-quantile prediction head (multi-quantile head) combined with residual scaling to mitigate variance drift; and "
            "(iv) Extensive expanding walk-forward empirical evaluations on ", ('m', 'N_val'), " trading days directly within "
            "high geopolitical volatility regimes (test set volatility 1.9–2.9× training levels) with comprehensive reporting across "
            "MAE, RMSE, MAPE, ", ('m', 'R2'), ", and DA% metrics."
        ])
    else:
        set_body(p_intro1, [
            "Thị trường xăng dầu Việt Nam có đặc thù cấu trúc nguồn cung mang tính chiến lược: khoảng 70% tổng sản lượng tiêu thụ nội địa "
            "được cung ứng bởi hai nhà máy lọc hóa dầu trong nước (Dung Quất và Nghi Sơn), trong khi 30% nhu cầu còn lại bắt buộc phải nhập "
            "khẩu trực tiếp từ các thị trường quốc tế. Trong đó, thị trường Singapore là địa bàn nhập khẩu trọng yếu nhất, với giá giao dịch "
            "thành phẩm Mean of Platts Singapore (MOPS)—tiêu biểu là Mogas 95 (MG95) và Gasoil 0.001%S (DO 0.001%)—đóng vai trò là hệ quy chiếu "
            "định giá cơ sở cho mọi hợp đồng thương mại, tính toán giá vốn và quản lý chuỗi cung ứng. Tuy nhiên, giá Platts biến động cực kỳ "
            "phức tạp do phản ứng đồng thời với các cú sốc cung–cầu toàn cầu, biến động tỷ giá và rủi ro địa chính trị (GPR) [1], [2], [3]. "
            "Các giai đoạn sụp đổ giá dầu 2014–2016, chiến tranh giá OPEC+ năm 2020, xung đột Nga–Ukraine 2022 và căng thẳng Biển Đỏ 2023–2024 "
            "cho thấy chuỗi giá thành phẩm thường xuyên xuất hiện bước nhảy phi tuyến, dịch chuyển chế độ và phân cụm biến động mạnh. Do đó, "
            "nhu cầu dự báo chính xác giá xăng dầu Platts trong ngắn hạn (H1–H7) và trung hạn (H10–H60) là đòi hỏi cấp thiết phục vụ trực "
            "tiếp cho việc ra quyết định kinh doanh, tối ưu hóa kế hoạch mua hàng, quản trị tồn kho và phòng hộ rủi ro (hedging) của các doanh "
            "nghiệp đầu mối xăng dầu lớn như Petrolimex và PVOIL."
        ])
        set_body(p_intro2, [
            "Các kiến trúc chuỗi thời gian hiện đại như PatchTST [8], iTransformer [9], TimesNet [10], DLinear [11], Mamba [15] và Chronos [16] "
            "đã cải thiện đáng kể hiệu năng trên các benchmark tổng quát, nhưng phần lớn đều xử lý toàn bộ các biến đầu vào trong một không gian "
            "biểu diễn tương đối đồng nhất. Với dữ liệu năng lượng, ba nhóm tín hiệu có bản chất cơ bản khác nhau: động lượng giá tần số cao, "
            "trạng thái vĩ mô biến đổi chậm và phản ứng phi tuyến nhạy cú sốc biên độ lớn đòi hỏi các inductive bias chuyên biệt. Các mạng "
            "kết hợp chuyên gia (MoE) truyền thống [20], [21], [22] thường đưa cùng một tập đặc trưng tới mọi expert, làm suy giảm mức độ chuyên môn hóa."
        ])
        set_body(p_intro_contrib, [
            "Mô hình đề xuất GUMNetHet giải quyết điểm nghẽn này bằng kỹ thuật phân vùng đặc trưng (feature partitioning): "
            "nhóm giá và benchmark được xử lý bởi CNN-1D đa tỷ lệ; nhóm vĩ mô và chỉ số GPR bởi GRU-Attention; nhóm tỷ lệ crack-spread "
            "và độ biến động bởi Wavelet-KAN. Ba biểu diễn được hợp nhất linh hoạt thông qua một bộ định tuyến (gating router) phụ thuộc "
            "vào horizon dự báo và ngữ cảnh thị trường. Đóng góp chính của bài báo gồm: "
            "(i) Kiến trúc MoE dị thể với phân vùng đặc trưng theo bản chất kinh tế và miền tần số; "
            "(ii) Bộ định tuyến nhận biết horizon (horizon-aware routing); "
            "(iii) Đầu ra đa phân vị (multi-quantile head) kết hợp residual scaling để kiểm soát độ trôi phương sai; và "
            "(iv) Đánh giá thực nghiệm mở rộng walk-forward trên ", ('m', 'N_val'), " ngày giao dịch trực tiếp trong các giai đoạn "
            "biến động địa chính trị cao (độ biến động tập test cao gấp 1,9–2,9 lần tập train) với bảng kết quả đầy đủ các chỉ số "
            "MAE, RMSE, MAPE, ", ('m', 'R2'), " và DA%."
        ])

    # 5. Section 2: Related Work
    set_h1(p_h1_related, "2. RELATED WORK" if is_en else "2. NGHIÊN CỨU LIÊN QUAN")
    if is_en:
        set_body(p_related1, [
            "Kilian [2] and Baumeister & Kilian [3] established foundational insights into decomposing structural shocks in global crude oil "
            "markets. Caldara & Iacoviello [1] introduced the Geopolitical Risk (GPR) index, establishing a standardized quantitative metric "
            "for measuring external political tensions."
        ])
        set_body(p_related2, [
            "Regarding MoE, Jacobs et al. [20] introduced adaptive mixtures of local experts, while Shazeer et al. [21] and Fedus et al. [22] "
            "scaled MoE to large-scale deep neural networks. In representation learning, Kolmogorov–Arnold Networks (KAN) [23] and Wav-KAN [24] "
            "offer compelling non-linear approximation capabilities through learnable spline and continuous wavelet basis functions. "
            "Foundation time series models (Chronos [16], MOIRAI [17], UniTS [18], Time-MoE [26]) provide powerful generalized representations; "
            "however, energy forecasting under extreme geopolitical shocks still requires architectural inductive biases tailored specifically "
            "to market mechanics."
        ])
    else:
        set_body(p_related1, [
            "Kilian [2] và Baumeister & Kilian [3] đặt nền tảng cho việc phân tách cú sốc cung–cầu trong thị trường dầu, trong khi chỉ số GPR "
            "của Caldara & Iacoviello [1] cung cấp thước đo định lượng cho rủi ro địa chính trị. Ở phía mô hình, PatchTST [8], iTransformer [9] "
            "và TimesNet [10] đại diện cho các thiết kế Transformer/biến thiên thời gian hiện đại; DLinear [11] cho thấy mô hình tuyến tính có "
            "phân rã vẫn là baseline mạnh. TFT [12], N-BEATS [13] và N-HiTS [14] mở rộng dự báo đa horizon; Mamba [15] đưa state-space chọn lọc "
            "vào sequence modeling; Chronos [16], TimesFM [17], MOIRAI [18] và TTM [19] mở rộng sang foundation time-series models."
        ])
        set_body(p_related2, [
            "Về MoE, Jacobs et al. [20] giới thiệu adaptive mixtures of local experts; Shazeer et al. [21] và Switch Transformer [22] "
            "phát triển sparse gating ở quy mô lớn. KAN [23] và Wav-KAN [24] thay thế activation cố định bằng hàm cơ sở có thể học, trong đó "
            "wavelet phù hợp với tín hiệu cục bộ đa độ phân giải [25]. Time-MoE [26] và TimeMixer++ [27] tiếp tục cho thấy giá trị của "
            "expert routing trong chuỗi thời gian. Khác với các thiết kế expert đồng nhất, GUMNetHet gán các loại kiến trúc khác nhau cho "
            "các nhóm biến khác nhau và điều kiện hóa router trên horizon."
        ])

    # 6. Section 3: Proposed Methodology
    set_h1(p_h1_method, "3. PROPOSED METHODOLOGY: GUMNetHet" if is_en else "3. PHƯƠNG PHÁP GUMNetHet")
    set_h2(p_h2_prob, "3.1. Problem Formulation" if is_en else "3.1. Phát biểu bài toán")
    
    if is_en:
        set_body(p_form1_txt, [
            "Let ", ('m', 'X_price'), ", ", ('m', 'X_macro'), ", ", ('m', 'X_shock'), " denote past input feature matrices with lookback length ",
            ('m', 'L_30'), " trading days. The forecasting objective is to directly estimate the cumulative log-return vector ", ('m', 'r_thc'),
            " across horizons ", ('m', 'h_set'), " for product basket ", ('m', 'c_set'), ", defined in (1):"
        ])
        set_body(p_form2_txt, [
            "The future forecasted price level ", ('m', 'P_hat_thc'), " is subsequently reconstructed exactly via the deterministic inverse mapping "
            "in (2). This direct cumulative return formulation transforms the non-stationary price series into stationary log-returns ",
            ('m', 'p_val'), " (empirically validated in Section 4.2), while completely eliminating autoregressive error compounding across multi-step horizons."
        ])
    else:
        set_body(p_form1_txt, [
            "Gọi ", ('m', 'X_price'), ", ", ('m', 'X_macro'), ", ", ('m', 'X_shock'), " là các ma trận đặc trưng đầu vào quá khứ với lookback ",
            ('m', 'L_30'), " ngày giao dịch. Mục tiêu dự báo là ước lượng trực tiếp vectơ lợi suất log tích lũy ", ('m', 'r_thc'),
            " qua chu kỳ ", ('m', 'h_set'), " cho nhóm sản phẩm ", ('m', 'c_set'), " được định nghĩa theo (1):"
        ])
        set_body(p_form2_txt, [
            "Mức giá dự báo tương lai ", ('m', 'P_hat_thc'), " sau đó được phục hồi chính xác thông qua phép biến đổi nghịch đảo tiền định theo (2). "
            "Thiết kế dự báo lợi suất tích lũy trực tiếp (direct cumulative) này giúp đưa chuỗi giá không dừng về chuỗi lợi suất dừng ",
            ('m', 'p_val'), " (chứng minh ở Mục 4.2), đồng thời triệt tiêu hoàn toàn hiện tượng tích lũy sai số đệ quy (autoregressive error compounding) "
            "khi dự báo qua nhiều bước thời gian."
        ])

    set_h2(p_h2_arch, "3.2. Overall System Architecture" if is_en else "3.2. Khung kiến trúc hệ thống GUMNetHet")
    if is_en:
        set_body(p_arch_txt, [
            "Fig. 1 illustrates the neural architecture of GUMNetHet. Unlike conventional MoE networks that feed all features into every expert—which "
            "often degrades specialization—GUMNetHet partitions features across frequency domains and economic characteristics: high-frequency "
            "price dynamics to multi-scale 1D-CNN, slowly-evolving macroeconomic regimes to GRU-Attention, and fat-tailed crack-spread/shock indicators "
            "to Wavelet-KAN. This mechanism enforces strong domain inductive bias, effectively preventing expert degeneration."
        ])
        set_caption(p_cap_fig1, "Fig. 1.", "Architectural overview of GUMNetHet: feature partitioning, three heterogeneous experts, horizon-aware router, and multi-quantile head.", "1", is_table=False)
    else:
        set_body(p_arch_txt, [
            "Hình 1 mô tả chi tiết kiến trúc nơ-ron của GUMNetHet. Khác với các mạng MoE truyền thống nạp toàn bộ đặc trưng vào mọi expert gây suy giảm "
            "chuyên môn hóa, GUMNetHet phân vùng đặc trưng theo miền tần số và bản chất kinh tế: nhóm giá tần số cao cho CNN-1D đa tỷ lệ, "
            "nhóm vĩ mô biến đổi chậm cho GRU-Attention, và nhóm tỷ lệ crack-spread/cú sốc đuôi dày cho Wavelet-KAN. Cơ chế này ép buộc inductive bias "
            "chuyên biệt, ngăn chặn triệt để hiện tượng thoái hóa chuyên gia (expert degeneration)."
        ])
        set_caption(p_cap_fig1, "Hình 1.", "Kiến trúc GUMNetHet: phân vùng đặc trưng, ba expert dị thể, horizon-aware router và multi-quantile head.", "1", is_table=False)

    set_h2(p_h2_cnn, "3.2.1. Price Momentum Expert: Multi-Scale 1D-CNN" if is_en else "3.2.1. Chuyên gia Động lượng Giá: 1D-CNN Đa tỷ lệ")
    if is_en:
        set_body(p_cnn_txt, [
            "The price expert employs three parallel 1D convolutional layers with kernel sizes ", ('m', 'k_set'),
            " to extract multi-resolution temporal features, integrated with layer normalization and temporal attention according to (3) and (4):"
        ])
    else:
        set_body(p_cnn_txt, [
            "Chuyên gia giá sử dụng ba tầng tích chập 1D song song với kích thước kernel ", ('m', 'k_set'),
            " để trích xuất đặc trưng thời gian đa độ phân giải, kết hợp chuẩn hóa tầng và cơ chế chú ý theo thời gian theo (3) và (4):"
        ])

    set_h2(p_h2_gru, "3.2.2. Macroeconomic Regime Expert: GRU-Attention" if is_en else "3.2.2. Chuyên gia Chế độ Vĩ mô: GRU-Attention")
    if is_en:
        set_body(p_gru_txt, [
            "The macroeconomic expert captures low-frequency trend signals via a 2-layer stacked GRU network with dropout = 0.1, "
            "extracting the final hidden state representation according to (5):"
        ])
    else:
        set_body(p_gru_txt, [
            "Chuyên gia kinh tế vĩ mô xử lý các tín hiệu xu hướng tần số thấp thông qua mạng GRU 2 tầng xếp chồng với dropout = 0.1, "
            "trích xuất biểu diễn trạng thái ẩn cuối cùng theo (5):"
        ])

    set_h2(p_h2_wkan, "3.2.3. Non-Linear Shock-Absorption Expert: Wavelet-KAN" if is_en else "3.2.3. Chuyên gia Triệt tiêu Cú sốc Phi tuyến: Wavelet-KAN")
    if is_en:
        set_body(p_wkan_txt, [
            "To capture severe non-linear structural breaks induced by extreme geopolitical risk events, the third expert implements a "
            "Kolmogorov–Arnold Network integrated with Mexican Hat wavelets ", ('m', 'mexican_hat'),
            " parameterized by learnable translation and dilation coefficients according to (6) and (7):"
        ])
    else:
        set_body(p_wkan_txt, [
            "Để nắm bắt các đứt gãy phi tuyến tính nghiêm trọng do các sự kiện rủi ro địa chính trị cực đoan, chuyên gia thứ ba triển khai Mạng "
            "Kolmogorov-Arnold tích hợp sóng con Mexican Hat ", ('m', 'mexican_hat'),
            " với các tham số dịch chuyển và co giãn có thể học theo (6) và (7):"
        ])

    set_h2(p_h2_router, "3.2.4. Horizon-Aware Dynamic Gating Router" if is_en else "3.2.4. Bộ Định tuyến Cổng Động Nhận biết Chu kỳ")
    if is_en:
        set_body(p_router_txt, [
            "To dynamically modulate expert contributions across diverse forecasting horizons and varying market regimes, the gating router "
            "accepts concatenated expert representations, horizon positional embeddings, and global summary context statistics ",
            ('m', 'x_ctx'), " according to (8), (9), and (10):"
        ])
    else:
        set_body(p_router_txt, [
            "Nhằm điều tiết linh hoạt đóng góp của từng chuyên gia qua các chu kỳ dự báo và các chế độ thị trường khác nhau, bộ định tuyến nhận "
            "đầu vào là các biểu diễn chuyên gia được ghép nối, nhúng vị trí chu kỳ và thống kê tóm tắt ngữ cảnh toàn cục ",
            ('m', 'x_ctx'), " theo (8), (9) và (10):"
        ])

    set_h2(p_h2_res, "3.2.5. Residual Scaling Bounding & Multi-Quantile Prediction Head" if is_en else "3.2.5. Cơ chế Co giãn Phần dư & Đầu ra Dự báo Đa Phân vị")
    if is_en:
        set_body(p_res_txt, [
            "To prevent variance explosion at long horizons (such as H60), GUMNetHet incorporates a learnable residual scaling vector ",
            ('m', 'gamma_h'), " initialized at 0.1, outputting multi-quantile return predictions for ", ('m', 'q_set'), " according to (11):"
        ])
    else:
        set_body(p_res_txt, [
            "Để ngăn chặn sự bùng nổ phương sai ở chu kỳ dài hạn (như H60), GUMNetHet tích hợp vectơ co giãn phần dư ",
            ('m', 'gamma_h'), " được khởi tạo tại 0.1, cung cấp các dự báo lợi suất đa phân vị cho các mức ", ('m', 'q_set'), " theo (11):"
        ])

    set_h2(p_h2_loss, "3.2.6. Dual-Loss Optimization" if is_en else "3.2.6. Tối ưu hóa Hàm Mất mát Kép")
    if is_en:
        set_body(p_loss_txt, [
            "The overall objective function combines the multi-quantile pinball loss with a load-balancing regularization penalty to mitigate "
            "routing collapse according to (12), (13), and (14):"
        ])
    else:
        set_body(p_loss_txt, [
            "Hàm mục tiêu tổng thể kết hợp giữa mất mát pinball đa phân vị và điều chuẩn cân bằng tải nhằm chống sụp đổ cổng định tuyến theo (12), (13) và (14):"
        ])

    # 7. Section 4: Experimental Setup
    set_h1(p_h1_exp, "4. EXPERIMENTAL SETUP" if is_en else "4. THIẾT LẬP THỰC NGHIỆM")
    set_h2(p_h2_data, "4.1. Data and Walk-Forward Protocol" if is_en else "4.1. Dữ liệu và giao thức walk-forward")
    
    if is_en:
        set_body(p_data_txt, [
            "The dataset covers 03/11/2008–30/04/2026 with ", ('m', 'N_val'), " trading-day observations. The two reported primary targets are MG95 "
            "and DO 0.001% per Platts Singapore; exogenous covariates include inter-product Platts benchmarks, WTI, Brent, GPR ", ('ref', '1'),
            ", DXY, crude oil production, crack-spread ratios, realized volatility, and calendar features."
        ])
        set_body(p_data_wf_txt, [
            "To prevent look-ahead bias, daily variables are indexed at ", ('m', 't_minus_1'), "; GPR is lagged by 30 calendar days; crude oil production "
            "is lagged by 7 days; and rolling features are computed strictly after applying lags. All scalers are fit exclusively on the training partition "
            "at each walk-forward step. Lookback length is ", ('m', 'L_30'), "; train/validation ratio is 85/15 per expansion. Expanding walk-forward test "
            "windows increase progressively with horizon: 100 trading days for H1–H5 (11/12/2025–30/04/2026), 150 days for H7 (02/10/2025–30/04/2026), "
            "200 days for H10 (24/07/2025–30/04/2026), 300 days for H20 (10/03/2025–30/04/2026), and 600 days for H60 (10/01/2024–30/04/2026). "
            "All test windows fall entirely within a period of clustered extreme energy shocks (Red Sea crisis, Middle East escalations, "
            "Russia–Ukraine war, and OPEC+ quota shifts). Specifically, annualized realized return volatility during the short-term test window "
            "(100 days) surges to 73.04% for MG95 (1.90× historical training volatility of 38.45%) and 96.29% for DO 0.001% (2.90× training volatility "
            "of 33.16%). The Geopolitical Risk (GPR) index averages 225.66 during the test period (nearly doubling the historical mean of 114.60) and "
            "peaks at 500.81 (90th percentile at 376.48). Price swings in the H60 test window record extreme peak-to-trough variations: 70.34 to 170.52 USD/bbl "
            "(+142.4%) for MG95 and 75.90 to 292.82 USD/bbl (+285.8%) for DO 0.001%. Comprehensive train/test splits and volatility characteristics "
            "across horizons are summarized in Table 1. This configuration serves as a rigorous stress-testing benchmark, ensuring models are evaluated "
            "on their capacity to adapt under fat-tailed regime shifts rather than static market conditions."
        ])
    else:
        set_body(p_data_txt, [
            "Dữ liệu bao phủ 03/11/2008–30/04/2026 với ", ('m', 'N_val'), " quan sát ngày giao dịch. Hai target được báo cáo là MG95 và "
            "DO 0.001% theo Platts; các biến ngoại sinh gồm Platts liên sản phẩm, WTI, Brent, GPR ", ('ref', '1'),
            ", DXY, sản lượng dầu, crack-spread, realized volatility và biến lịch."
        ])
        set_body(p_data_wf_txt, [
            "Để tránh look-ahead bias, các biến ngày được dùng ở ", ('m', 't_minus_1'), "; GPR được áp lag 30 ngày lịch; sản lượng dầu áp lag 7 ngày; "
            "các biến rolling được tính sau khi lag. Mọi scaler chỉ fit trên train tại từng bước walk-forward. Lookback ", ('m', 'L_30'),
            "; train/validation=85/15 ở mỗi lần mở rộng. Cửa sổ test expanding walk-forward được thiết kế tăng dần theo horizon: 100 ngày giao dịch "
            "ở H1–H5 (11/12/2025–30/04/2026), 150 ngày ở H7 (02/10/2025–30/04/2026), 200 ngày ở H10 (24/07/2025–30/04/2026), 300 ngày ở H20 "
            "(10/03/2025–30/04/2026) và 600 ngày ở H60 (10/01/2024–30/04/2026). Toàn bộ các cửa sổ test đều nằm trọn trong giai đoạn thị trường "
            "năng lượng chịu chuỗi cú sốc địa chính trị dồn dập và phân cụm biến động cực đoan (khủng hoảng Biển Đỏ, căng thẳng Trung Đông, "
            "xung đột Nga–Ukraine và biến động hạn ngạch OPEC+). Cụ thể, độ biến động thực tế hàng năm (annualized volatility) của lợi suất trong "
            "cửa sổ test ngắn hạn (100 ngày) tăng vọt lên 73,04% đối với MG95 (gấp 1,90 lần mức 38,45% của tập train lịch sử) và 96,29% đối với "
            "DO 0.001% (gấp 2,90 lần mức 33,16% của train). Chỉ số rủi ro địa chính trị GPR trong giai đoạn test đạt trung bình 225,66 "
            "(gần gấp đôi mức trung bình toàn kỳ 114,60) và đạt đỉnh 500,81 (phân vị 90% đạt 376,48). Biên độ dao động giá trong cửa sổ test H60 "
            "ghi nhận mức chênh lệch đỉnh–đáy cực lớn: từ 70,34 đến 170,52 USD/thùng (+142,4%) đối với MG95 và từ 75,90 đến 292,82 USD/thùng "
            "(+285,8%) đối với DO 0.001%. Chi tiết phân chia tập train/test và các đặc trưng biến động theo từng horizon được tổng hợp trong Bảng 1. "
            "Thiết lập này đóng vai trò như một bài kiểm tra áp lực (stress-testing) thực thụ, bảo đảm các mô hình được đánh giá về khả năng thích ứng "
            "dưới các chế độ biến động đuôi dày (fat-tailed regimes) thay vì chỉ học trong điều kiện thị trường tĩnh."
        ])

    set_h2(p_h2_diag, "4.2. Econometric Diagnostics and Training Configuration" if is_en else "4.2. Kiểm định thống kê và cấu hình huấn luyện")
    if is_en:
        set_body(p_diag_txt, [
            "ADF and KPSS unit root tests in Table 2 provide evidence of non-stationarity in raw price levels; upon log-return transformation, "
            "ADF tests reject the unit-root null hypothesis (", ('m', 'p_val'), ") across all series. Return kurtosis is exceptionally high "
            "(213.35 for WTI and 17.51 for MG95), substantiating fat tails and extreme jumps in energy price series."
        ])
        set_caption(p_cap_tbl2, "Table 2.", "Condensed econometric statistical diagnostics for primary energy series.", "2", is_table=True)
        set_body(p_impl_txt, [
            "Implementation utilizes PyTorch 2.1 / Python 3.10 with AdamW optimizer (lr = 10⁻³, weight_decay = 10⁻⁴), "
            "CosineAnnealingLR scheduler (T_max = 60, η_min = 10⁻⁵), batch size = 64, early stopping patience = 15, and 60 training epochs."
        ])
    else:
        set_body(p_diag_txt, [
            "Kiểm định ADF và KPSS trong Bảng 2 cho bằng chứng về tính dừng ở mức giá; sau chuyển log-return, kiểm định ADF bác bỏ giả thuyết nghiệm đơn vị (",
            ('m', 'p_val'), ") trên tất cả các chuỗi. Độ nhọn (Kurtosis) của lợi suất rất cao (đạt 213,35 ở WTI và 17,51 ở MG95), khẳng định tính chất đuôi dày và các cú sốc cực đoan trong chuỗi giá năng lượng."
        ])
        set_caption(p_cap_tbl2, "Bảng 2.", "Chẩn đoán thống kê rút gọn của các chuỗi chính.", "2", is_table=True)
        set_body(p_impl_txt, [
            "Cài đặt trong bài báo sử dụng PyTorch 2.1/Python 3.10, AdamW (lr=10⁻³, weight_decay=10⁻⁴), CosineAnnealingLR (T_max=60, η_min=10⁻⁵), batch size=64, early stopping patience=15, và 60 epoch."
        ])

    set_h2(p_h2_base, "4.3. Baseline Methods and Evaluation Metrics" if is_en else "4.3. Baseline và chỉ số đánh giá")
    if is_en:
        set_body(p_base_txt, [
            "Evaluated baselines comprise six representative state-of-the-art benchmarks: PatchTST [8], iTransformer [9], TimesNet [10], "
            "DLinear [11], TFT [12], and N-HiTS [14]. Evaluation metrics include MAE, RMSE, MAPE (%), ", ('m', 'R2'), ", and Directional Accuracy (DA, %). "
            "Probabilistic forecasts are quantified using Prediction Interval Coverage Probability (PICP) and Prediction Interval Normalized "
            "Average Width (PINAW) on ", ('m', 'q_interval'), "."
        ])
    else:
        set_body(p_base_txt, [
            "Các baseline dùng để so sánh gồm 6 mô hình đại diện tiêu biểu: PatchTST [8], iTransformer [9], TimesNet [10], "
            "DLinear [11], TFT [12], và N-HiTS [14]. Các chỉ số đánh giá gồm MAE, RMSE, MAPE (%), ", ('m', 'R2'), ", và Directional Accuracy (DA, %). "
            "Đánh giá phân vị sử dụng PICP (Prediction Interval Coverage Probability) và PINAW (Prediction Interval Normalized Average Width) trên ",
            ('m', 'q_interval'), "."
        ])

    # 8. Section 5: Results and Discussion
    set_h1(p_h1_res, "5. RESULTS AND DISCUSSION" if is_en else "5. KẾT QUẢ VÀ THẢO LUẬN")
    set_h2(p_h2_point, "5.1. Multi-Horizon Point Forecasting Performance" if is_en else "5.1. Hiệu năng điểm đa horizon")
    
    if is_en:
        set_body(p_point_txt, [
            "Fig. 2 alongside Table 3 and Table 4 demonstrates that GUMNetHet consistently maintains lower MAE than all reported baselines "
            "across both products over all horizons, with pronounced advantages at H20–H60 where the test set undergoes massive price swings "
            "of up to +142.4% (MG95) and +285.8% (DO). At H60, gasoline MAE is 4.847 compared to 6.933 for the best baseline (a 30.1% reduction); "
            "for diesel, MAE is 7.066 versus 9.167 (a 22.9% reduction). However, ", ('m', 'R2'), " at H60 moderates to 0.155 (gasoline) and −0.007 (diesel), "
            "indicating that long-term performance represents superior price-level stabilization under high volatility rather than flawless trajectory forecasting."
        ])
        set_caption(p_cap_fig2, "Fig. 2.", "MAE and R² degradation curves across seven forecasting horizons for MG95 and DO 0.001% (Seed=42).", "2", is_table=False)
        set_caption(p_cap_tbl3, "Table 3.", "Detailed empirical forecasting performance on MG95 gasoline (Seed=42). MAE/RMSE in USD/bbl; DA in percent.", "3", is_table=True)
        set_caption(p_cap_tbl4, "Table 4.", "Detailed empirical forecasting performance on DO 0.001% diesel (Seed=42).", "4", is_table=True)
    else:
        set_body(p_point_txt, [
            "Hình 2 cùng Bảng 3 và Bảng 4 cho thấy GUMNetHet duy trì MAE thấp hơn nhóm baseline được báo cáo ở cả hai sản phẩm trên toàn bộ các horizon, "
            "đặc biệt vượt trội ở H20–H60 trong bối cảnh tập test trải qua biên độ dao động giá lên tới +142,4% (MG95) và +285,8% (DO). Ở H60, MAE "
            "của xăng là 4,847 so với 6,933 của baseline tốt nhất (giảm 30,1%); với dầu là 7,066 so với 9,167 (giảm 22,9%). Tuy nhiên, ",
            ('m', 'R2'), " tại H60 giảm còn 0,155 (xăng) và −0,007 (dầu), vì vậy kết quả dài hạn nên được hiểu là ổn định sai số mức giá tốt hơn "
            "baseline trong môi trường biến động mạnh, không phải dự báo quỹ đạo dài hạn hoàn hảo."
        ])
        set_caption(p_cap_fig2, "Hình 2.", "Đường cong MAE và R² qua bảy horizon cho MG95 và DO 0.001% (Seed=42).", "2", is_table=False)
        set_caption(p_cap_tbl3, "Bảng 3.", "Kết quả chi tiết trên MG95 (Seed=42). MAE/RMSE tính theo USD/thùng; giá trị DA là phần trăm.", "3", is_table=True)
        set_caption(p_cap_tbl4, "Bảng 4.", "Kết quả chi tiết trên DO 0.001% (Seed=42).", "4", is_table=True)

    set_h2(p_h2_da, "5.2. Directional Accuracy Analysis" if is_en else "5.2. Độ chính xác xu hướng")
    if is_en:
        set_body(p_da_txt, [
            "Directional accuracy in Fig. 3 reveals distinct behavior between short and long horizons. For gasoline (Table 3), GUMNetHet attains "
            "high DA between 90.95% and 95.56% at H1–H7; for diesel (Table 4), DA spans 76.65% to 84.92%. At H20, DA remains substantial "
            "(91.65% for gasoline; 71.11% for diesel). Conversely, at H10 and H60, DA drops below 50% (42.24%/27.95% for gasoline and 32.29%/19.10% for diesel). "
            "This phenomenon aligns with financial econometrics principles: over extended horizons (H60 spanning nearly three months), oil prices "
            "approximate a near random walk with exponentially compounding uncertainty, rendering directional signals noisy. Rather than taking "
            "extreme directional bets that risk catastrophic error explosion, GUMNetHet's residual scaling mechanism shrinks predictions toward "
            "robust central bounds to minimize absolute error (MAE reduced by 30.1% at H60). Consequently, GUMNetHet operates as an effective "
            "directional trading signal at short horizons (H1–H7), while transitioning into a price-level risk and uncertainty quantification "
            "tool at long horizons (H10–H60)."
        ])
        set_caption(p_cap_fig3, "Fig. 3.", "Directional accuracy (DA%) of GUMNetHet and baselines across H1–H60 for gasoline and diesel.", "3", is_table=False)
    else:
        set_body(p_da_txt, [
            "Độ chính xác hướng trong Hình 3 cho thấy một hành vi khác biệt rõ rệt giữa ngắn hạn và dài hạn. Với xăng (Bảng 3), GUMNetHet đạt DA "
            "rất cao từ 90,95% đến 95,56% ở H1–H7; với dầu (Bảng 4) tương ứng là 76,65% đến 84,92%. Ở H20, DA vẫn duy trì ở mức cao (91,65% xăng; 71,11% dầu). "
            "Tuy nhiên, tại H10 và H60, DA giảm dưới 50% (lần lượt 42,24%/27,95% cho xăng và 32,29%/19,10% cho dầu). Hiện tượng này hoàn toàn phù hợp với "
            "nguyên lý kinh tế lượng tài chính: ở horizon dài (H60 tương đương gần 3 tháng), giá dầu tiệm cận bước đi ngẫu nhiên (near random walk) với "
            "tính bất định tích lũy lũy thừa, khiến việc đoán hướng trở nên nhiễu. Thay vì phán đoán hướng cực đoan dễ dẫn đến bùng nổ sai số, "
            "cơ chế residual scaling của GUMNetHet chủ động co dự báo về vùng giá trị an toàn nhằm tối thiểu hóa sai số tuyệt đối (MAE giảm 30,1% ở H60). "
            "Do đó, GUMNetHet đóng vai trò như công cụ sinh tín hiệu giao dịch hướng ở ngắn hạn (H1–H7), và chuyển sang vai trò công cụ quản trị sai số mức giá "
            "cùng định lượng biên bất định ở dài hạn (H10–H60)."
        ])
        set_caption(p_cap_fig3, "Hình 3.", "Directional accuracy (DA%) của GUMNetHet và các baseline qua H1–H60 cho xăng và dầu.", "3", is_table=False)

    set_h2(p_h2_prob_res, "5.3. Probabilistic Forecasting, Ablation, and Router Behavior" if is_en else "5.3. Dự báo xác suất, ablation và hành vi router")
    if is_en:
        set_body(p_prob_res_txt, [
            "The ", ('m', 'q_interval'), " prediction interval achieves an empirical coverage probability of PICP=82.4% "
            "(exceeding nominal 80%) with a normalized average width of PINAW=0.142. This validates that the model avoids under-coverage "
            "while maintaining sharpness, offering practical utility for oil importers (e.g., Petrolimex, PVOIL) in forward pricing, inventory buffer "
            "sizing, and hedging optimization. Fig. 4 illustrates that prediction intervals dynamically widen during heightened market volatility. "
            "Ablation results in Table 5 indicate that replacing Wav-KAN with MLP causes the steepest degradation among expert variants; a uniform "
            "router similarly inflates MAE, corroborating the essential roles of expert specialization and adaptive routing."
        ])
        set_caption(p_cap_tbl5, "Table 5.", "Condensed ablation study of GUMNetHet variants (Seed=42).", "5", is_table=True)
        set_body(p_res_scale_txt, [
            "Empirical results reveal that removing residual scaling increases MAE by approximately 8.5%/6.3% at H20 and 14.1%/11.8% at H60 "
            "for gasoline/diesel. Router analysis in Fig. 4 indicates that under low GPR and short horizons, the CNN expert receives an average "
            "weight of ~0.48; when GPR exceeds the 90th percentile, the Wav-KAN weight surges from ~0.29 to 0.61 at medium–long horizons while "
            "the CNN weight contracts to ~0.21. These findings demonstrate that the router executes true regime-switching rather than trivial output averaging."
        ])
        set_caption(p_cap_fig4, "Fig. 4.", "Top: Multi-quantile fan chart under high volatility. Bottom: Router gating weights across low-GPR and high-GPR regimes.", "4", is_table=False)
    else:
        set_body(p_prob_res_txt, [
            "Khoảng phân vị ", ('m', 'q_interval'), " đạt tỷ lệ bao phủ thực tế PICP=82,4% (vượt mức danh định 80%) cùng độ rộng dải chuẩn hóa PINAW=0,142. "
            "Kết quả này khẳng định mô hình không bị ước lượng thiếu bất định (under-coverage), đồng thời duy trì dải dự báo sắc nét (sharpness), "
            "hỗ trợ đắc lực cho các doanh nghiệp đầu mối xăng dầu (Petrolimex, PVOIL) trong việc định giá hợp đồng kỳ hạn, thiết lập mức tồn kho đệm an toàn "
            "và tối ưu hóa chi phí phòng hộ rủi ro (hedging). Hình 4 cho thấy biên bất định tự động mở rộng tương ứng khi biến động thị trường gia tăng. "
            "Ablation trong Bảng 5 chứng minh thay Wav-KAN bằng MLP gây suy giảm lớn nhất trong các biến thể expert; router đồng nhất cũng làm MAE tăng "
            "đáng kể, củng cố vai trò của chuyên môn hóa và định tuyến thích ứng."
        ])
        set_caption(p_cap_tbl5, "Bảng 5.", "Ablation rút gọn của GUMNetHet (Seed=42).", "5", is_table=True)
        set_body(p_res_scale_txt, [
            "Kết quả thực nghiệm ghi nhận rằng loại bỏ residual scaling làm MAE tăng khoảng 8,5%/6,3% ở H20 và 14,1%/11,8% ở H60 đối với xăng/dầu. "
            "Phân tích router trong Hình 4 cho thấy trong điều kiện GPR thấp và horizon ngắn, expert CNN nhận trọng số trung bình ~0,48; "
            "khi GPR vượt phân vị 90%, trọng số Wav-KAN tăng vọt từ ~0,29 lên 0,61 ở các horizon trung và dài hạn, trong khi trọng số CNN co lại còn ~0,21. "
            "Điều này chứng tỏ router thực hiện chuyển đổi chế độ thực chất (regime-switching) chứ không chỉ đơn thuần bình quân hóa đầu ra."
        ])
        set_caption(p_cap_fig4, "Hình 4.", "Trên: fan chart đa phân vị dưới biến động mạnh. Dưới: trọng số router trong chế độ GPR thấp và GPR cao.", "4", is_table=False)

    # 9. Section 6: Conclusion
    set_h1(p_h1_concl, "6. CONCLUSION" if is_en else "6. KẾT LUẬN")
    if is_en:
        set_body(p_concl1, [
            "This paper proposed GUMNetHet, a heterogeneous mixture-of-experts architecture for multi-horizon probabilistic refined oil price "
            "forecasting. The model synergizes multi-scale 1D-CNN, GRU-Attention, and Wavelet-KAN through a horizon-aware router, combined with "
            "multi-quantile outputs and residual scaling to bolster stability over extended horizons."
        ])
        set_body(p_concl2, [
            "Expanding walk-forward experiments on MG95 and DO 0.001% demonstrate that GUMNetHet achieves the lowest MAE across all evaluated "
            "baselines throughout H1–H60, yielding MAE reductions at H60 of 30.1% for MG95 and 22.9% for DO 0.001%. Ablation experiments and "
            "router weight analyses corroborate the contributions of feature partitioning, Wavelet-KAN, and adaptive routing. However, the "
            "attenuation of ", ('m', 'R2'), " and directional accuracy at long horizons underscores that model superiority lies in price-level error containment "
            "and uncertainty quantification, rather than long-range directional forecasting."
        ])
        set_body(p_concl3, [
            "Future work will focus on multi-seed benchmarking, expanding probabilistic baseline comparisons, formal statistical significance "
            "testing for multi-step forecasts, and integrating real-time multimodal geopolitical news representations."
        ])
        set_h1(p_h1_refs, "REFERENCES")
    else:
        set_body(p_concl1, [
            "Bài báo đề xuất GUMNetHet, một kiến trúc mixture-of-experts dị thể cho dự báo xác suất đa chu kỳ giá xăng dầu thành phẩm. "
            "Mô hình kết hợp CNN-1D đa tỷ lệ, GRU-Attention và Wavelet-KAN thông qua bộ định tuyến nhận biết horizon, đồng thời tích hợp đầu ra "
            "đa phân vị và residual scaling để tăng cường độ ổn định ở các chu kỳ dài hạn."
        ])
        set_body(p_concl2, [
            "Thực nghiệm expanding walk-forward trên MG95 và DO 0.001% cho thấy GUMNetHet đạt MAE thấp nhất trong nhóm baseline được đánh giá ở toàn bộ "
            "H1–H60, với mức giảm MAE tại H60 đạt 30,1% cho MG95 và 22,9% cho DO 0.001%. Kết quả ablation và phân tích trọng số router cho thấy "
            "phân vùng đặc trưng, Wavelet-KAN và định tuyến thích ứng đều đóng góp đáng kể vào hiệu năng. Tuy nhiên, sự suy giảm của ",
            ('m', 'R2'), " và directional accuracy ở các horizon dài cho thấy ưu thế của mô hình chủ yếu nằm ở độ ổn định sai số mức giá và định lượng "
            "bất định, không nên diễn giải như khả năng dự báo hướng dài hạn đáng tin cậy."
        ])
        set_body(p_concl3, [
            "Trong tương lai, nghiên cứu sẽ tập trung vào đánh giá multi-seed, mở rộng so sánh với các baseline xác suất, kiểm định ý nghĩa thống kê "
            "chính thức cho dự báo đa bước, và tích hợp các biểu diễn tin tức địa chính trị đa phương thức theo thời gian thực."
        ])
        set_h1(p_h1_refs, "TÀI LIỆU THAM KHẢO")

    # 10. References [1] to [28]
    for p_ref in doc.paragraphs[67:]:
        if p_ref.text.strip():
            p_ref.style = 'Refs'
            p_ref.paragraph_format.line_spacing = 1.05
            p_ref.paragraph_format.space_before = Pt(0)
            p_ref.paragraph_format.space_after = Pt(2)
            p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            m_r = re.match(r'\[(\d+)\]', p_ref.text)
            if m_r:
                r_num = m_r.group(1)
                bm_name = f"ref_{r_num}"
                bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{100+int(r_num)}" w:name="{bm_name}"/>')
                bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{100+int(r_num)}"/>')
                p_ref._p.insert(0, bm_start)
                p_ref._p.append(bm_end)

    # 11. Insert Table 1
    if is_en:
        t1_data = [
            ["Horizon", "Test Window", "Test Date Range", "Ann. Vol MG95 (Train→Test)", "Ann. Vol DO (Train→Test)", "Test GPR (Mean / Max)", "Test Price Range (USD/bbl)"],
            ["H1, H3, H5", "100 days", "11/12/2025 – 30/04/2026", "38.45% → 73.04% (1.90×)", "33.16% → 96.29% (2.90×)", "225.66 / 500.81", "MG95: [70.58, 170.52]\nDO: [77.11, 292.82]"],
            ["H7", "150 days", "02/10/2025 – 30/04/2026", "38.61% → 60.75% (1.57×)", "33.20% → 80.59% (2.43×)", "197.51 / 500.81", "MG95: [70.58, 170.52]\nDO: [77.11, 292.82]"],
            ["H10", "200 days", "24/07/2025 – 30/04/2026", "38.77% → 53.60% (1.38×)", "33.31% → 70.63% (2.12×)", "184.98 / 500.81", "MG95: [70.58, 170.52]\nDO: [77.11, 292.82]"],
            ["H20", "300 days", "10/03/2025 – 30/04/2026", "38.95% → 47.24% (1.21×)", "33.35% → 60.50% (1.81×)", "184.47 / 540.16", "MG95: [70.34, 170.52]\nDO: [75.90, 292.82]"],
            ["H60", "600 days", "10/01/2024 – 30/04/2026", "39.81% → 37.86% (0.95×)", "33.93% → 46.20% (1.36×)", "165.51 / 540.16", "MG95: [70.34, 170.52]\nDO: [75.90, 292.82]"]
        ]
        cap_t1_text = "Expanding walk-forward train/test split configurations and empirical volatility characteristics across forecasting horizons."
        cap_t1_prefix = "Table 1."
    else:
        t1_data = [
            ["Horizon", "Cửa sổ Test", "Khoảng thời gian Test", "Ann. Vol MG95 (Train→Test)", "Ann. Vol DO (Train→Test)", "GPR Test (TB / Max)", "Biên độ giá Test (USD/thùng)"],
            ["H1, H3, H5", "100 ngày", "11/12/2025 – 30/04/2026", "38,45% → 73,04% (1,90×)", "33,16% → 96,29% (2,90×)", "225,66 / 500,81", "MG95: [70,58; 170,52]\nDO: [77,11; 292,82]"],
            ["H7", "150 ngày", "02/10/2025 – 30/04/2026", "38,61% → 60,75% (1,57×)", "33,20% → 80,59% (2,43×)", "197,51 / 500,81", "MG95: [70,58; 170,52]\nDO: [77,11; 292,82]"],
            ["H10", "200 ngày", "24/07/2025 – 30/04/2026", "38,77% → 53,60% (1,38×)", "33,31% → 70,63% (2,12×)", "184,98 / 500,81", "MG95: [70,58; 170,52]\nDO: [77,11; 292,82]"],
            ["H20", "300 ngày", "10/03/2025 – 30/04/2026", "38,95% → 47,24% (1,21×)", "33,35% → 60,50% (1,81×)", "184,47 / 540,16", "MG95: [70,34; 170,52]\nDO: [75,90; 292,82]"],
            ["H60", "600 ngày", "10/01/2024 – 30/04/2026", "39,81% → 37,86% (0,95×)", "33,93% → 46,20% (1,36×)", "165,51 / 540,16", "MG95: [70,34; 170,52]\nDO: [75,90; 292,82]"]
        ]
        cap_t1_text = "Cấu hình phân chia tập huấn luyện/kiểm thử expanding walk-forward và đặc trưng biến động theo từng horizon."
        cap_t1_prefix = "Bảng 1."

    p38_elem = p_data_wf_txt._p
    new_tbl = doc.add_table(rows=len(t1_data), cols=len(t1_data[0]))
    for r_idx, row in enumerate(t1_data):
        for c_idx, val in enumerate(row):
            cell = new_tbl.cell(r_idx, c_idx)
            cell.text = val
            
    style_table(new_tbl)
    
    new_caption_p = doc.add_paragraph()
    set_caption(new_caption_p, cap_t1_prefix, cap_t1_text, "1", is_table=True)
    
    p38_elem.addnext(new_caption_p._p)
    p38_elem.addnext(new_tbl._tbl)

    # 12. Format Tables (Header translation & bolding)
    if is_en:
        # Table 2: Econometric Diagnostics (doc.tables[9])
        t2_headers = ["Series", "ADF Level p", "KPSS Level p", "ADF Return p", "Kurtosis"]
        for c_idx, h in enumerate(t2_headers):
            doc.tables[9].rows[0].cells[c_idx].text = h
            
        # Table 3: MG95 (doc.tables[10])
        t3_headers = ["H", "Model", "MAE", "RMSE", "MAPE (%)", "R²", "DA (%)"]
        for c_idx, h in enumerate(t3_headers):
            doc.tables[10].rows[0].cells[c_idx].text = h

        # Table 4: DO (doc.tables[11])
        t4_headers = ["H", "Model", "MAE", "RMSE", "MAPE (%)", "R²", "DA (%)"]
        for c_idx, h in enumerate(t4_headers):
            doc.tables[11].rows[0].cells[c_idx].text = h

        # Table 5: Ablation (doc.tables[12])
        t5_headers = ["Variant", "MG95-H3 MAE", "MG95-H3 R²", "DO-H5 MAE", "DO-H5 R²"]
        for c_idx, h in enumerate(t5_headers):
            doc.tables[12].rows[0].cells[c_idx].text = h
        
        ablation_variant_map = {
            "GUMNetHet đầy đủ": "Full GUMNetHet",
            "w/o Wavelet-KAN (thay bằng MLP)": "w/o Wavelet-KAN (MLP replacement)",
            "w/o 1D-CNN": "w/o 1D-CNN",
            "w/o GRU-Attention": "w/o GRU-Attention",
            "Định tuyến đồng nhất (Không Router)": "Uniform Routing (No Router)",
            "w/o Residual Scaling": "w/o Residual Scaling"
        }
        for row in doc.tables[12].rows[1:]:
            curr_txt = row.cells[0].text.strip()
            if curr_txt in ablation_variant_map:
                row.cells[0].text = ablation_variant_map[curr_txt]

    for tbl_idx, t in enumerate(doc.tables):
        style_table(t)
        # Bold header row ONLY
        for cell in t.rows[0].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
        
        # Bold GUMNetHet row
        for row in t.rows[1:]:
            first_cell_text = row.cells[0].text.strip()
            if "GUMNetHet" in first_cell_text or "Full GUMNetHet" in first_cell_text:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.bold = True

    doc.save(out_path)
    print(f"Successfully generated clean manuscript at {out_path}!")

if __name__ == '__main__':
    # Build English version as main GUMNETHet_FAIRv4_final.docx
    build_manuscript(is_en=True, out_path="GUMNETHet_FAIRv4_final.docx")
    # Also build Vietnamese version as GUMNETHet_FAIRv4_final_VI.docx
    build_manuscript(is_en=False, out_path="GUMNETHet_FAIRv4_final_VI.docx")
