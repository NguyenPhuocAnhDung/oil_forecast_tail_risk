import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re
import os
import sys

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

from scripts.table_builder_helper import M, create_inline_omml, style_table, make_caption, add_styled_paragraph

def style_heading(p, text, is_h1=True, font_size=10.0):
    p.text = ""
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(6 if is_h1 else 4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def style_caption(p, prefix, text, tag_num, is_table=True):
    p.text = ""
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    
    r_pre = p.add_run(prefix)
    r_pre.font.name = "Times New Roman"
    r_pre.font.size = Pt(8.5)
    r_pre.font.bold = True
    r_pre.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    r_txt = p.add_run(" " + text)
    r_txt.font.name = "Times New Roman"
    r_txt.font.size = Pt(8.5)
    r_txt.font.italic = False
    r_txt.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_table else WD_ALIGN_PARAGRAPH.LEFT
    
    bm_name = f"tbl_{tag_num}" if is_table else f"fig_{tag_num}"
    bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{tag_num}" w:name="{bm_name}"/>')
    bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{tag_num}"/>')
    p._p.insert(0, bm_start)
    p._p.append(bm_end)

def style_abstract_para(p, prefix, text):
    p.text = ""
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_pre = p.add_run(prefix)
    r_pre.font.name = "Times New Roman"
    r_pre.font.size = Pt(9.0)
    r_pre.font.bold = True
    r_pre.font.italic = True
    
    r_txt = p.add_run(text)
    r_txt.font.name = "Times New Roman"
    r_txt.font.size = Pt(9.0)
    r_txt.font.italic = False

def build_perfect_manuscript_clean(target_path):
    print(f"Reading base backup document: GUMNETHet_FAIRv3 - Copy.backup.docx...")
    doc = docx.Document('GUMNETHet_FAIRv3 - Copy.backup.docx')
    
    # Save paragraph objects into exact variables BEFORE any modifications
    p_title = doc.paragraphs[0]
    p_authors = doc.paragraphs[1]
    p_abstract = doc.paragraphs[5]
    p_keywords = doc.paragraphs[6]
    
    p_h1_intro = doc.paragraphs[7]
    p_intro_contrib = doc.paragraphs[10]
    
    p_h1_related = doc.paragraphs[11]
    
    p_h1_method = doc.paragraphs[14]
    p_h2_prob = doc.paragraphs[15]
    p_form1_txt = doc.paragraphs[16]
    p_form2_txt = doc.paragraphs[17]
    p_h2_arch = doc.paragraphs[18]
    p_arch_txt = doc.paragraphs[19]
    p_dr1 = doc.paragraphs[20] # DRAWING 1
    p_cap_fig1 = doc.paragraphs[21]
    
    p_h2_cnn = doc.paragraphs[22]
    p_h2_gru = doc.paragraphs[24]
    p_h2_wkan = doc.paragraphs[26]
    p_h2_router = doc.paragraphs[28]
    p_h2_res = doc.paragraphs[30]
    p_h2_loss = doc.paragraphs[32]
    
    p_h1_exp = doc.paragraphs[35]
    p_h2_data = doc.paragraphs[36]
    p_data_txt = doc.paragraphs[37]
    p_data_wf_txt = doc.paragraphs[38]
    
    p_h2_diag = doc.paragraphs[39]
    p_diag_txt = doc.paragraphs[40]
    p_cap_tbl2 = doc.paragraphs[41]
    
    p_h2_base = doc.paragraphs[43]
    
    p_h1_res = doc.paragraphs[45]
    p_h2_point = doc.paragraphs[46]
    p_point_txt = doc.paragraphs[47]
    p_dr2 = doc.paragraphs[48] # DRAWING 2
    p_cap_fig2 = doc.paragraphs[49]
    p_cap_tbl3 = doc.paragraphs[50]
    p_cap_tbl4 = doc.paragraphs[51]
    
    p_h2_da = doc.paragraphs[52]
    p_da_txt = doc.paragraphs[53]
    p_dr3 = doc.paragraphs[54] # DRAWING 3
    p_cap_fig3 = doc.paragraphs[55]
    
    p_h2_prob_res = doc.paragraphs[56]
    p_prob_res_txt = doc.paragraphs[57]
    p_cap_tbl5 = doc.paragraphs[58]
    p_res_scale_txt = doc.paragraphs[59]
    p_dr4 = doc.paragraphs[60] # DRAWING 4
    p_cap_fig4 = doc.paragraphs[61]
    
    p_h1_concl = doc.paragraphs[62]
    p_h1_refs = doc.paragraphs[66]

    # --- 1. TITLE & AUTHORS ---
    p_title.runs[0].font.bold = True
    p_title.runs[0].font.size = Pt(14.0)
    p_authors.runs[0].font.bold = True
    p_authors.runs[0].font.size = Pt(10.0)

    # --- 2. ABSTRACT & KEYWORDS ---
    style_abstract_para(p_abstract, "Tóm tắt—", 
        "Bài báo đề xuất GUMNetHet (Heterogeneous Gated Unified Mixture Network) cho dự báo xác suất đa chu kỳ "
        "giá xăng dầu thành phẩm dưới biến động địa chính trị. Mô hình phân vùng đặc trưng thành ba nhóm và xử lý bằng "
        "các expert chuyên biệt: CNN-1D đa tỷ lệ cho động lượng giá, GRU-Attention cho trạng thái vĩ mô và Wavelet-KAN cho "
        "các quan hệ phi tuyến nhạy cú sốc. Các biểu diễn được kết hợp bởi bộ định tuyến nhận biết horizon và ngữ cảnh thị trường; "
        "đầu ra đa phân vị q ∈ {0.1, 0.5, 0.9} được tối ưu bằng pinball loss và residual scaling. Trên dữ liệu đa nguồn 11/2008–04/2026 "
        "(N = 4.512 quan sát) theo giao thức expanding walk-forward trong các giai đoạn thị trường biến động khốc liệt "
        "(độ biến động annualized volatility tập test đạt 73,04–96,29%, gấp 1,9–2,9 lần tập train lịch sử), GUMNetHet đạt MAE "
        "thấp nhất trong nhóm baseline được báo cáo ở cả bảy horizon H1–H60 cho MG95 và DO 0.001%. Tại H60, MAE giảm lần lượt 30,1% "
        "và 22,9% so với baseline tốt nhất. Độ chính xác xu hướng đạt mức cao ở H1–H7 nhưng giảm rõ ở H10 và H60, cho thấy sự khác biệt "
        "giữa độ chính xác mức giá và dự báo hướng ở horizon dài. Khoảng dự báo 80% đạt PICP=82,4% với PINAW=0,142. Kết quả ablation "
        "và phân tích router cho thấy chuyên môn hóa expert và định tuyến thích ứng đóng góp đáng kể vào hiệu năng của mô hình."
    )
    
    style_abstract_para(p_keywords, "Từ khóa—", 
        "dự báo giá năng lượng; mixture-of-experts; Wavelet-KAN; rủi ro địa chính trị; dự báo phân vị; expanding walk-forward."
    )

    # --- 3. SECTION 1: GIỚI THIỆU ---
    style_heading(p_h1_intro, "1. GIỚI THIỆU", is_h1=True, font_size=10.0)
    add_styled_paragraph(p_intro_contrib, [
        "Mô hình đề xuất GUMNetHet giải quyết điểm nghẽn này bằng kỹ thuật phân vùng đặc trưng (feature partitioning): "
        "nhóm giá và benchmark được xử lý bởi CNN-1D đa tỷ lệ; nhóm vĩ mô và chỉ số GPR bởi GRU-Attention; nhóm tỷ lệ crack-spread "
        "và độ biến động bởi Wavelet-KAN. Ba biểu diễn được hợp nhất linh hoạt thông qua một bộ định tuyến (gating router) phụ thuộc "
        "vào horizon dự báo và ngữ cảnh thị trường. Đóng góp chính của bài báo gồm: ",
        ('bold', "(i) Kiến trúc MoE dị thể "), "với phân vùng đặc trưng theo bản chất kinh tế và miền tần số; ",
        ('bold', "(ii) Bộ định tuyến nhận biết horizon "), "(horizon-aware routing); ",
        ('bold', "(iii) Đầu ra đa phân vị "), "(multi-quantile head) kết hợp residual scaling để kiểm soát độ trôi phương sai; và ",
        ('bold', "(iv) Đánh giá thực nghiệm mở rộng walk-forward "), "trên ", ('m', 'N_val'), " ngày giao dịch trực tiếp trong các giai đoạn "
        "biến động địa chính trị cao (độ biến động tập test cao gấp 1,9–2,9 lần tập train) với bảng kết quả đầy đủ các chỉ số MAE, RMSE, MAPE, ",
        ('m', 'R2'), " và DA%."
    ])

    # --- 4. SECTION 2: NGHIÊN CỨU LIÊN QUAN ---
    style_heading(p_h1_related, "2. NGHIÊN CỨU LIÊN QUAN", is_h1=True, font_size=10.0)

    # --- 5. SECTION 3: PHƯƠNG PHÁP GUMNetHet ---
    style_heading(p_h1_method, "3. PHƯƠNG PHÁP GUMNetHet", is_h1=True, font_size=10.0)
    style_heading(p_h2_prob, "3.1. Phát biểu bài toán", is_h1=False, font_size=9.5)
    
    add_styled_paragraph(p_form1_txt, [
        "Gọi ", ('m', 'X_price'), ", ", ('m', 'X_macro'), ", ", ('m', 'X_shock'), " là các ma trận đặc trưng đầu vào quá khứ với lookback ",
        ('m', 'L_30'), " ngày giao dịch. Mục tiêu dự báo là ước lượng trực tiếp vectơ lợi suất log tích lũy ", ('m', 'r_thc'),
        " qua chu kỳ ", ('m', 'h_set'), " cho nhóm sản phẩm ", ('m', 'c_set'), " được định nghĩa theo (1):"
    ])
    
    add_styled_paragraph(p_form2_txt, [
        "Mức giá dự báo tương lai ", ('m', 'P_hat_thc'), " sau đó được phục hồi chính xác thông qua phép biến đổi nghịch đảo tiền định theo (2). "
        "Thiết kế dự báo lợi suất tích lũy trực tiếp (direct cumulative) này giúp đưa chuỗi giá không dừng về chuỗi lợi suất dừng ",
        ('m', 'p_val'), " (chứng minh ở Mục 4.2), đồng thời triệt tiêu hoàn toàn hiện tượng tích lũy sai số đệ quy (autoregressive error compounding) "
        "khi dự báo qua nhiều bước thời gian."
    ])
    
    style_heading(p_h2_arch, "3.2. Khung kiến trúc hệ thống GUMNetHet", is_h1=False, font_size=9.5)
    add_styled_paragraph(p_arch_txt, [
        "Hình 1 mô tả chi tiết kiến trúc nơ-ron của GUMNetHet. Khác với các mạng MoE truyền thống nạp toàn bộ đặc trưng vào mọi expert gây suy giảm "
        "chuyên môn hóa, GUMNetHet phân vùng đặc trưng theo miền tần số và bản chất kinh tế: nhóm giá tần số cao cho CNN-1D đa tỷ lệ, "
        "nhóm vĩ mô biến đổi chậm cho GRU-Attention, và nhóm tỷ lệ crack-spread/cú sốc đuôi dày cho Wavelet-KAN. Cơ chế này ép buộc inductive bias "
        "chuyên biệt, ngăn chặn triệt để hiện tượng thoái hóa chuyên gia (expert degeneration)."
    ])
    
    style_caption(p_cap_fig1, "Hình 1.", "Kiến trúc GUMNetHet: phân vùng đặc trưng, ba expert dị thể, horizon-aware router và multi-quantile head.", "1", is_table=False)
    
    style_heading(p_h2_cnn, "3.2.1. Chuyên gia Động lượng Giá: 1D-CNN Đa tỷ lệ", is_h1=False, font_size=9.5)
    style_heading(p_h2_gru, "3.2.2. Chuyên gia Chế độ Vĩ mô: GRU-Attention", is_h1=False, font_size=9.5)
    style_heading(p_h2_wkan, "3.2.3. Chuyên gia Triệt tiêu Cú sốc Phi tuyến: Wavelet-KAN", is_h1=False, font_size=9.5)
    style_heading(p_h2_router, "3.2.4. Bộ Định tuyến Cổng Động Nhận biết Chu kỳ", is_h1=False, font_size=9.5)
    style_heading(p_h2_res, "3.2.5. Cơ chế Co giãn Phần dư & Đầu ra Dự báo Đa Phân vị", is_h1=False, font_size=9.5)
    style_heading(p_h2_loss, "3.2.6. Tối ưu hóa Hàm Mất mát Kép", is_h1=False, font_size=9.5)

    # --- 6. SECTION 4: THIẾT LẬP THỰC NGHIỆM ---
    style_heading(p_h1_exp, "4. THIẾT LẬP THỰC NGHIỆM", is_h1=True, font_size=10.0)
    style_heading(p_h2_data, "4.1. Dữ liệu và giao thức walk-forward", is_h1=False, font_size=9.5)
    
    add_styled_paragraph(p_data_txt, [
        "Dữ liệu bao phủ 03/11/2008–30/04/2026 với ", ('m', 'N_val'), " quan sát ngày giao dịch. Hai target được báo cáo là MG95 và "
        "DO 0.001% theo Platts; các biến ngoại sinh gồm Platts liên sản phẩm, WTI, Brent, GPR ", ('ref', '1'),
        ", DXY, sản lượng dầu, crack-spread, realized volatility và biến lịch."
    ])
    
    add_styled_paragraph(p_data_wf_txt, [
        "Để tránh look-ahead bias, các biến ngày được dùng ở ", ('m', 't_minus_1'), "; GPR được áp lag 30 ngày lịch; sản lượng dầu áp lag 7 ngày; "
        "các biến rolling được tính sau khi lag. Mọi scaler chỉ fit trên train tại từng bước walk-forward. Lookback ", ('m', 'L_30'),
        "; train/validation=85/15 ở mỗi lần mở rộng. Cửa sổ test expanding walk-forward được thiết kế tăng dần theo horizon: 100 ngày giao dịch "
        "ở H1–H5 (11/12/2025–30/04/2026), 150 ngày ở H7 (02/10/2025–30/04/2026), 200 ngày ở H10 (24/07/2025–30/04/2026), 300 ngày ở H20 "
        "(10/03/2025–30/04/2026) và 600 ngày ở H60 (10/01/2024–30/04/2026). Toàn bộ các cửa sổ test đều nằm trọn trong giai đoạn thị trường "
        "năng lượng chịu chuỗi cú sốc địa chính trị dồn dập và phân cụm biến động cực đoan (khủng hoảng Biển Đỏ, căng thẳng Trung Đông, "
        "xung đột Nga–Ukraine và biến động hạn ngạch OPEC+). Cụ thể, độ biến động thực tế hàng năm (annualized volatility) của lợi suất trong "
        "cửa sổ test ngắn hạn (100 ngày) tăng vọt lên 73,04% đối với MG95 (gấp 1,90 lần mức 38,45% của tập train lịch sử) và 96,29% đối với "
        "DO 0.001% (gấp 2,90 lần mức 33,16% của train). Chỉ số rủi ro địa chính trị GPR trong giai đoạn test đạt trung bình 225,66 "
        "(gần gấp đôi mức trung bình toàn kỳ 114,60) và đạt đỉnh 500,81 (phân vị 90% đạt 376,48). Biên độ dao động giá trong cửa sổ test H60 "
        "ghi nhận mức chênh lệch đỉnh–đáy cực lớn: từ 70,34 đến 170,52 USD/thùng (+142,4%) đối với MG95 và từ 75,90 đến 292,82 USD/thùng "
        "(+285,8%) đối với DO 0.001%. Chi tiết phân chia tập train/test và các đặc trưng biến động theo từng horizon được tổng hợp trong Bảng 1. "
        "Thiết lập này đóng vai trò như một bài kiểm tra áp lực (stress-testing) thực thụ, bảo đảm các mô hình được đánh giá về khả năng thích ứng "
        "dưới các chế độ biến động đuôi dày (fat-tailed regimes) thay vì chỉ học trong điều kiện thị trường tĩnh."
    ])

    style_heading(p_h2_diag, "4.2. Kiểm định thống kê và cấu hình huấn luyện", is_h1=False, font_size=9.5)
    add_styled_paragraph(p_diag_txt, [
        "Kiểm định ADF và KPSS trong Bảng 2 cho bằng chứng về tính dừng ở mức giá; sau chuyển log-return, kiểm định ADF bác bỏ giả thuyết nghiệm đơn vị (",
        ('m', 'p_val'), ") trên tất cả các chuỗi. Độ nhọn (Kurtosis) của lợi suất rất cao (đạt 213,35 ở WTI và 17,51 ở MG95), khẳng định tính chất đuôi dày và các cú sốc cực đoan trong chuỗi giá năng lượng."
    ])
    style_caption(p_cap_tbl2, "Bảng 2.", "Chẩn đoán thống kê rút gọn của các chuỗi chính.", "2", is_table=True)

    style_heading(p_h2_base, "4.3. Baseline và chỉ số đánh giá", is_h1=False, font_size=9.5)

    # --- 7. SECTION 5: KẾT QUẢ VÀ THẢO LUẬN ---
    style_heading(p_h1_res, "5. KẾT QUẢ VÀ THẢO LUẬN", is_h1=True, font_size=10.0)
    style_heading(p_h2_point, "5.1. Hiệu năng điểm đa horizon", is_h1=False, font_size=9.5)
    
    add_styled_paragraph(p_point_txt, [
        "Hình 2 cùng Bảng 3 và Bảng 4 cho thấy GUMNetHet duy trì MAE thấp hơn nhóm baseline được báo cáo ở cả hai sản phẩm trên toàn bộ các horizon, "
        "đặc biệt vượt trội ở H20–H60 trong bối cảnh tập test trải qua biên độ dao động giá lên tới +142,4% (MG95) và +285,8% (DO). Ở H60, MAE "
        "của xăng là 4,847 so với 6,933 của baseline tốt nhất (giảm 30,1%); với dầu là 7,066 so với 9,167 (giảm 22,9%). Tuy nhiên, ",
        ('m', 'R2'), " tại H60 giảm còn 0,155 (xăng) và −0,007 (dầu), vì vậy kết quả dài hạn nên được hiểu là ổn định sai số mức giá tốt hơn "
        "baseline trong môi trường biến động mạnh, không phải dự báo quỹ đạo dài hạn hoàn hảo."
    ])

    style_caption(p_cap_fig2, "Hình 2.", "Đường cong MAE và R² qua bảy horizon cho MG95 và DO 0.001% (Seed=42).", "2", is_table=False)
    style_caption(p_cap_tbl3, "Bảng 3.", "Kết quả chi tiết trên MG95 (Seed=42). MAE/RMSE tính theo USD/thùng; giá trị DA là phần trăm.", "3", is_table=True)
    style_caption(p_cap_tbl4, "Bảng 4.", "Kết quả chi tiết trên DO 0.001% (Seed=42).", "4", is_table=True)

    style_heading(p_h2_da, "5.2. Độ chính xác xu hướng", is_h1=False, font_size=9.5)
    add_styled_paragraph(p_da_txt, [
        "Độ chính xác hướng trong Hình 3 cho thấy một hành vi khác biệt rõ rệt giữa ngắn hạn và dài hạn. Với xăng (Bảng 3), GUMNetHet đạt DA "
        "rất cao từ 90,95% đến 95,56% ở H1–H7; với dầu (Bảng 4) tương ứng là 76,65% đến 84,92%. Ở H20, DA vẫn duy trì ở mức cao (91,65% xăng; 71,11% dầu). "
        "Tuy nhiên, tại H10 và H60, DA giảm dưới 50% (lần lượt 42,24%/27,95% cho xăng và 32,29%/19,10% cho dầu). Hiện tượng này hoàn toàn phù hợp với "
        "nguyên lý kinh tế lượng tài chính: ở horizon dài (H60 tương đương gần 3 tháng), giá dầu tiệm cận bước đi ngẫu nhiên (near random walk) với "
        "tính bất định tích lũy lũy thừa, khiến việc đoán hướng trở nên nhiễu. Thay vì phán đoán hướng cực đoan dễ dẫn đến bùng nổ sai số, "
        "cơ chế residual scaling của GUMNetHet chủ động co dự báo về vùng giá trị an toàn nhằm tối thiểu hóa sai số tuyệt đối (MAE giảm 30,1% ở H60). "
        "Do đó, GUMNetHet đóng vai trò như công cụ sinh tín hiệu giao dịch hướng ở ngắn hạn (H1–H7), và chuyển sang vai trò công cụ quản trị sai số mức giá "
        "cùng định lượng biên bất định ở dài hạn (H10–H60)."
    ])

    style_caption(p_cap_fig3, "Hình 3.", "Directional accuracy (DA%) của GUMNetHet và các baseline qua H1–H60 cho xăng và dầu.", "3", is_table=False)

    style_heading(p_h2_prob_res, "5.3. Dự báo xác suất, ablation và hành vi router", is_h1=False, font_size=9.5)
    add_styled_paragraph(p_prob_res_txt, [
        "Khoảng phân vị ", ('m', 'q_interval'), " đạt tỷ lệ bao phủ thực tế ", ('bold', "PICP=82,4% "), "(vượt mức danh định 80%) cùng độ rộng dải chuẩn hóa ",
        ('bold', "PINAW=0,142"), ". Kết quả này khẳng định mô hình không bị ước lượng thiếu bất định (under-coverage), đồng thời duy trì dải dự báo sắc nét (sharpness), "
        "hỗ trợ đắc lực cho các doanh nghiệp đầu mối xăng dầu (Petrolimex, PVOIL) trong việc định giá hợp đồng kỳ hạn, thiết lập mức tồn kho đệm an toàn "
        "và tối ưu hóa chi phí phòng hộ rủi ro (hedging). Hình 4 cho thấy biên bất định tự động mở rộng tương ứng khi biến động thị trường gia tăng. "
        "Ablation trong Bảng 5 chứng minh thay Wav-KAN bằng MLP gây suy giảm lớn nhất trong các biến thể expert; router đồng nhất cũng làm MAE tăng "
        "đáng kể, củng cố vai trò của chuyên môn hóa và định tuyến thích ứng."
    ])

    style_caption(p_cap_tbl5, "Bảng 5.", "Ablation rút gọn của GUMNetHet (Seed=42).", "5", is_table=True)
    style_caption(p_cap_fig4, "Hình 4.", "Trên: fan chart đa phân vị dưới biến động mạnh. Dưới: trọng số router trong chế độ GPR thấp và GPR cao.", "4", is_table=False)

    # --- 8. SECTION 6: KẾT LUẬN & TÀI LIỆU THAM KHẢO ---
    style_heading(p_h1_concl, "6. KẾT LUẬN", is_h1=True, font_size=10.0)
    style_heading(p_h1_refs, "TÀI LIỆU THAM KHẢO", is_h1=True, font_size=10.0)

    # --- 9. INSERT TABLE 1 RIGHT AFTER P_DATA_WF_TXT ---
    t1_data = [
        ["Horizon", "Cửa sổ Test", "Khoảng thời gian Test", "Ann. Vol MG95 (Train→Test)", "Ann. Vol DO (Train→Test)", "GPR Test (TB / Max)", "Biên độ giá Test (USD/thùng)"],
        ["H1, H3, H5", "100 ngày", "11/12/2025 – 30/04/2026", "38,45% → 73,04% (1,90×)", "33,16% → 96,29% (2,90×)", "225,66 / 500,81", "MG95: [70,58; 170,52]\nDO: [77,11; 292,82]"],
        ["H7", "150 ngày", "02/10/2025 – 30/04/2026", "38,61% → 60,75% (1,57×)", "33,20% → 80,59% (2,43×)", "197,51 / 500,81", "MG95: [70,58; 170,52]\nDO: [77,11; 292,82]"],
        ["H10", "200 ngày", "24/07/2025 – 30/04/2026", "38,77% → 53,60% (1,38×)", "33,31% → 70,63% (2,12×)", "184,98 / 500,81", "MG95: [70,58; 170,52]\nDO: [77,11; 292,82]"],
        ["H20", "300 ngày", "10/03/2025 – 30/04/2026", "38,95% → 47,24% (1,21×)", "33,35% → 60,50% (1,81×)", "184,47 / 540,16", "MG95: [70,34; 170,52]\nDO: [75,90; 292,82]"],
        ["H60", "600 ngày", "10/01/2024 – 30/04/2026", "39,81% → 37,86% (0,95×)", "33,93% → 46,20% (1,36×)", "165,51 / 540,16", "MG95: [70,34; 170,52]\nDO: [75,90; 292,82]"]
    ]

    p38_elem = p_data_wf_txt._p
    
    new_tbl = doc.add_table(rows=len(t1_data), cols=len(t1_data[0]))
    for r_idx, row in enumerate(t1_data):
        for c_idx, val in enumerate(row):
            cell = new_tbl.cell(r_idx, c_idx)
            cell.text = val
            
    style_table(new_tbl)
    
    new_caption_p = doc.add_paragraph()
    style_caption(new_caption_p, "Bảng 1.", "Cấu hình phân chia tập huấn luyện/kiểm thử expanding walk-forward và đặc trưng biến động theo từng horizon.", "1", is_table=True)
    
    p38_elem.addnext(new_caption_p._p)
    p38_elem.addnext(new_tbl._tbl)

    # --- 10. BOLD ALL TABLE HEADERS & BEST MODELS IN TABLES ---
    for tbl_idx, t in enumerate(doc.tables):
        style_table(t)
        # Bold header row (Row 0)
        for cell in t.rows[0].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
        
        # In Table 3 (MG95), Table 4 (DO 0.001%), Table 5 (Ablation), bold GUMNetHet model row / best results
        for row in t.rows[1:]:
            first_cell_text = row.cells[0].text.strip()
            if "GUMNetHet" in first_cell_text or "GUMNet" in first_cell_text:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.bold = True

    doc.save(target_path)
    print(f"Successfully generated clean perfect manuscript at {target_path}!")

if __name__ == '__main__':
    target = sys.argv[1] if len(sys.argv) > 1 else 'GUMNETHet_FAIRv3.docx'
    build_perfect_manuscript_clean(target)
