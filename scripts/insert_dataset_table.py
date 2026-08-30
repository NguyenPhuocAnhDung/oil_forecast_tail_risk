import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re
import os
import sys

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

from scripts.table_builder_helper import M, create_inline_omml, style_table, make_caption, add_styled_paragraph

def build_manuscript_with_dataset_table(target_path):
    print(f"Reading base template/copy...")
    doc = docx.Document('GUMNETHet_FAIRv3 - Copy.backup.docx')
    
    # 1. P[5] Abstract
    add_styled_paragraph(doc.paragraphs[5], [
        "Tóm tắt—Bài báo đề xuất GUMNetHet (Heterogeneous Gated Unified Mixture Network) cho dự báo xác suất đa chu kỳ "
        "giá xăng dầu thành phẩm dưới biến động địa chính trị. Mô hình phân vùng đặc trưng thành ba nhóm và xử lý bằng "
        "các expert chuyên biệt: CNN-1D đa tỷ lệ cho động lượng giá, GRU-Attention cho trạng thái vĩ mô và Wavelet-KAN cho "
        "các quan hệ phi tuyến nhạy cú sốc. Các biểu diễn được kết hợp bởi bộ định tuyến nhận biết horizon và ngữ cảnh thị trường; "
        "đầu ra đa phân vị ", ('m', 'q_set'), " được tối ưu bằng pinball loss và residual scaling. Trên dữ liệu đa nguồn 11/2008–04/2026 (",
        ('m', 'N_val'), " quan sát) theo giao thức expanding walk-forward trong các giai đoạn thị trường biến động khốc liệt "
        "(độ biến động annualized volatility tập test đạt 73,04–96,29%, gấp 1,9–2,9 lần tập train lịch sử), GUMNetHet đạt MAE "
        "thấp nhất trong nhóm baseline được báo cáo ở cả bảy horizon H1–H60 cho MG95 và DO 0.001%. Tại H60, MAE giảm lần lượt 30,1% "
        "và 22,9% so với baseline tốt nhất. Độ chính xác xu hướng đạt mức cao ở H1–H7 nhưng giảm rõ ở H10 và H60, cho thấy sự khác biệt "
        "giữa độ chính xác mức giá và dự báo hướng ở horizon dài. Khoảng dự báo 80% đạt PICP=82,4% với PINAW=0,142. Kết quả ablation "
        "và phân tích router cho thấy chuyên môn hóa expert và định tuyến thích ứng đóng góp đáng kể vào hiệu năng của mô hình."
    ], line_spacing=1.10, space_after=4)

    # 2. P[10] Introduction Contributions
    add_styled_paragraph(doc.paragraphs[10], [
        "Mô hình đề xuất GUMNetHet giải quyết điểm nghẽn này bằng kỹ thuật phân vùng đặc trưng (feature partitioning): "
        "nhóm giá và benchmark được xử lý bởi CNN-1D đa tỷ lệ; nhóm vĩ mô và chỉ số GPR bởi GRU-Attention; nhóm tỷ lệ crack-spread "
        "và độ biến động bởi Wavelet-KAN. Ba biểu diễn được hợp nhất linh hoạt thông qua một bộ định tuyến (gating router) phụ thuộc "
        "vào horizon dự báo và ngữ cảnh thị trường. Đóng góp chính của bài báo gồm: (i) Kiến trúc MoE dị thể với phân vùng đặc trưng "
        "theo bản chất kinh tế và miền tần số; (ii) Bộ định tuyến nhận biết horizon (horizon-aware routing); (iii) Đầu ra đa phân vị "
        "(multi-quantile head) kết hợp residual scaling để kiểm soát độ trôi phương sai; và (iv) Đánh giá thực nghiệm mở rộng walk-forward "
        "trên ", ('m', 'N_val'), " ngày giao dịch trực tiếp trong các giai đoạn biến động địa chính trị cao (độ biến động tập test cao "
        "gấp 1,9–2,9 lần tập train) với bảng kết quả đầy đủ các chỉ số MAE, RMSE, MAPE, ", ('m', 'R2'), " và DA%."
    ])

    # 3. P[37] Data Section
    add_styled_paragraph(doc.paragraphs[37], [
        "Dữ liệu bao phủ 03/11/2008–30/04/2026 với ", ('m', 'N_val'), " quan sát ngày giao dịch. Hai target được báo cáo là MG95 và "
        "DO 0.001% theo Platts; các biến ngoại sinh gồm Platts liên sản phẩm, WTI, Brent, GPR ", ('ref', '1'),
        ", DXY, sản lượng dầu, crack-spread, realized volatility và biến lịch."
    ])

    # 4. P[38] Test set & Volatility details (Refers to Bảng 1)
    add_styled_paragraph(doc.paragraphs[38], [
        "Để tránh look-ahead bias, các biến ngày được dùng ở ", ('m', 't_minus_1'), "; GPR được áp lag 30 ngày lịch; sản lượng dầu áp lag 7 ngày; "
        "các biến rolling được tính sau khi lag. Mọi scaler chỉ fit trên train tại từng bước walk-forward. Lookback ", ('m', 'L_30'),
        "; train/validation=85/15 ở mỗi lần mở rộng. Cửa sổ test expanding walk-forward được thiết kế tăng dần theo horizon: 100 ngày giao dịch "
        "ở H1–H5 (11/12/2025–30/04/2026), 150 ngày ở H7 (02/10/2025–30/04/2026), 200 ngày ở H10 (24/07/2025–30/04/2026), 300 ngày ở H20 "
        "(10/03/2025–30/04/2026) và 600 ngày ở H60 (10/01/2024–30/04/2026). Toàn bộ các cửa sổ test đều nằm trọn trong giai đoạn thị trường "
        "năng lượng chịu chuỗi cú sốc địa chính trị dồn dập và phân cụm biến động cực đoan (khủng hoảng Biển Đỏ, căng thẳng Trung Đông, "
        "xung đột Nga–Ukraine và biến động hạn ngạch OPEC+). Cụ thể, độ biến động thực tế hàng năm (annualized volatility) của lợi suất trong "
        "cửa sổ test ngắn hạn (100 ngày) tăng vọt lên 73,04% đối với MG95 (gấp 1,90 lần mức 38,45% của tập train lịch sử) và 96,29% đối với "
        "DO 0.001% (gấp 2,90 lần mức 33,16% của train). Chỉ số rủi ro địa chính trị GPR trong giai đoạn test đạt trung bình 225,66 "
        "(gần gấp đôi mức trung bình toàn kỳ 114,60) và đạt đỉnh 500,81 (phân vị 90% đạt 376,48). Biên độ dao động giá trong cửa sổ test H60 "
        "ghi nhận mức chênh lệch đỉnh–đáy cực lớn: từ 70,34 đến 170,52 USD/thùng (+142,4%) đối với MG95 và từ 75,90 đến 292,82 USD/thùng "
        "(+285,8%) đối với DO 0.001%. Chi tiết phân chia tập train/test và các đặc trưng biến động theo từng horizon được tổng hợp trong Bảng 1. "
        "Thiết lập này đóng vai trò như một bài kiểm tra áp lực (stress-testing) thực thụ, bảo đảm các mô hình được đánh giá về khả năng thích ứng "
        "dưới các chế độ biến động đuôi dày (fat-tailed regimes) thay vì chỉ học trong điều kiện thị trường tĩnh."
    ])

    # 5. Insert New Table 1 after P[38]
    t1_data = [
        ["Horizon", "Cửa sổ Test", "Khoảng thời gian Test", "Ann. Vol MG95 (Train→Test)", "Ann. Vol DO (Train→Test)", "GPR Test (TB / Max)", "Biên độ giá Test (USD/thùng)"],
        ["H1, H3, H5", "100 ngày", "11/12/2025 – 30/04/2026", "38,45% → 73,04% (1,90×)", "33,16% → 96,29% (2,90×)", "225,66 / 500,81", "MG95: [70,58; 170,52]\nDO: [77,11; 292,82]"],
        ["H7", "150 ngày", "02/10/2025 – 30/04/2026", "38,61% → 60,75% (1,57×)", "33,20% → 80,59% (2,43×)", "197,51 / 500,81", "MG95: [70,58; 170,52]\nDO: [77,11; 292,82]"],
        ["H10", "200 ngày", "24/07/2025 – 30/04/2026", "38,77% → 53,60% (1,38×)", "33,31% → 70,63% (2,12×)", "184,98 / 500,81", "MG95: [70,58; 170,52]\nDO: [77,11; 292,82]"],
        ["H20", "300 ngày", "10/03/2025 – 30/04/2026", "38,95% → 47,24% (1,21×)", "33,35% → 60,50% (1,81×)", "184,47 / 540,16", "MG95: [70,34; 170,52]\nDO: [75,90; 292,82]"],
        ["H60", "600 ngày", "10/01/2024 – 30/04/2026", "39,81% → 37,86% (0,95×)", "33,93% → 46,20% (1,36×)", "165,51 / 540,16", "MG95: [70,34; 170,52]\nDO: [75,90; 292,82]"]
    ]

    p38_elem = doc.paragraphs[38]._p
    
    new_tbl = doc.add_table(rows=len(t1_data), cols=len(t1_data[0]))
    for r_idx, row in enumerate(t1_data):
        for c_idx, val in enumerate(row):
            cell = new_tbl.cell(r_idx, c_idx)
            cell.text = val
            
    style_table(new_tbl)
    
    new_caption_p = doc.add_paragraph()
    make_caption(new_caption_p, "Bảng 1. Cấu hình phân chia tập huấn luyện/kiểm thử expanding walk-forward và đặc trưng biến động theo từng horizon.", "1", is_table=True)
    
    p38_elem.addnext(new_caption_p._p)
    p38_elem.addnext(new_tbl._tbl)

    # 6. Update Section 4.2 Diagnostics (Refers to Bảng 2)
    add_styled_paragraph(doc.paragraphs[40], [
        "Kiểm định ADF và KPSS trong Bảng 2 cho bằng chứng về tính dừng ở mức giá; sau chuyển log-return, kiểm định ADF bác bỏ giả thuyết nghiệm đơn vị (",
        ('m', 'p_val'), ") trên tất cả các chuỗi. Độ nhọn (Kurtosis) của lợi suất rất cao (đạt 213,35 ở WTI và 17,51 ở MG95), khẳng định tính chất đuôi dày và các cú sốc cực đoan trong chuỗi giá năng lượng."
    ])
    make_caption(doc.paragraphs[41], "Bảng 2. Chẩn đoán thống kê rút gọn của các chuỗi chính.", "2", is_table=True)
    
    # Table 3: MG95
    make_caption(doc.paragraphs[50], "Bảng 3. Kết quả chi tiết trên MG95 (Seed=42). MAE/RMSE tính theo USD/thùng; giá trị DA là phần trăm.", "3", is_table=True)
    
    # Table 4: DO
    make_caption(doc.paragraphs[51], "Bảng 4. Kết quả chi tiết trên DO 0.001% (Seed=42).", "4", is_table=True)
    
    # Table 5: Ablation
    make_caption(doc.paragraphs[58], "Bảng 5. Ablation rút gọn của GUMNetHet (Seed=42).", "5", is_table=True)

    # P47: "Hình 2 cùng Bảng 3 và Bảng 4 cho thấy..."
    add_styled_paragraph(doc.paragraphs[47], [
        "Hình 2 cùng Bảng 3 và Bảng 4 cho thấy GUMNetHet duy trì MAE thấp hơn nhóm baseline được báo cáo ở cả hai sản phẩm trên toàn bộ các horizon, "
        "đặc biệt vượt trội ở H20–H60 trong bối cảnh tập test trải qua biên độ dao động giá lên tới +142,4% (MG95) và +285,8% (DO). Ở H60, MAE "
        "của xăng là 4,847 so với 6,933 của baseline tốt nhất (giảm 30,1%); với dầu là 7,066 so với 9,167 (giảm 22,9%). Tuy nhiên, ",
        ('m', 'R2'), " tại H60 giảm còn 0,155 (xăng) và −0,007 (dầu), vì vậy kết quả dài hạn nên được hiểu là ổn định sai số mức giá tốt hơn "
        "baseline trong môi trường biến động mạnh, không phải dự báo quỹ đạo dài hạn hoàn hảo."
    ])

    # P53: "Với xăng (Bảng 3)... với dầu (Bảng 4)..."
    add_styled_paragraph(doc.paragraphs[53], [
        "Độ chính xác hướng trong Hình 3 cho thấy một hành vi khác với sai số mức. Với xăng (Bảng 3), GUMNetHet đạt 91,46%, 91,37%, 90,95% và 95,56% "
        "ở H1, H3, H5, H7; với dầu (Bảng 4) tương ứng là 84,92%, 76,65%, 76,88% và 83,28%. Ở H20, DA vẫn cao (91,65% xăng; 71,11% dầu). Ngược lại, "
        "H10 và H60 giảm dưới 50%, lần lượt 42,24%/27,95% cho xăng và 32,29%/19,10% cho dầu. Do đó, GUMNetHet phù hợp hơn như mô hình dự báo mức "
        "và biên bất định ở horizon dài, thay vì công cụ sinh tín hiệu hướng."
    ])

    # P57: "Ablation trong Bảng 5 cho thấy..."
    add_styled_paragraph(doc.paragraphs[57], [
        "Khoảng ", ('m', 'q_interval'), " đạt PICP=82,4% so với danh định 80% và PINAW=0,142. Hình 4 cho thấy biên bất định mở rộng khi biến động tăng. "
        "Ablation trong Bảng 5 cho thấy thay Wav-KAN bằng MLP gây suy giảm lớn nhất trong các biến thể expert; router đồng nhất cũng làm MAE tăng "
        "đáng kể, củng cố vai trò của chuyên môn hóa và định tuyến thích ứng."
    ])

    # Figure 1: P[21]
    make_caption(doc.paragraphs[21], "Hình 1. Kiến trúc GUMNetHet: phân vùng đặc trưng, ba expert dị thể, horizon-aware router và multi-quantile head.", "1", is_table=False)
    # Figure 2: P[49]
    make_caption(doc.paragraphs[49], "Hình 2. Đường cong MAE và R² qua bảy horizon cho MG95 và DO 0.001% (Seed=42).", "2", is_table=False)
    # Figure 3: P[55]
    make_caption(doc.paragraphs[55], "Hình 3. Directional accuracy (DA%) của GUMNetHet và các baseline qua H1–H60 cho xăng và dầu.", "3", is_table=False)
    # Figure 4: P[61]
    make_caption(doc.paragraphs[61], "Hình 4. Trên: fan chart đa phân vị dưới biến động mạnh. Dưới: trọng số router trong chế độ GPR thấp và GPR cao.", "4", is_table=False)

    doc.save(target_path)
    print(f"Successfully generated {target_path}!")

if __name__ == '__main__':
    target = sys.argv[1] if len(sys.argv) > 1 else 'GUMNETHet_FAIRv3.docx'
    build_manuscript_with_dataset_table(target)
