# -*- coding: utf-8 -*-
import docx
import zipfile
import re
import sys
from xml.etree import ElementTree

# Reconfigure stdout to support Vietnamese character printing on Windows
sys.stdout.reconfigure(encoding='utf-8')

DOC_PATH = 'docs/Bản_thảo_GUMNET_v3.docx'

def check_xml_colors_are_black(docx_path):
    print("Checking XML colors...")
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            xml_content = z.read('word/document.xml')
            root = ElementTree.fromstring(xml_content)
            
            # namespaces
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            color_elements = root.findall('.//w:color', ns)
            
            non_black = []
            for el in color_elements:
                val = el.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if val is not None and val.upper() not in ['000000', 'AUTO']:
                    non_black.append(val)
                    
            if non_black:
                print(f"FAIL: Found {len(non_black)} non-black color tags in XML: {set(non_black)}")
                return False
            print("PASS: All w:color elements are black (000000 or AUTO).")
            return True
    except Exception as e:
        print(f"FAIL: Error reading XML colors: {e}")
        return False

def check_no_placeholders(docx_path):
    print("Checking for 'XEM LẠI' placeholders...")
    try:
        doc = docx.Document(docx_path)
        placeholders_found = 0
        
        # Check paragraphs
        for i, p in enumerate(doc.paragraphs):
            if 'XEM LẠI' in p.text:
                print(f"  Found 'XEM LẠI' in paragraph {i}: {p.text[:100]}...")
                placeholders_found += 1
                
        # Check tables
        for t_idx, t in enumerate(doc.tables):
            for r_idx, r in enumerate(t.rows):
                for c_idx, cell in enumerate(r.cells):
                    if 'XEM LẠI' in cell.text:
                        print(f"  Found 'XEM LẠI' in Table {t_idx}, Row {r_idx}, Col {c_idx}: {cell.text[:100]}...")
                        placeholders_found += 1
                        
        if placeholders_found > 0:
            print(f"FAIL: Found {placeholders_found} 'XEM LẠI' placeholders in the document.")
            return False
        print("PASS: No 'XEM LẠI' placeholders found.")
        return True
    except Exception as e:
        print(f"FAIL: Error checking placeholders: {e}")
        return False

def check_ablation_rows_count(docx_path):
    print("Checking Ablation table rows count...")
    try:
        doc = docx.Document(docx_path)
        for t_idx, t in enumerate(doc.tables):
            if len(t.rows) > 0 and 'Mô hình Ablation' in t.cell(0, 0).text:
                rows_count = len(t.rows)
                print(f"  Ablation table found (Table index {t_idx}) with {rows_count} rows.")
                if rows_count == 9:
                    print("PASS: Ablation table has exactly 9 rows (1 header + 8 variants).")
                    return True
                else:
                    print(f"FAIL: Ablation table has {rows_count} rows, expected 9.")
                    return False
        print("FAIL: Ablation table not found in document.")
        return False
    except Exception as e:
        print(f"FAIL: Error checking ablation rows: {e}")
        return False

def check_table3_splits(docx_path):
    print("Checking Table 3 (Data Allocation) splits and seq_len...")
    try:
        doc = docx.Document(docx_path)
        t3 = None
        for t in doc.tables:
            if len(t.rows) > 0 and 'Chân trời (H)' in t.cell(0, 0).text:
                t3 = t
                break
        if t3 is None:
            print("FAIL: Table 3 (Data Allocation) not found.")
            return False
            
        # Check split headers
        row0_texts = [cell.text.strip() for cell in t3.rows[0].cells]
        required_headers = ['Train (70%)', 'Validation (15%)', 'Test (15%)']
        for h in required_headers:
            if h not in row0_texts:
                print(f"FAIL: Header '{h}' not found in Table 3. Found headers: {row0_texts}")
                return False
                
        # Check seq_len column and values
        if 'seq_len' not in row0_texts:
            print(f"FAIL: 'seq_len' column not found in Table 3. Found headers: {row0_texts}")
            return False
            
        seq_len_idx = row0_texts.index('seq_len')
        expected_seq_lens = {
            'H1': '10',
            'H3': '20',
            'H5': '30',
            'H10': '60',
            'H60': '180'
        }
        for r_idx in range(1, len(t3.rows)):
            row = t3.rows[r_idx]
            h_val = row.cells[0].text.strip()
            seq_val = row.cells[seq_len_idx].text.strip()
            
            # Find which horizon this row is for
            for h_key, expected_val in expected_seq_lens.items():
                if h_val.split(' ')[0] == h_key:
                    if seq_val != expected_val:
                        print(f"FAIL: Expected seq_len {expected_val} for {h_key}, but found '{seq_val}'.")
                        return False
                        
        print("PASS: Table 3 has correct split headers and seq_len values.")
        return True
    except Exception as e:
        print(f"FAIL: Error checking Table 3 splits: {e}")
        return False

def check_dm_test_tables(docx_path):
    print("Checking Bảng 9a and 9b (DM Test) presence and coverage...")
    try:
        doc = docx.Document(docx_path)
        t9a = None
        t9b = None
        for t in doc.tables:
            if len(t.rows) > 0:
                cell_text = t.cell(0, 0).text.strip()
                if 'Mục tiêu' in cell_text and 'GUM-Net DA' in [c.text.strip() for c in t.rows[0].cells]:
                    t9a = t
                elif 'Chân trời (H)' in cell_text and 'DM Stat (Xăng)' in [c.text.strip() for c in t.rows[0].cells]:
                    t9b = t
                    
        if t9a is None:
            print("FAIL: Bảng 9a (Directional Accuracy DM Test) not found.")
            return False
        if t9b is None:
            print("FAIL: Bảng 9b (Forecasting MSE DM Test) not found.")
            return False
            
        # Check that t9a covers all 5 horizons (H1, H3, H5, H10, H60) for both targets (XANG, DAU)
        da_horizons = [t9a.cell(r, 1).text.strip() for r in range(1, len(t9a.rows))]
        expected_horizons = ['H1', 'H3', 'H5', 'H10', 'H60']
        for h in expected_horizons:
            count = da_horizons.count(h)
            if count != 2:
                print(f"FAIL: Expected horizon '{h}' to appear 2 times in Bảng 9a, but found {count} times.")
                return False
                
        # Check that t9b covers all 5 horizons
        mse_horizons = [t9b.cell(r, 0).text.strip() for r in range(1, len(t9b.rows))]
        for h in expected_horizons:
            if h not in mse_horizons:
                print(f"FAIL: Horizon '{h}' not found in Bảng 9b.")
                return False
                
        print("PASS: Bảng 9a and 9b both exist and cover all 5 horizons.")
        return True
    except Exception as e:
        print(f"FAIL: Error checking DM test tables: {e}")
        return False

def main():
    print("=" * 60)
    print(" RUNNING FINAL DOCUMENT VALIDATION ")
    print("=" * 60)
    
    # Check if document exists
    import os
    if not os.path.exists(DOC_PATH):
        print(f"FAIL: Output document {DOC_PATH} does not exist!")
        sys.exit(1)
        
    try:
        # Check opening
        doc = docx.Document(DOC_PATH)
        print("PASS: File opens correctly in python-docx.")
    except Exception as e:
        print(f"FAIL: Cannot open file: {e}")
        sys.exit(1)
        
    c1 = check_xml_colors_are_black(DOC_PATH)
    c2 = check_no_placeholders(DOC_PATH)
    c3 = check_ablation_rows_count(DOC_PATH)
    c4 = check_table3_splits(DOC_PATH)
    c5 = check_dm_test_tables(DOC_PATH)
    
    print("=" * 60)
    if c1 and c2 and c3 and c4 and c5:
        print(" SUCCESS: ALL VALIDATION CHECKS PASSED!")
        print("=" * 60)
        sys.exit(0)
    else:
        print(" FAILURE: SOME VALIDATION CHECKS FAILED!")
        print("=" * 60)
        sys.exit(1)

if __name__ == '__main__':
    main()
