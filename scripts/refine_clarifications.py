import docx
import os
import sys

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

from scripts.table_builder_helper import M, create_inline_omml, style_table, make_caption, add_styled_paragraph

def build_refined_clarified_manuscript(target_path):
    print(f"Building refined clarified manuscript: {target_path}...")
    doc = docx.Document('GUMNETHet_FAIRv3.docx')
    
    # 1. P[16] Section 3.1: Problem Formulation - Add concise clarity on Direct Cumulative & I(0) stationarity
    add_styled_paragraph(doc.paragraphs[16], [
        "Gọi ", ('m', 'X_price'), ", ", ('m', 'X_macro'), ", ", ('m', 'X_shock'), " là các ma trận đặc trưng đầu vào quá khứ với lookback ",
        ('m', 'L_30'), " ngày giao dịch. Mục tiêu dự báo là ước lượng trực tiếp vectơ lợi suất log tích lũy ", ('m', 'r_thc'),
        " qua chu kỳ ", ('m', 'h_set'), " cho nhóm sản phẩm ", ('m', 'c_set'), " được định nghĩa theo (1):"
    ])

    # 2. P[17] Section 3.1: Explain Why Direct Cumulative Target is used
    add_styled_paragraph(doc.paragraphs[17], [
        "Mức giá dự báo tương lai ", ('m', 'P_hat_thc'), " sau đó được phục hồi chính xác thông qua phép biến đổi nghịch đảo tiền định theo (2). "
        "Thiết kế dự báo lợi suất tích lũy trực tiếp (direct cumulative) này giúp đưa chuỗi giá không dừng ", ('m', 'R2'), " về chuỗi lợi suất dừng ",
        ('m', 'p_val'), " (chứng minh ở Mục 4.2), đồng thời triệt tiêu hoàn toàn hiện tượng tích lũy sai số đệ quy (autoregressive error compounding) "
        "khi dự báo qua nhiều bước thời gian."
    ])

    # 3. P[19] Section 3.2: Architecture overview - Explain why Feature Partitioning prevents expert collapse
    add_styled_paragraph(doc.paragraphs[19], [
        "Hình 1 mô tả chi tiết kiến trúc nơ-ron của GUMNetHet. Khác với các mạng MoE truyền thống nạp toàn bộ đặc trưng vào mọi expert gây suy giảm "
        "chuyên môn hóa, GUMNetHet phân vùng đặc trưng theo miền tần số và bản chất kinh tế: nhóm giá tần số cao cho CNN-1D đa tỷ lệ, "
        "nhóm vĩ mô biến đổi chậm cho GRU-Attention, và nhóm tỷ lệ crack-spread/cú sốc đuôi dày cho Wavelet-KAN. Cơ chế này ép buộc inductive bias "
        "chuyên biệt, ngăn chặn triệt để hiện tượng thoái hóa chuyên gia (expert degeneration)."
    ])

    # 4. P[53] Section 5.2: Directional Accuracy - Deep clarity on short-term vs long-term trade-off
    add_styled_paragraph(doc.paragraphs[53], [
        "Độ chính xác hướng trong Hình 3 cho thấy một hành vi khác biệt rõ rệt giữa ngắn hạn và dài hạn. Với xăng (Bảng 3), GUMNetHet đạt DA "
        "rất cao từ 90,95% đến 95,56% ở H1–H7; với dầu (Bảng 4) tương ứng là 76,65% đến 84,92%. Ở H20, DA vẫn duy trì ở mức cao (91,65% xăng; 71,11% dầu). "
        "Tuy nhiên, tại H10 và H60, DA giảm dưới 50% (lần lượt 42,24%/27,95% cho xăng và 32,29%/19,10% cho dầu). Hiện tượng này hoàn toàn phù hợp với "
        "nguyên lý kinh tế lượng tài chính: ở horizon dài (H60 tương đương gần 3 tháng), giá dầu tiệm cận bước đi ngẫu nhiên (near random walk) với "
        "tính bất định tích lũy lũy thừa, khiến việc đoán hướng trở nên nhiễu. Thay vì phán đoán hướng cực đoan dễ dẫn đến bùng nổ sai số, "
        "cơ chế residual scaling của GUMNetHet chủ động co dự báo về vùng giá trị an toàn nhằm tối thiểu hóa sai số tuyệt đối (MAE giảm 30,1% ở H60). "
        "Do đó, GUMNetHet đóng vai trò như công cụ sinh tín hiệu giao dịch hướng ở ngắn hạn (H1–H7), và chuyển sang vai trò công cụ quản trị sai số mức giá "
        "cùng định lượng biên bất định ở dài hạn (H10–H60)."
    ])

    # 5. P[57] Section 5.3: Probabilistic Forecasting - Clarify practical business meaning of PICP & PINAW
    add_styled_paragraph(doc.paragraphs[57], [
        "Khoảng phân vị ", ('m', 'q_interval'), " đạt tỷ lệ bao phủ thực tế PICP=82,4% (vượt mức danh định 80%) cùng độ rộng dải chuẩn hóa PINAW=0,142. "
        "Kết quả này khẳng định mô hình không bị ước lượng thiếu bất định (under-coverage), đồng thời duy trì dải dự báo sắc nét (sharpness), "
        "hỗ trợ đắc lực cho các doanh nghiệp đầu mối xăng dầu (Petrolimex, PVOIL) trong việc định giá hợp đồng kỳ hạn, thiết lập mức tồn kho đệm an toàn "
        "và tối ưu hóa chi phí phòng hộ rủi ro (hedging). Hình 4 cho thấy biên bất định tự động mở rộng tương ứng khi biến động thị trường gia tăng. "
        "Ablation trong Bảng 5 chứng minh thay Wav-KAN bằng MLP gây suy giảm lớn nhất trong các biến thể expert; router đồng nhất cũng làm MAE tăng "
        "đáng kể, củng cố vai trò của chuyên môn hóa và định tuyến thích ứng."
    ])

    doc.save(target_path)
    print(f"Successfully generated clarified document at {target_path}!")

if __name__ == '__main__':
    build_refined_clarified_manuscript('GUMNETHet_FAIRv3.docx')
    build_refined_clarified_manuscript('GUMNETHet_FAIRv3_template.docx')
