import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import sys
import re

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

from scripts.table_builder_helper import M, create_inline_omml, style_table

def set_para_heading(p, text, style_name='H1x', outline_lvl=0):
    p.text = ""
    p.style = style_name
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(6 if outline_lvl == 0 else 4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add outline level XML if needed
    pPr = p._p.get_or_add_pPr()
    # Check if outlineLvl exists
    existing_lvl = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
    if existing_lvl is None:
        pPr.append(parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{outline_lvl}"/>'))
    else:
        existing_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(outline_lvl))
        
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.0 if outline_lvl == 0 else 9.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_para_caption(p, prefix, text, tag_num, is_table=True):
    p.text = ""
    p.style = 'Caption'
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_table else WD_ALIGN_PARAGRAPH.LEFT
    
    r_pre = p.add_run(prefix)
    r_pre.font.name = "Times New Roman"
    r_pre.font.size = Pt(8.5)
    r_pre.font.bold = True
    r_pre.font.color.rgb = RGBColor(0, 0, 0)
    
    r_txt = p.add_run(" " + text)
    r_txt.font.name = "Times New Roman"
    r_txt.font.size = Pt(8.5)
    r_txt.font.bold = False
    r_txt.font.italic = False
    r_txt.font.color.rgb = RGBColor(0, 0, 0)
    
    bm_name = f"tbl_{tag_num}" if is_table else f"fig_{tag_num}"
    bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{tag_num}" w:name="{bm_name}"/>')
    bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{tag_num}"/>')
    p._p.insert(0, bm_start)
    p._p.append(bm_end)

def set_para_body(p, segments, line_spacing=1.15, space_after=4, is_en=True):
    p.text = ""
    p.style = 'Normal'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    
    pat = r'(Table \d+|Fig\. \d+|\[\d+\])' if is_en else r'(Bảng \d+|Hình \d+|\[\d+\])'
    
    for seg in segments:
        if isinstance(seg, str):
            parts = re.split(pat, seg)
            for part in parts:
                if not part:
                    continue
                m_tbl = re.match(r'(?:Table|Bảng) (\d+)', part)
                m_fig = re.match(r'(?:Fig\.|Hình) (\d+)', part)
                m_ref = re.match(r'\[(\d+)\]', part)
                if m_tbl:
                    t_id = f"tbl_{m_tbl.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{t_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                elif m_fig:
                    f_id = f"fig_{m_fig.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{f_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                elif m_ref:
                    r_id = f"ref_{m_ref.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                else:
                    run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9.5)
                    run.font.bold = False
        elif isinstance(seg, tuple) and seg[0] == 'm':
            m_xml = M[seg[1]]
            p._p.append(parse_xml(create_inline_omml(m_xml)))
        elif isinstance(seg, tuple) and seg[0] == 'ref':
            r_id = f"ref_{seg[1]}"
            hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                           f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                           f'<w:t>[{seg[1]}]</w:t></w:r></w:hyperlink>')
            p._p.append(hl)

def set_para_abstract(p, prefix, text_fragments):
    p.text = ""
    p.style = 'Abstract'
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_pre = p.add_run(prefix)
    r_pre.font.name = "Times New Roman"
    r_pre.font.size = Pt(9.0)
    r_pre.font.bold = True
    r_pre.font.italic = True
    
    for item in text_fragments:
        if isinstance(item, str):
            r = p.add_run(item)
            r.font.name = "Times New Roman"
            r.font.size = Pt(9.0)
            r.font.bold = False
            r.font.italic = False
        elif isinstance(item, tuple) and item[0] == 'm':
            m_xml = M[item[1]]
            p._p.append(parse_xml(create_inline_omml(m_xml)))

def build_english_doc(target_path):
    print(f"Building English document: {target_path}...")
    doc = docx.Document('GUMNETHet_FAIRv3 - Copy.backup.docx')
    
    # Map paragraph objects
    p_title = doc.paragraphs[0]
    p_authors = doc.paragraphs[1]
    p_affil1 = doc.paragraphs[2]
    p_affil2 = doc.paragraphs[3]
    p_corr = doc.paragraphs[4]
    
    p_abstract = doc.paragraphs[5]
    p_keywords = doc.paragraphs[6]
    
    p_h1_intro = doc.paragraphs[7]
    p_intro1 = doc.paragraphs[8]
    p_intro2 = doc.paragraphs[9]
    p_intro_contrib = doc.paragraphs[10]
    
    p_h1_related = doc.paragraphs[11]
    p_related1 = doc.paragraphs[12]
    p_related2 = doc.paragraphs[13]
    
    p_h1_method = doc.paragraphs[14]
    p_h2_prob = doc.paragraphs[15]
    p_form1_txt = doc.paragraphs[16]
    p_form2_txt = doc.paragraphs[17]
    p_h2_arch = doc.paragraphs[18]
    p_arch_txt = doc.paragraphs[19]
    p_dr1 = doc.paragraphs[20] # DRAWING 1
    p_cap_fig1 = doc.paragraphs[21]
    
    p_h2_cnn = doc.paragraphs[22]
    p_cnn_txt = doc.paragraphs[23]
    p_h2_gru = doc.paragraphs[24]
    p_gru_txt = doc.paragraphs[25]
    p_h2_wkan = doc.paragraphs[26]
    p_wkan_txt = doc.paragraphs[27]
    p_h2_router = doc.paragraphs[28]
    p_router_txt = doc.paragraphs[29]
    p_h2_res = doc.paragraphs[30]
    p_res_txt = doc.paragraphs[31]
    p_h2_loss = doc.paragraphs[32]
    p_loss_txt = doc.paragraphs[33]
    
    p_h1_exp = doc.paragraphs[35]
    p_h2_data = doc.paragraphs[36]
    p_data_txt = doc.paragraphs[37]
    p_data_wf_txt = doc.paragraphs[38]
    
    p_h2_diag = doc.paragraphs[39]
    p_diag_txt = doc.paragraphs[40]
    p_cap_tbl2 = doc.paragraphs[41]
    p_impl_txt = doc.paragraphs[42]
    
    p_h2_base = doc.paragraphs[43]
    p_base_txt = doc.paragraphs[44]
    
    p_h1_res = doc.paragraphs[45]
    p_h2_point = doc.paragraphs[46]
    p_point_txt = doc.paragraphs[47]
    p_dr2 = doc.paragraphs[48] # DRAWING 2
    p_cap_fig2 = doc.paragraphs[49]
    p_cap_tbl3 = doc.paragraphs[50]
    p_cap_tbl4 = doc.paragraphs[51]
    
    p_h2_da = doc.paragraphs[52]
    p_da_txt = doc.paragraphs[53]
    p_dr3 = doc.paragraphs[54] # DRAWING 3
    p_cap_fig3 = doc.paragraphs[55]
    
    p_h2_prob_res = doc.paragraphs[56]
    p_prob_res_txt = doc.paragraphs[57]
    p_cap_tbl5 = doc.paragraphs[58]
    p_res_scale_txt = doc.paragraphs[59]
    p_dr4 = doc.paragraphs[60] # DRAWING 4
    p_cap_fig4 = doc.paragraphs[61]
    
    p_h1_concl = doc.paragraphs[62]
    p_concl1 = doc.paragraphs[63]
    p_concl2 = doc.paragraphs[64]
    p_concl3 = doc.paragraphs[65]
    p_h1_refs = doc.paragraphs[66]

    # --- 1. TITLE & AUTHORS ---
    p_title.text = "ROBUST PROBABILISTIC ENERGY FORECASTING UNDER GEOPOLITICAL SHOCKS: AN ADAPTIVE MIXTURE OF LOCAL-GLOBAL EXPERTS"
    p_title.style = 'PaperTitle'
    p_title.runs[0].font.name = "Times New Roman"
    p_title.runs[0].font.bold = True
    p_title.runs[0].font.size = Pt(14.0)
    
    p_authors.text = "Phuoc Anh Dung Nguyen¹, Danh Huong Bui¹*, Van Quy Hoang²"
    p_authors.style = 'Authors'
    p_authors.runs[0].font.name = "Times New Roman"
    p_authors.runs[0].font.bold = True
    p_authors.runs[0].font.size = Pt(10.0)
    
    p_affil1.text = "¹Faculty of Information Technology, Ho Chi Minh City University of Technology (HUTECH), Ho Chi Minh City, Vietnam"
    p_affil1.style = 'Affil'
    p_affil2.text = "²Faculty of Information Technology, Thuyloi University (TLU), Hanoi, Vietnam"
    p_affil2.style = 'Affil'
    p_corr.text = "*Corresponding author: bd.huong@hutech.edu.vn"
    p_corr.style = 'Affil'

    # --- 2. ABSTRACT & KEYWORDS ---
    set_para_abstract(p_abstract, "Abstract—", [
        "This paper proposes GUMNetHet (Heterogeneous Gated Unified Mixture Network) for multi-horizon probabilistic "
        "forecasting of refined oil prices under geopolitical volatility. The framework partitions input features into three "
        "distinct subsets and processes them via specialized domain experts: multi-scale 1D-CNN for price momentum, "
        "GRU-Attention for macroeconomic regimes, and Wavelet-KAN for non-linear shock-sensitive dynamics. These representations "
        "are dynamically fused by a horizon-aware and market-context gating router; multi-quantile outputs ", ('m', 'q_set'),
        " are optimized using pinball loss combined with residual scaling. On multi-source data spanning 11/2008–04/2026 (",
        ('m', 'N_val'), " observations) under an expanding walk-forward protocol during severe market turmoil (annualized test set "
        "volatility reaching 73.04–96.29%, 1.9–2.9× historical training levels), GUMNetHet achieves the lowest MAE across all "
        "reported baselines over all seven horizons H1–H60 for both MG95 gasoline and DO 0.001% diesel. At H60, MAE is reduced "
        "by 30.1% and 22.9% relative to the strongest baseline, respectively. Directional accuracy remains high at H1–H7 but declines "
        "noticeably at H10 and H60, highlighting the structural dichotomy between price-level precision and long-term directional forecasting. "
        "The 80% prediction interval attains PICP=82.4% with PINAW=0.142. Ablation studies and router analysis confirm that expert "
        "specialization and adaptive routing contribute substantially to model performance."
    ])
    
    set_para_abstract(p_keywords, "Keywords—", [
        "energy price forecasting; mixture-of-experts; Wavelet-KAN; geopolitical risk; quantile forecasting; expanding walk-forward."
    ])

    # --- 3. SECTION 1: INTRODUCTION ---
    set_para_heading(p_h1_intro, "1. INTRODUCTION", style_name='H1x', outline_lvl=0)
    
    set_para_body(p_intro1, [
        "The Vietnamese refined oil market exhibits distinctive structural supply characteristics, being heavily dependent "
        "on domestic refinery output and international refined imports pegged to Platts Singapore benchmarks. The price adjustment "
        "mechanism has evolved progressively toward shorter revision cycles (from 30 days, to 15 days, 10 days, and currently 7 days), "
        "escalating the operational demand for highly accurate multi-horizon price projections to safeguard domestic supply security "
        "and optimize commercial hedging."
    ], is_en=True)
    
    set_para_body(p_intro2, [
        "Modern time series architectures such as PatchTST [8], iTransformer [9], TimesNet [10], and Mamba [15] demonstrate remarkable "
        "capabilities in representation learning. Nevertheless, they frequently confront significant challenges in energy markets "
        "characterized by structural breaks, fat tails, and regime shifts triggered by geopolitical risk events (such as the Russia–Ukraine "
        "conflict or Red Sea disruptions). Furthermore, homogeneous neural networks often encounter difficulty in simultaneously "
        "handling short-term price momentum, long-term macroeconomic trends, and abrupt non-linear shocks."
    ], is_en=True)
    
    set_para_body(p_intro_contrib, [
        "The proposed GUMNetHet model addresses this bottleneck via heterogeneous feature partitioning: price and benchmark features "
        "are processed by a multi-scale 1D-CNN; macroeconomic and GPR indices by GRU-Attention; and crack-spread ratios alongside "
        "realized volatility by Wavelet-KAN. These three representations are flexibly aggregated through a gating router conditioned "
        "on the forecasting horizon and global market context. The principal contributions of this paper are fourfold: "
        "(i) A heterogeneous MoE architecture with feature partitioning grounded in economic principles and frequency domains; "
        "(ii) A horizon-aware routing mechanism (horizon-aware routing); "
        "(iii) A multi-quantile prediction head (multi-quantile head) combined with residual scaling to mitigate variance drift; and "
        "(iv) Extensive expanding walk-forward empirical evaluations on ", ('m', 'N_val'), " trading days directly within "
        "high geopolitical volatility regimes (test set volatility 1.9–2.9× training levels) with comprehensive reporting across "
        "MAE, RMSE, MAPE, ", ('m', 'R2'), ", and DA% metrics."
    ], is_en=True)

    # --- 4. SECTION 2: RELATED WORK ---
    set_para_heading(p_h1_related, "2. RELATED WORK", style_name='H1x', outline_lvl=0)
    
    set_para_body(p_related1, [
        "Kilian [2] and Baumeister & Kilian [3] established foundational insights into decomposing structural shocks in global crude oil "
        "markets. Caldara & Iacoviello [1] introduced the Geopolitical Risk (GPR) index, establishing a standardized quantitative metric "
        "for measuring external political tensions."
    ], is_en=True)
    
    set_para_body(p_related2, [
        "Regarding MoE, Jacobs et al. [20] introduced adaptive mixtures of local experts, while Shazeer et al. [21] and Fedus et al. [22] "
        "scaled MoE to large-scale deep neural networks. In representation learning, Kolmogorov–Arnold Networks (KAN) [23] and Wav-KAN [24] "
        "offer compelling non-linear approximation capabilities through learnable spline and continuous wavelet basis functions. "
        "Foundation time series models (Chronos [16], MOIRAI [17], UniTS [18], Time-MoE [26]) provide powerful generalized representations; "
        "however, energy forecasting under extreme geopolitical shocks still requires architectural inductive biases tailored specifically "
        "to market mechanics."
    ], is_en=True)

    # --- 5. SECTION 3: PROPOSED METHODOLOGY: GUMNetHet ---
    set_para_heading(p_h1_method, "3. PROPOSED METHODOLOGY: GUMNetHet", style_name='H1x', outline_lvl=0)
    set_para_heading(p_h2_prob, "3.1. Problem Formulation", style_name='H2x', outline_lvl=1)
    
    set_para_body(p_form1_txt, [
        "Let ", ('m', 'X_price'), ", ", ('m', 'X_macro'), ", ", ('m', 'X_shock'), " denote past input feature matrices with lookback length ",
        ('m', 'L_30'), " trading days. The forecasting objective is to directly estimate the cumulative log-return vector ", ('m', 'r_thc'),
        " across horizons ", ('m', 'h_set'), " for product basket ", ('m', 'c_set'), ", defined in (1):"
    ], is_en=True)
    
    set_para_body(p_form2_txt, [
        "The future forecasted price level ", ('m', 'P_hat_thc'), " is subsequently reconstructed exactly via the deterministic inverse mapping "
        "in (2). This direct cumulative return formulation transforms the non-stationary price series into stationary log-returns ",
        ('m', 'p_val'), " (empirically validated in Section 4.2), while completely eliminating autoregressive error compounding across multi-step horizons."
    ], is_en=True)
    
    set_para_heading(p_h2_arch, "3.2. Overall System Architecture", style_name='H2x', outline_lvl=1)
    set_para_body(p_arch_txt, [
        "Fig. 1 illustrates the neural architecture of GUMNetHet. Unlike conventional MoE networks that feed all features into every expert—which "
        "often degrades specialization—GUMNetHet partitions features across frequency domains and economic characteristics: high-frequency "
        "price dynamics to multi-scale 1D-CNN, slowly-evolving macroeconomic regimes to GRU-Attention, and fat-tailed crack-spread/shock indicators "
        "to Wavelet-KAN. This mechanism enforces strong domain inductive bias, effectively preventing expert degeneration."
    ], is_en=True)
    
    set_para_caption(p_cap_fig1, "Fig. 1.", "Architectural overview of GUMNetHet: feature partitioning, three heterogeneous experts, horizon-aware router, and multi-quantile head.", "1", is_table=False)
    
    set_para_heading(p_h2_cnn, "3.2.1. Price Momentum Expert: Multi-Scale 1D-CNN", style_name='H2x', outline_lvl=2)
    set_para_body(p_cnn_txt, [
        "The price expert employs three parallel 1D convolutional layers with kernel sizes ", ('m', 'k_set'),
        " to extract multi-resolution temporal features, integrated with layer normalization and temporal attention according to (3) and (4):"
    ], is_en=True)
    
    set_para_heading(p_h2_gru, "3.2.2. Macroeconomic Regime Expert: GRU-Attention", style_name='H2x', outline_lvl=2)
    set_para_body(p_gru_txt, [
        "The macroeconomic expert captures low-frequency trend signals via a 2-layer stacked GRU network with dropout = 0.1, "
        "extracting the final hidden state representation according to (5):"
    ], is_en=True)
    
    set_para_heading(p_h2_wkan, "3.2.3. Non-Linear Shock-Absorption Expert: Wavelet-KAN", style_name='H2x', outline_lvl=2)
    set_para_body(p_wkan_txt, [
        "To capture severe non-linear structural breaks induced by extreme geopolitical risk events, the third expert implements a "
        "Kolmogorov–Arnold Network integrated with Mexican Hat wavelets ", ('m', 'mexican_hat'),
        " parameterized by learnable translation and dilation coefficients according to (6) and (7):"
    ], is_en=True)
    
    set_para_heading(p_h2_router, "3.2.4. Horizon-Aware Dynamic Gating Router", style_name='H2x', outline_lvl=2)
    set_para_body(p_router_txt, [
        "To dynamically modulate expert contributions across diverse forecasting horizons and varying market regimes, the gating router "
        "accepts concatenated expert representations, horizon positional embeddings, and global summary context statistics ",
        ('m', 'x_ctx'), " according to (8), (9), and (10):"
    ], is_en=True)
    
    set_para_heading(p_h2_res, "3.2.5. Residual Scaling Bounding & Multi-Quantile Prediction Head", style_name='H2x', outline_lvl=2)
    set_para_body(p_res_txt, [
        "To prevent variance explosion at long horizons (such as H60), GUMNetHet incorporates a learnable residual scaling vector ",
        ('m', 'gamma_h'), " initialized at 0.1, outputting multi-quantile return predictions for ", ('m', 'q_set'), " according to (11):"
    ], is_en=True)
    
    set_para_heading(p_h2_loss, "3.2.6. Dual-Loss Optimization", style_name='H2x', outline_lvl=2)
    set_para_body(p_loss_txt, [
        "The overall objective function combines the multi-quantile pinball loss with a load-balancing regularization penalty to mitigate "
        "routing collapse according to (12), (13), and (14):"
    ], is_en=True)

    # --- 6. SECTION 4: EXPERIMENTAL SETUP ---
    set_para_heading(p_h1_exp, "4. EXPERIMENTAL SETUP", style_name='H1x', outline_lvl=0)
    set_para_heading(p_h2_data, "4.1. Data and Walk-Forward Protocol", style_name='H2x', outline_lvl=1)
    
    set_para_body(p_data_txt, [
        "The dataset covers 03/11/2008–30/04/2026 with ", ('m', 'N_val'), " trading-day observations. The two reported primary targets are MG95 "
        "and DO 0.001% per Platts Singapore; exogenous covariates include inter-product Platts benchmarks, WTI, Brent, GPR ", ('ref', '1'),
        ", DXY, crude oil production, crack-spread ratios, realized volatility, and calendar features."
    ], is_en=True)
    
    set_para_body(p_data_wf_txt, [
        "To prevent look-ahead bias, daily variables are indexed at ", ('m', 't_minus_1'), "; GPR is lagged by 30 calendar days; crude oil production "
        "is lagged by 7 days; and rolling features are computed strictly after applying lags. All scalers are fit exclusively on the training partition "
        "at each walk-forward step. Lookback length is ", ('m', 'L_30'), "; train/validation ratio is 85/15 per expansion. Expanding walk-forward test "
        "windows increase progressively with horizon: 100 trading days for H1–H5 (11/12/2025–30/04/2026), 150 days for H7 (02/10/2025–30/04/2026), "
        "200 days for H10 (24/07/2025–30/04/2026), 300 days for H20 (10/03/2025–30/04/2026), and 600 days for H60 (10/01/2024–30/04/2026). "
        "All test windows fall entirely within a period of clustered extreme energy shocks (Red Sea crisis, Middle East escalations, "
        "Russia–Ukraine war, and OPEC+ quota shifts). Specifically, annualized realized return volatility during the short-term test window "
        "(100 days) surges to 73.04% for MG95 (1.90× historical training volatility of 38.45%) and 96.29% for DO 0.001% (2.90× training volatility "
        "of 33.16%). The Geopolitical Risk (GPR) index averages 225.66 during the test period (nearly doubling the historical mean of 114.60) and "
        "peaks at 500.81 (90th percentile at 376.48). Price swings in the H60 test window record extreme peak-to-trough variations: 70.34 to 170.52 USD/bbl "
        "(+142.4%) for MG95 and 75.90 to 292.82 USD/bbl (+285.8%) for DO 0.001%. Comprehensive train/test splits and volatility characteristics "
        "across horizons are summarized in Table 1. This configuration serves as a rigorous stress-testing benchmark, ensuring models are evaluated "
        "on their capacity to adapt under fat-tailed regime shifts rather than static market conditions."
    ], is_en=True)

    set_para_heading(p_h2_diag, "4.2. Econometric Diagnostics and Training Configuration", style_name='H2x', outline_lvl=1)
    set_para_body(p_diag_txt, [
        "ADF and KPSS unit root tests in Table 2 provide evidence of non-stationarity in raw price levels; upon log-return transformation, "
        "ADF tests reject the unit-root null hypothesis (", ('m', 'p_val'), ") across all series. Return kurtosis is exceptionally high "
        "(213.35 for WTI and 17.51 for MG95), substantiating fat tails and extreme jumps in energy price series."
    ], is_en=True)
    set_para_caption(p_cap_tbl2, "Table 2.", "Condensed econometric statistical diagnostics for primary energy series.", "2", is_table=True)
    
    set_para_body(p_impl_txt, [
        "Implementation utilizes PyTorch 2.1 / Python 3.10 with AdamW optimizer (lr = 10⁻³, weight_decay = 10⁻⁴), "
        "CosineAnnealingLR scheduler (T_max = 60, η_min = 10⁻⁵), batch size = 64, early stopping patience = 15, and 60 training epochs."
    ], is_en=True)

    set_para_heading(p_h2_base, "4.3. Baseline Methods and Evaluation Metrics", style_name='H2x', outline_lvl=1)
    set_para_body(p_base_txt, [
        "Evaluated baselines comprise six representative state-of-the-art benchmarks: PatchTST [8], iTransformer [9], TimesNet [10], "
        "DLinear [11], TFT [12], and N-HiTS [14]. Evaluation metrics include MAE, RMSE, MAPE (%), ", ('m', 'R2'), ", and Directional Accuracy (DA, %). "
        "Probabilistic forecasts are quantified using Prediction Interval Coverage Probability (PICP) and Prediction Interval Normalized "
        "Average Width (PINAW) on ", ('m', 'q_interval'), "."
    ], is_en=True)

    # --- 7. SECTION 5: RESULTS AND DISCUSSION ---
    set_para_heading(p_h1_res, "5. RESULTS AND DISCUSSION", style_name='H1x', outline_lvl=0)
    set_para_heading(p_h2_point, "5.1. Multi-Horizon Point Forecasting Performance", style_name='H2x', outline_lvl=1)
    
    set_para_body(p_point_txt, [
        "Fig. 2 alongside Table 3 and Table 4 demonstrates that GUMNetHet consistently maintains lower MAE than all reported baselines "
        "across both products over all horizons, with pronounced advantages at H20–H60 where the test set undergoes massive price swings "
        "of up to +142.4% (MG95) and +285.8% (DO). At H60, gasoline MAE is 4.847 compared to 6.933 for the best baseline (a 30.1% reduction); "
        "for diesel, MAE is 7.066 versus 9.167 (a 22.9% reduction). However, ", ('m', 'R2'), " at H60 moderates to 0.155 (gasoline) and −0.007 (diesel), "
        "indicating that long-term performance represents superior price-level stabilization under high volatility rather than flawless trajectory forecasting."
    ], is_en=True)

    set_para_caption(p_cap_fig2, "Fig. 2.", "MAE and R² degradation curves across seven forecasting horizons for MG95 and DO 0.001% (Seed=42).", "2", is_table=False)
    set_para_caption(p_cap_tbl3, "Table 3.", "Detailed empirical forecasting performance on MG95 gasoline (Seed=42). MAE/RMSE in USD/bbl; DA in percent.", "3", is_table=True)
    set_para_caption(p_cap_tbl4, "Table 4.", "Detailed empirical forecasting performance on DO 0.001% diesel (Seed=42).", "4", is_table=True)

    set_para_heading(p_h2_da, "5.2. Directional Accuracy Analysis", style_name='H2x', outline_lvl=1)
    set_para_body(p_da_txt, [
        "Directional accuracy in Fig. 3 reveals distinct behavior between short and long horizons. For gasoline (Table 3), GUMNetHet attains "
        "high DA between 90.95% and 95.56% at H1–H7; for diesel (Table 4), DA spans 76.65% to 84.92%. At H20, DA remains substantial "
        "(91.65% for gasoline; 71.11% for diesel). Conversely, at H10 and H60, DA drops below 50% (42.24%/27.95% for gasoline and 32.29%/19.10% for diesel). "
        "This phenomenon aligns with financial econometrics principles: over extended horizons (H60 spanning nearly three months), oil prices "
        "approximate a near random walk with exponentially compounding uncertainty, rendering directional signals noisy. Rather than taking "
        "extreme directional bets that risk catastrophic error explosion, GUMNetHet's residual scaling mechanism shrinks predictions toward "
        "robust central bounds to minimize absolute error (MAE reduced by 30.1% at H60). Consequently, GUMNetHet operates as an effective "
        "directional trading signal at short horizons (H1–H7), while transitioning into a price-level risk and uncertainty quantification "
        "tool at long horizons (H10–H60)."
    ], is_en=True)

    set_para_caption(p_cap_fig3, "Fig. 3.", "Directional accuracy (DA%) of GUMNetHet and baselines across H1–H60 for gasoline and diesel.", "3", is_table=False)

    set_para_heading(p_h2_prob_res, "5.3. Probabilistic Forecasting, Ablation, and Router Behavior", style_name='H2x', outline_lvl=1)
    set_para_body(p_prob_res_txt, [
        "The ", ('m', 'q_interval'), " prediction interval achieves an empirical coverage probability of PICP=82.4% "
        "(exceeding nominal 80%) with a normalized average width of PINAW=0.142. This validates that the model avoids under-coverage "
        "while maintaining sharpness, offering practical utility for oil importers (e.g., Petrolimex, PVOIL) in forward pricing, inventory buffer "
        "sizing, and hedging optimization. Fig. 4 illustrates that prediction intervals dynamically widen during heightened market volatility. "
        "Ablation results in Table 5 indicate that replacing Wav-KAN with MLP causes the steepest degradation among expert variants; a uniform "
        "router similarly inflates MAE, corroborating the essential roles of expert specialization and adaptive routing."
    ], is_en=True)

    set_para_caption(p_cap_tbl5, "Table 5.", "Condensed ablation study of GUMNetHet variants (Seed=42).", "5", is_table=True)
    
    set_para_body(p_res_scale_txt, [
        "Empirical results reveal that removing residual scaling increases MAE by approximately 8.5%/6.3% at H20 and 14.1%/11.8% at H60 "
        "for gasoline/diesel. Router analysis in Fig. 4 indicates that under low GPR and short horizons, the CNN expert receives an average "
        "weight of ~0.48; when GPR exceeds the 90th percentile, the Wav-KAN weight surges from ~0.29 to 0.61 at medium–long horizons while "
        "the CNN weight contracts to ~0.21. These findings demonstrate that the router executes true regime-switching rather than trivial output averaging."
    ], is_en=True)
    
    set_para_caption(p_cap_fig4, "Fig. 4.", "Top: Multi-quantile fan chart under high volatility. Bottom: Router gating weights across low-GPR and high-GPR regimes.", "4", is_table=False)

    # --- 8. SECTION 6: CONCLUSION & REFERENCES ---
    set_para_heading(p_h1_concl, "6. CONCLUSION", style_name='H1x', outline_lvl=0)
    
    set_para_body(p_concl1, [
        "This paper proposed GUMNetHet, a heterogeneous mixture-of-experts architecture for multi-horizon probabilistic refined oil price "
        "forecasting. The model synergizes multi-scale 1D-CNN, GRU-Attention, and Wavelet-KAN through a horizon-aware router, combined with "
        "multi-quantile outputs and residual scaling to bolster stability over extended horizons."
    ], is_en=True)
    
    set_para_body(p_concl2, [
        "Expanding walk-forward experiments on MG95 and DO 0.001% demonstrate that GUMNetHet achieves the lowest MAE across all evaluated "
        "baselines throughout H1–H60, yielding MAE reductions at H60 of 30.1% for MG95 and 22.9% for DO 0.001%. Ablation experiments and "
        "router weight analyses corroborate the contributions of feature partitioning, Wavelet-KAN, and adaptive routing. However, the "
        "attenuation of ", ('m', 'R2'), " and directional accuracy at long horizons underscores that model superiority lies in price-level error containment "
        "and uncertainty quantification, rather than long-range directional forecasting."
    ], is_en=True)
    
    set_para_body(p_concl3, [
        "Future work will focus on multi-seed benchmarking, expanding probabilistic baseline comparisons, formal statistical significance "
        "testing for multi-step forecasts, and integrating real-time multimodal geopolitical news representations."
    ], is_en=True)
    
    set_para_heading(p_h1_refs, "REFERENCES", style_name='H1x', outline_lvl=0)

    # --- 9. INSERT TABLE 1 RIGHT AFTER P_DATA_WF_TXT ---
    t1_data = [
        ["Horizon", "Test Window", "Test Date Range", "Ann. Vol MG95 (Train→Test)", "Ann. Vol DO (Train→Test)", "Test GPR (Mean / Max)", "Test Price Range (USD/bbl)"],
        ["H1, H3, H5", "100 days", "11/12/2025 – 30/04/2026", "38.45% → 73.04% (1.90×)", "33.16% → 96.29% (2.90×)", "225.66 / 500.81", "MG95: [70.58, 170.52]\nDO: [77.11, 292.82]"],
        ["H7", "150 days", "02/10/2025 – 30/04/2026", "38.61% → 60.75% (1.57×)", "33.20% → 80.59% (2.43×)", "197.51 / 500.81", "MG95: [70.58, 170.52]\nDO: [77.11, 292.82]"],
        ["H10", "200 days", "24/07/2025 – 30/04/2026", "38.77% → 53.60% (1.38×)", "33.31% → 70.63% (2.12×)", "184.98 / 500.81", "MG95: [70.58, 170.52]\nDO: [77.11, 292.82]"],
        ["H20", "300 days", "10/03/2025 – 30/04/2026", "38.95% → 47.24% (1.21×)", "33.35% → 60.50% (1.81×)", "184.47 / 540.16", "MG95: [70.34, 170.52]\nDO: [75.90, 292.82]"],
        ["H60", "600 days", "10/01/2024 – 30/04/2026", "39.81% → 37.86% (0.95×)", "33.93% → 46.20% (1.36×)", "165.51 / 540.16", "MG95: [70.34, 170.52]\nDO: [75.90, 292.82]"]
    ]

    p38_elem = p_data_wf_txt._p
    
    new_tbl = doc.add_table(rows=len(t1_data), cols=len(t1_data[0]))
    for r_idx, row in enumerate(t1_data):
        for c_idx, val in enumerate(row):
            cell = new_tbl.cell(r_idx, c_idx)
            cell.text = val
            
    style_table(new_tbl)
    
    new_caption_p = doc.add_paragraph()
    set_para_caption(new_caption_p, "Table 1.", "Expanding walk-forward train/test split configurations and empirical volatility characteristics across forecasting horizons.", "1", is_table=True)
    
    p38_elem.addnext(new_caption_p._p)
    p38_elem.addnext(new_tbl._tbl)

    # --- 10. TRANSLATE & BOLD ALL TABLES ---
    
    # Table 2: Econometric Diagnostics (Table index 9)
    t2_headers = ["Series", "ADF Level p", "KPSS Level p", "ADF Return p", "Kurtosis"]
    for c_idx, h in enumerate(t2_headers):
        doc.tables[9].rows[0].cells[c_idx].text = h
        
    # Table 3: MG95 (Table index 10)
    t3_headers = ["H", "Model", "MAE", "RMSE", "MAPE (%)", "R²", "DA (%)"]
    for c_idx, h in enumerate(t3_headers):
        doc.tables[10].rows[0].cells[c_idx].text = h

    # Table 4: DO (Table index 11)
    t4_headers = ["H", "Model", "MAE", "RMSE", "MAPE (%)", "R²", "DA (%)"]
    for c_idx, h in enumerate(t4_headers):
        doc.tables[11].rows[0].cells[c_idx].text = h

    # Table 5: Ablation (Table index 12)
    t5_headers = ["Variant", "MG95-H3 MAE", "MG95-H3 R²", "DO-H5 MAE", "DO-H5 R²"]
    for c_idx, h in enumerate(t5_headers):
        doc.tables[12].rows[0].cells[c_idx].text = h
    
    ablation_variant_map = {
        "GUMNetHet đầy đủ": "Full GUMNetHet",
        "w/o Wavelet-KAN (thay bằng MLP)": "w/o Wavelet-KAN (MLP replacement)",
        "w/o 1D-CNN": "w/o 1D-CNN",
        "w/o GRU-Attention": "w/o GRU-Attention",
        "Định tuyến đồng nhất (Không Router)": "Uniform Routing (No Router)",
        "w/o Residual Scaling": "w/o Residual Scaling"
    }
    for row in doc.tables[12].rows[1:]:
        curr_txt = row.cells[0].text.strip()
        if curr_txt in ablation_variant_map:
            row.cells[0].text = ablation_variant_map[curr_txt]

    for tbl_idx, t in enumerate(doc.tables):
        style_table(t)
        # Bold header row ONLY
        for cell in t.rows[0].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
        
        # In Table 3, Table 4, Table 5, bold GUMNetHet row
        for row in t.rows[1:]:
            first_cell_text = row.cells[0].text.strip()
            if "GUMNetHet" in first_cell_text or "Full GUMNetHet" in first_cell_text:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.bold = True

    doc.save(target_path)
    print(f"Successfully generated clean styled English manuscript at {target_path}!")

if __name__ == '__main__':
    target = sys.argv[1] if len(sys.argv) > 1 else 'GUMNETHet_FAIRv3_EN.docx'
    build_english_doc(target)
