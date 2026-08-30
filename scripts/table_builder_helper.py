import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

# Helper to create inline OMML formula
def create_inline_omml(xml_body):
    return f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{xml_body}</m:oMath>'

# Math snippets
M = {
    'P_tc': '<m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t,c</m:t></m:r></m:sub></m:sSub>',
    'c_set': '<m:r><m:t>c ∈ {MG95, DO 0.001%}</m:t></m:r>',
    'L_30': '<m:r><m:t>L = 30</m:t></m:r>',
    'h_set': '<m:r><m:t>h ∈ {1, 3, 5, 7, 10, 20, 60}</m:t></m:r>',
    'r_thc': '<m:sSub><m:e><m:r><m:t>r</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t> = ln(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t>) − ln(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t,c</m:t></m:r></m:sub></m:sSub><m:r><m:t>)</m:t></m:r>',
    'P_hat_thc': '<m:sSub><m:e><m:r><m:t>P̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t,c</m:t></m:r></m:sub></m:sSub><m:r><m:t> · exp(</m:t></m:r><m:sSub><m:e><m:r><m:t>r̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t>)</m:t></m:r>',
    'X_price': '<m:sSub><m:e><m:r><m:t>X</m:t></m:r></m:e><m:sub><m:r><m:t>price</m:t></m:r></m:sub></m:sSub>',
    'X_macro': '<m:sSub><m:e><m:r><m:t>X</m:t></m:r></m:e><m:sub><m:r><m:t>macro</m:t></m:r></m:sub></m:sSub>',
    'X_shock': '<m:sSub><m:e><m:r><m:t>X</m:t></m:r></m:e><m:sub><m:r><m:t>shock</m:t></m:r></m:sub></m:sSub>',
    'k_set': '<m:r><m:t>k ∈ {3, 7, 15}</m:t></m:r>',
    'router_in': '<m:r><m:t>[</m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>CNN</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>GRU</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>KAN</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ Emb(h) ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>ctx</m:t></m:r></m:sub></m:sSub><m:r><m:t>]</m:t></m:r>',
    'w_h': '<m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> = Softmax(MLP(·))</m:t></m:r>',
    'f_fused': '<m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>fused</m:t></m:r></m:sub></m:sSub><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>∑</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>h,i</m:t></m:r></m:sub></m:sSub><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>',
    'x_ctx': '<m:sSub><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>ctx</m:t></m:r></m:sub></m:sSub>',
    'q_set': '<m:r><m:t>q ∈ {0.1, 0.5, 0.9}</m:t></m:r>',
    'gamma_h_set': '<m:sSub><m:e><m:r><m:t>γ</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> ∈ (0, 1)</m:t></m:r>',
    'r_hat_q': '<m:sSubSup><m:e><m:r><m:t>r̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub><m:sup><m:r><m:t>(q)</m:t></m:r></m:sup></m:sSubSup><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>Head</m:t></m:r></m:e><m:sub><m:r><m:t>q</m:t></m:r></m:sub></m:sSub><m:r><m:t>(</m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>fused</m:t></m:r></m:sub></m:sSub><m:r><m:t>) + </m:t></m:r><m:sSub><m:e><m:r><m:t>γ</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> · </m:t></m:r><m:sSubSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub><m:sup><m:r><m:t>target</m:t></m:r></m:sup></m:sSubSup>',
    'L_tot': '<m:r><m:t>ℒ = </m:t></m:r><m:sSub><m:e><m:r><m:t>ℒ</m:t></m:r></m:e><m:sub><m:r><m:t>pinball</m:t></m:r></m:sub></m:sSub><m:r><m:t> + α</m:t></m:r><m:sSub><m:e><m:r><m:t>ℒ</m:t></m:r></m:e><m:sub><m:r><m:t>balance</m:t></m:r></m:sub></m:sSub>',
    'alpha_val': '<m:r><m:t>α = 0.01</m:t></m:r>',
    'L_bal': '<m:sSub><m:e><m:r><m:t>ℒ</m:t></m:r></m:e><m:sub><m:r><m:t>balance</m:t></m:r></m:sub></m:sSub>',
    'gamma_h': '<m:sSub><m:e><m:r><m:t>γ</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub>',
    't_minus_1': '<m:r><m:t>t − 1</m:t></m:r>',
    'N_val': '<m:r><m:t>N = 4.512</m:t></m:r>',
    'R2': '<m:sSup><m:e><m:r><m:t>R</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>',
    'DA_formula': '<m:r><m:t>DA = </m:t></m:r><m:f><m:num><m:r><m:t>1</m:t></m:r></m:num><m:den><m:r><m:t>N</m:t></m:r></m:den></m:f><m:sSubSup><m:e><m:r><m:t>∑</m:t></m:r></m:e><m:sub><m:r><m:t>t=1</m:t></m:r></m:sub><m:sup><m:r><m:t>N</m:t></m:r></m:sup></m:sSubSup><m:r><m:t> 𝕀[sign(</m:t></m:r><m:sSub><m:e><m:r><m:t>P̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub></m:sSub><m:r><m:t> − </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub></m:sSub><m:r><m:t>) = sign(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub></m:sSub><m:r><m:t> − </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub></m:sSub><m:r><m:t>)]</m:t></m:r>',
    'q_interval': '<m:r><m:t>[</m:t></m:r><m:sSub><m:e><m:r><m:t>q</m:t></m:r></m:e><m:sub><m:r><m:t>0.1</m:t></m:r></m:sub></m:sSub><m:r><m:t>, </m:t></m:r><m:sSub><m:e><m:r><m:t>q</m:t></m:r></m:e><m:sub><m:r><m:t>0.9</m:t></m:r></m:sub></m:sSub><m:r><m:t>]</m:t></m:r>',
    'p_val': '<m:r><m:t>p &lt; 0.001</m:t></m:r>',
    'mexican_hat': '<m:r><m:t>ψ(z) = (1 − </m:t></m:r><m:sSup><m:e><m:r><m:t>z</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup><m:r><m:t>) · </m:t></m:r><m:sSup><m:e><m:r><m:t>e</m:t></m:r></m:e><m:sup><m:r><m:t>−0.5z²</m:t></m:r></m:sup></m:sSup>'
}

def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    for old_b in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old_b)
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                        f'<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                        f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
                        f'<w:insideV w:val="none"/>'
                        f'<w:left w:val="none"/>'
                        f'<w:right w:val="none"/>'
                        f'</w:tblBorders>')
    tblPr.append(borders)
    for row_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        if row_idx == 0:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for old_shd in tcPr.findall(qn('w:shd')):
                tcPr.remove(old_shd)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="FFFFFF"/>')
            tcPr.append(shd)
            mar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                            f'<w:top w:w="80" w:type="dxa"/>'
                            f'<w:bottom w:w="80" w:type="dxa"/>'
                            f'<w:left w:w="80" w:type="dxa"/>'
                            f'<w:right w:w="80" w:type="dxa"/>'
                            f'</w:tcMar>')
            tcPr.append(mar)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(8.0)
                    r.bold = False

def make_caption(p, text, item_id, is_table=True):
    p.style = 'Caption'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if is_table else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6 if is_table else 3)
    p.paragraph_format.space_after = Pt(3 if is_table else 6)
    p.text = ""
    prefix = "tbl_" if is_table else "fig_"
    bm_id = f"{300 if is_table else 400}{item_id}"
    bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="{prefix}{item_id}"/>')
    bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>')
    p._p.append(bm_start)
    p._p.append(bm_end)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8.5 if is_table else 8.0)
    run.bold = False

def add_styled_paragraph(p, segments, line_spacing=1.15, space_after=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p.text = ""
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    
    for seg in segments:
        if isinstance(seg, str):
            parts = re.split(r'(Bảng \d+|Hình \d+|\[\d+\])', seg)
            for part in parts:
                if not part:
                    continue
                m_tbl = re.match(r'Bảng (\d+)', part)
                m_fig = re.match(r'Hình (\d+)', part)
                m_ref = re.match(r'\[(\d+)\]', part)
                if m_tbl:
                    t_id = f"tbl_{m_tbl.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{t_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                elif m_fig:
                    f_id = f"fig_{m_fig.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{f_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                elif m_ref:
                    r_id = f"ref_{m_ref.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                else:
                    run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9.5)
        elif isinstance(seg, tuple) and seg[0] == 'm':
            m_xml = M[seg[1]]
            p._p.append(parse_xml(create_inline_omml(m_xml)))
        elif isinstance(seg, tuple) and seg[0] == 'ref':
            r_id = f"ref_{seg[1]}"
            hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                           f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                           f'<w:t>[{seg[1]}]</w:t></w:r></w:hyperlink>')
            p._p.append(hl)
        elif isinstance(seg, tuple) and seg[0] == 'refs':
            for idx, r_num in enumerate(seg[1]):
                r_id = f"ref_{r_num}"
                hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                               f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                               f'<w:t>[{r_num}]</w:t></w:r></w:hyperlink>')
                p._p.append(hl)
                if idx < len(seg[1]) - 1:
                    r_c = p.add_run(", ")
                    r_c.font.name = 'Times New Roman'
                    r_c.font.size = Pt(9.5)

print("Setup completed successfully.")
