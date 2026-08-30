# -*- coding: utf-8 -*-
import docx
import re
import os
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_math_run(p, text, is_sub=False, is_sup=False):
    run = p.add_run(text)
    if is_sub:
        run.font.subscript = True
    if is_sup:
        run.font.superscript = True
    if text.isalpha() and len(text) == 1:
        run.font.italic = True
    return run

path = 'docs/B\u1ea3n_th\u1ea3o_GUMNET_v2.docx'
doc = docx.Document(path)

# 1. Author Block
for p in doc.paragraphs:
    if 'Huong Bui^1' in p.text:
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Huong Bui")
        r_sup = p.add_run("1")
        r_sup.font.superscript = True
        p.add_run(", Phuoc Anh Dung Nguyen")
        r_sup2 = p.add_run("1")
        r_sup2.font.superscript = True
        p.add_run(", Van Quy Hoang")
        r_sup3 = p.add_run("2*")
        r_sup3.font.superscript = True
    elif '^1 Faculty of Information Technology' in p.text:
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sup = p.add_run("1")
        r_sup.font.superscript = True
        r_it = p.add_run("Faculty of Information Technology, HUTECH University, Ho Chi Minh City, Vietnam")
        r_it.font.italic = True
    elif '^2 Thuy Loi University' in p.text:
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sup = p.add_run("2")
        r_sup.font.superscript = True
        r_it = p.add_run("Thuy Loi University (TLU), Hanoi, Vietnam")
        r_it.font.italic = True

# 2. Email Block -> Footer
email_paras_idx = []
for i, p in enumerate(doc.paragraphs):
    if p.text == '______________________________':
        email_paras_idx.append(i)
    if '* Corresponding author.' in p.text:
        email_paras_idx.append(i)
    if 'E-mail addresses: bd.huong@hutech' in p.text:
        email_paras_idx.append(i)

for idx in sorted(email_paras_idx, reverse=True):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)

sec = doc.sections[0]
sec.different_first_page_header_footer = True
footer = sec.first_page_footer
if len(footer.paragraphs) == 0:
    footer.add_paragraph()
footer.paragraphs[0].text = '______________________________'
p2 = footer.add_paragraph('* Corresponding author.')
p3 = footer.add_paragraph('E-mail addresses: bd.huong@hutech.edu.vn (H. Bui), anhdungnguyen955@gmail.com (P.A.D. Nguyen), hoangvanquy@tlu.edu.vn (V.Q. Hoang)')

# 3. Format Math
for p in doc.paragraphs:
    # 3.1 Inline equations P_{t+h} = P_t
    if 'P_{t+h} = P_t' in p.text:
        parts = p.text.split('P_{t+h} = P_t')
        p.text = ""
        p.add_run(parts[0])
        add_math_run(p, 'P', is_sub=False)
        add_math_run(p, 't+h', is_sub=True)
        add_math_run(p, ' = ')
        add_math_run(p, 'P', is_sub=False)
        add_math_run(p, 't', is_sub=True)
        if len(parts) > 1:
            p.add_run(parts[1])
            
    # 3.2 Block equation f_final
    if 'f_{final} = w_1 f_{cnn} + w_2 f_{gru} + w_3 f_{kan}' in p.text:
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_math_run(p, 'f')
        add_math_run(p, 'final', is_sub=True)
        add_math_run(p, ' = ')
        add_math_run(p, 'w')
        add_math_run(p, '1', is_sub=True)
        add_math_run(p, ' f')
        add_math_run(p, 'cnn', is_sub=True)
        add_math_run(p, ' + ')
        add_math_run(p, 'w')
        add_math_run(p, '2', is_sub=True)
        add_math_run(p, ' f')
        add_math_run(p, 'gru', is_sub=True)
        add_math_run(p, ' + ')
        add_math_run(p, 'w')
        add_math_run(p, '3', is_sub=True)
        add_math_run(p, ' f')
        add_math_run(p, 'kan', is_sub=True)

    # 3.3 Target variable R_t -> t+h
    if '(R\u0302_t \u2192 t+h)' in p.text:
        parts = p.text.split('(R\u0302_t \u2192 t+h)')
        p.text = ""
        p.add_run(parts[0])
        p.add_run('(')
        add_math_run(p, 'R\u0302')
        add_math_run(p, 't \u2192 t+h', is_sub=True)
        p.add_run(')')
        if len(parts) > 1:
            p.add_run(parts[1])

    # 3.4 f_cnn, f_gru, f_kan
    if 'f_{cnn}' in p.text or 'f_{gru}' in p.text or 'f_{kan}' in p.text:
        # replace them manually by splitting one by one
        def replace_token(para, token, math_base, math_sub):
            if token in para.text:
                parts = para.text.split(token)
                para.text = ""
                for i, part in enumerate(parts):
                    para.add_run(part)
                    if i < len(parts) - 1:
                        add_math_run(para, math_base)
                        add_math_run(para, math_sub, is_sub=True)
        replace_token(p, 'f_{cnn}', 'f', 'cnn')
        replace_token(p, 'f_{gru}', 'f', 'gru')
        replace_token(p, 'f_{kan}', 'f', 'kan')
        replace_token(p, 'Pos_h', 'Pos', 'h')

doc.save('docs/B\u1ea3n_th\u1ea3o_GUMNET_v2_perfect.docx')
print('Done!')
