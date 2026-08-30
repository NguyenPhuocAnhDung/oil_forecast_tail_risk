# -*- coding: utf-8 -*-
import docx
import re
import os
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_math_run(p, text, is_sub=False, is_sup=False):
    run = p.add_run(text)
    if is_sub:
        run.font.subscript = True
    if is_sup:
        run.font.superscript = True
    # If it's a single letter or specific math variables, italicize
    if (text.isalpha() and len(text) == 1) or text in ['seq_len', 'Pos', 'cnn', 'gru', 'kan', 'final']:
        if text not in ['Pos', 'cnn', 'gru', 'kan', 'final']: # only italicize single letters and seq_len
            run.font.italic = True
        if text == 'seq_len': run.font.italic = True
    return run

def parse_math_in_paragraph(p):
    text = p.text
    if '$' not in text:
        return
    
    parts = text.split('$')
    p.text = ""
    
    for i, part in enumerate(parts):
        if i % 2 == 0:
            p.add_run(part)
        else:
            math_text = part.replace('\\in', ' ∈ ').replace('\\{', '{').replace('\\}', '}').replace('\\text{Pos}', 'Pos').replace('\\_', '_')
            
            # Simple state machine to parse _ and ^
            j = 0
            while j < len(math_text):
                if math_text[j] == '_':
                    j += 1
                    if j < len(math_text) and math_text[j] == '{':
                        end = math_text.find('}', j)
                        if end != -1:
                            add_math_run(p, math_text[j+1:end], is_sub=True)
                            j = end + 1
                        else:
                            add_math_run(p, '{', is_sub=True)
                            j += 1
                    elif j < len(math_text):
                        add_math_run(p, math_text[j], is_sub=True)
                        j += 1
                elif math_text[j] == '^':
                    j += 1
                    if j < len(math_text) and math_text[j] == '{':
                        end = math_text.find('}', j)
                        if end != -1:
                            add_math_run(p, math_text[j+1:end], is_sup=True)
                            j = end + 1
                        else:
                            add_math_run(p, '{', is_sup=True)
                            j += 1
                    elif j < len(math_text):
                        add_math_run(p, math_text[j], is_sup=True)
                        j += 1
                else:
                    next_sub = math_text.find('_', j)
                    next_sup = math_text.find('^', j)
                    next_idx = min([idx for idx in [next_sub, next_sup] if idx != -1] + [len(math_text)])
                    chunk = math_text[j:next_idx]
                    for char in chunk:
                        add_math_run(p, char)
                    j = next_idx

def fix_document():
    doc = docx.Document('docs/B\u1ea3n_th\u1ea3o_GUMNET_v2.docx')
    
    # 1. Remove " (Mới)", "BỔ SUNG", "SỬA"
    for p in doc.paragraphs:
        if '(Mới)' in p.text:
            p.text = p.text.replace(' (Mới)', '')
        if '[BỔ SUNG]' in p.text:
            p.text = p.text.replace('[BỔ SUNG] ', '').replace('[BỔ SUNG]', '')
        if '[SỬA]' in p.text:
            p.text = p.text.replace('[SỬA] ', '').replace('[SỬA]', '')

    # 2. Replace the text-based ADF results with an ACTUAL table
    # Find where I inserted the text lines for ADF
    # The lines were:
    # Bảng 1. Kết quả kiểm định ADF và KPSS...
    # - Xăng RON95...
    # - Xăng RON92/E5...
    # - Diesel DO 0.05%S...
    # - Diesel DO 0.001%S-V...
    
    adf_paras = []
    for i, p in enumerate(doc.paragraphs):
        if 'Bảng 1. Kết quả kiểm định ADF và KPSS' in p.text:
            adf_paras.append(i)
        elif 'Xăng RON95: ADF=-2.9376' in p.text:
            adf_paras.append(i)
        elif 'Xăng RON92/E5: ADF=-2.8569' in p.text:
            adf_paras.append(i)
        elif 'Diesel DO 0.05%S: ADF=-2.3898' in p.text:
            adf_paras.append(i)
        elif 'Diesel DO 0.001%S-V: ADF=-2.3772' in p.text:
            adf_paras.append(i)

    if adf_paras:
        # We'll use the title paragraph to insert the table after it,
        # then we delete the text lines.
        title_idx = None
        for idx in adf_paras:
            if 'Bảng 1. Kết quả kiểm định' in doc.paragraphs[idx].text:
                title_idx = idx
                break
        
        if title_idx is not None:
            p_title = doc.paragraphs[title_idx]
            table = doc.add_table(rows=5, cols=5)
            table.style = 'Table Grid'
            headers = ['Chuỗi', 'Thống kê ADF', 'p-value (ADF)', 'Thống kê KPSS', 'Kết luận (α=5%)']
            for c in range(5): table.cell(0, c).text = headers[c]
            data = [
                ['Xăng RON95', '-2.9376', '0.0411', '1.1240', 'Dừng'],
                ['Xăng RON92/E5', '-2.8569', '0.0506', '1.1859', 'Dừng'],
                ['Diesel DO 0.05%S', '-2.3898', '0.1446', '0.9930', 'Không dừng'],
                ['Diesel DO 0.001%S-V', '-2.3772', '0.1483', '0.9574', 'Không dừng']
            ]
            for r in range(4):
                for c in range(5):
                    table.cell(r+1, c).text = data[r][c]
            
            p_title._element.addnext(table._element)
            
            # Delete the text lines (reverse order to not mess up indices)
            for idx in sorted(adf_paras, reverse=True):
                if idx != title_idx:
                    p = doc.paragraphs[idx]
                    p._element.getparent().remove(p._element)

    # 3. Replace Descriptive Stats text with an ACTUAL table
    # I inserted lines like:
    # "Bảng Thống kê Mô tả Dữ liệu..."
    # "- Xăng RON95: Mean=..."
    desc_paras = []
    for i, p in enumerate(doc.paragraphs):
        if 'Bảng Thống kê Mô tả' in p.text:
            desc_paras.append(i)
        elif 'Xăng RON95: Mean=87.31' in p.text:
            desc_paras.append(i)
        elif 'Xăng RON92/E5: Mean=83.45' in p.text:
            desc_paras.append(i)
        elif 'Diesel DO 0.05%S: Mean=85.20' in p.text:
            desc_paras.append(i)
        elif 'Diesel DO 0.001%S-V: Mean=86.10' in p.text:
            desc_paras.append(i)
            
    if desc_paras:
        title_idx = None
        for idx in desc_paras:
            if 'Bảng Thống kê Mô tả' in doc.paragraphs[idx].text:
                title_idx = idx
                break
                
        if title_idx is not None:
            p_title = doc.paragraphs[title_idx]
            table = doc.add_table(rows=5, cols=6)
            table.style = 'Table Grid'
            headers = ['Mặt hàng', 'Mean', 'Std', 'Min', 'Max', 'Số lần điều chỉnh']
            for c in range(6): table.cell(0, c).text = headers[c]
            data = [
                ['Xăng RON95', '87.31', '18.52', '35.40', '150.21', '405'],
                ['Xăng RON92/E5', '83.45', '19.10', '32.10', '145.60', '405'],
                ['Diesel DO 0.05%S', '85.20', '21.05', '30.50', '160.40', '405'],
                ['Diesel DO 0.001%S-V', '86.10', '21.50', '31.20', '165.80', '405']
            ]
            for r in range(4):
                for c in range(6):
                    table.cell(r+1, c).text = data[r][c]
                    
            p_title._element.addnext(table._element)
            
            for idx in sorted(desc_paras, reverse=True):
                if idx != title_idx:
                    p = doc.paragraphs[idx]
                    p._element.getparent().remove(p._element)

    # 4. Format math in all paragraphs
    for p in doc.paragraphs:
        parse_math_in_paragraph(p)
        
    doc.save('docs/B\u1ea3n_th\u1ea3o_GUMNET_v2_perfect2.docx')
    print('Done!')

if __name__ == '__main__':
    fix_document()
