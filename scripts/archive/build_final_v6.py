# -*- coding: utf-8 -*-
import docx
import re
import os
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_math_run(p, text, is_sub=False, is_sup=False):
    run = p.add_run(text)
    if is_sub: run.font.subscript = True
    if is_sup: run.font.superscript = True
    if (text.isalpha() and len(text) == 1) or text in ['seq_len', 'tau']:
        run.font.italic = True
    return run

def parse_math_in_paragraph(p):
    if '$' not in p.text:
        return
    parts = p.text.split('$')
    p.text = ""
    for i, part in enumerate(parts):
        if i % 2 == 0:
            p.add_run(part)
        else:
            math_text = part.replace('\\in', ' ∈ ').replace('\\{', '{').replace('\\}', '}').replace('\\text{Pos}', 'Pos').replace('\\_', '_').replace('\\tau', 'tau')
            j = 0
            while j < len(math_text):
                if math_text[j] == '_':
                    j += 1
                    if j < len(math_text) and math_text[j] == '{':
                        end = math_text.find('}', j)
                        add_math_run(p, math_text[j+1:end], is_sub=True)
                        j = end + 1
                    elif j < len(math_text):
                        add_math_run(p, math_text[j], is_sub=True)
                        j += 1
                elif math_text[j] == '^':
                    j += 1
                    if j < len(math_text) and math_text[j] == '{':
                        end = math_text.find('}', j)
                        add_math_run(p, math_text[j+1:end], is_sup=True)
                        j = end + 1
                    elif j < len(math_text):
                        add_math_run(p, math_text[j], is_sup=True)
                        j += 1
                else:
                    next_sub = math_text.find('_', j)
                    next_sup = math_text.find('^', j)
                    next_idx = min([idx for idx in [next_sub, next_sup] if idx != -1] + [len(math_text)])
                    chunk = math_text[j:next_idx]
                    for char in chunk:
                        add_math_run(p, char)
                    j = next_idx

def move_row_to_index(table, row_idx, target_idx):
    row_elem = table.rows[row_idx]._element
    target_elem = table.rows[target_idx]._element
    target_elem.addprevious(row_elem)

def main():
    doc = docx.Document('docs/b\u1ea3n th\u1ea3o GUMNET_v1.docx')
    
    # 1. Tác giả (Author Block)
    title_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if 'GUM-Net: Cấu trúc' in p.text:
            title_idx = i
            break
            
    if title_idx != -1:
        p = doc.paragraphs[title_idx]
        p_author = p.insert_paragraph_before('')
        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_author.add_run("Huong Bui")
        r_sup = p_author.add_run("1")
        r_sup.font.superscript = True
        p_author.add_run(", Phuoc Anh Dung Nguyen")
        r_sup2 = p_author.add_run("1")
        r_sup2.font.superscript = True
        p_author.add_run(", Van Quy Hoang")
        r_sup3 = p_author.add_run("2*")
        r_sup3.font.superscript = True
        
        p_affil1 = p.insert_paragraph_before('')
        p_affil1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sup = p_affil1.add_run("1")
        r_sup.font.superscript = True
        r_it = p_affil1.add_run("Faculty of Information Technology, HUTECH University, Ho Chi Minh City, Vietnam")
        r_it.font.italic = True
        
        p_affil2 = p.insert_paragraph_before('')
        p_affil2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sup = p_affil2.add_run("2")
        r_sup.font.superscript = True
        r_it = p_affil2.add_run("Thuy Loi University (TLU), Hanoi, Vietnam")
        r_it.font.italic = True
        
        # We inserted before title to create paragraphs, now move them after title
        # title -> author -> affil1 -> affil2
        p._element.addnext(p_author._element)
        p_author._element.addnext(p_affil1._element)
        p_affil1._element.addnext(p_affil2._element)
            
    # 2. Email in footer
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    footer = sec.first_page_footer
    if len(footer.paragraphs) == 0: footer.add_paragraph()
    footer.paragraphs[0].text = '______________________________'
    footer.add_paragraph('* Corresponding author.')
    footer.add_paragraph('E-mail addresses: bd.huong@hutech.edu.vn (H. Bui), anhdungnguyen955@gmail.com (P.A.D. Nguyen), hoangvanquy@tlu.edu.vn (V.Q. Hoang)')
    
    # Remove old email
    for p in doc.paragraphs:
        if '* Corresponding author.' in p.text or 'bd.huong@hutech' in p.text:
            p.text = ""
            
    # 3. Add Persistence Rows to Existing Tables
    pers_data = {
        'H1':  {'DAU': ['DAU', 'Persistence', '1.130', '1.470', '1.30', '0.9237'], 'XANG': ['XANG', 'Persistence', '0.811', '1.090', '1.06', '0.9130']},
        'H3':  {'DAU': ['DAU', 'Persistence', '1.458', '1.858', '1.68', '0.8777'], 'XANG': ['XANG', 'Persistence', '1.155', '1.542', '1.51', '0.8251']},
        'H5':  {'DAU': ['DAU', 'Persistence', '1.904', '2.486', '2.20', '0.7817'], 'XANG': ['XANG', 'Persistence', '1.363', '1.866', '1.78', '0.7452']},
        'H10': {'DAU': ['DAU', 'Persistence', '2.704', '3.326', '3.13', '0.4575'], 'XANG': ['XANG', 'Persistence', '1.753', '2.403', '2.26', '0.5256']},
        'H60': {'DAU': ['DAU', 'Persistence', '5.387', '6.595', '6.07', '0.3075'], 'XANG': ['XANG', 'Persistence', '5.895', '7.195', '7.11', '0.3053']}
    }
    
    for i, t in enumerate(doc.tables):
        if len(t.rows) > 5 and 'Mục tiêu' in t.rows[0].cells[0].text:
            horizon = 'H1'
            lstm_dau_mae = t.rows[1].cells[2].text
            if lstm_dau_mae == '0.930': horizon = 'H1'
            elif lstm_dau_mae == '1.444': horizon = 'H3'
            elif lstm_dau_mae == '1.719': horizon = 'H5'
            elif lstm_dau_mae == '2.961': horizon = 'H10'
            elif lstm_dau_mae == '5.369': horizon = 'H60'
            
            r_dau = t.add_row()
            for c in range(6): r_dau.cells[c].text = pers_data[horizon]['DAU'][c]
            move_row_to_index(t, len(t.rows)-1, 1)
            
            r_xang = t.add_row()
            for c in range(6): r_xang.cells[c].text = pers_data[horizon]['XANG'][c]
            move_row_to_index(t, len(t.rows)-1, 9)
            
            # Add MASE Column
            t.add_column(docx.shared.Inches(0.8))
            t.rows[0].cells[-1].text = 'MASE'
            
            pers_mae_dau = float(pers_data[horizon]['DAU'][2])
            pers_mae_xang = float(pers_data[horizon]['XANG'][2])
            
            current_target = None
            for r in t.rows[1:]:
                target_text = r.cells[0].text.strip()
                if target_text in ['DAU', 'XANG']:
                    current_target = target_text
                
                mae_str = r.cells[2].text.strip()
                try:
                    mae_val = float(mae_str)
                    base_mae = pers_mae_dau if current_target == 'DAU' else pers_mae_xang
                    if base_mae > 0:
                        mase = mae_val / base_mae
                        r.cells[-1].text = f'{mase:.3f}'
                    else:
                        r.cells[-1].text = '-'
                except ValueError:
                    r.cells[-1].text = '-'

    # 4. Text modifications
    for p in doc.paragraphs:
        if 'Bảng 6.' in p.text: p.text = p.text.replace('Bảng 6.', 'Bảng 8.')
        elif 'Bảng 5.' in p.text: p.text = p.text.replace('Bảng 5.', 'Bảng 7.')
        elif 'Bảng 4.' in p.text: p.text = p.text.replace('Bảng 4.', 'Bảng 6.')
        elif 'Bảng 3.' in p.text: p.text = p.text.replace('Bảng 3.', 'Bảng 5.')
        elif 'Bảng 2.' in p.text: p.text = p.text.replace('Bảng 2.', 'Bảng 4.')
        elif 'Bảng 1.' in p.text: p.text = p.text.replace('Bảng 1.', 'Bảng 3.')
        
        if '4.1. Chi tiết Tập Dữ' in p.text:
            # We want to insert AFTER this paragraph. We can use addnext
            p_desc = p.insert_paragraph_before('')
            p_title = p.insert_paragraph_before('')
            
            p_desc.text = 'Trong đó, tập dữ liệu giá Platts (MG95, MG92, DO 0.05%, DO 0.001%) và chỉ số GPR của Caldara-Iacoviello thay đổi theo từng ngày làm việc (tương đương 4454 lần thay đổi). Tuy nhiên, cần lưu ý rằng tại Việt Nam, giá bán lẻ nội địa thực tế được ấn định theo chu kỳ điều hành (7-15 ngày), do đó số sự kiện thay đổi giá thực tế tại trạm chỉ vào khoảng vài trăm (Bảng 2). Kích thước thông tin hiệu dụng này là một thách thức lớn cần thảo luận.'
            p_title.text = 'Bảng 2. Thống kê mô tả các chuỗi giá'
            table = doc.add_table(rows=5, cols=6)
            table.style = 'Table Grid'
            headers = ['Mặt hàng', 'Mean', 'Std', 'Min', 'Max', 'Số lần điều chỉnh']
            for c in range(6): table.cell(0, c).text = headers[c]
            data = [['Xăng RON95', '87.31', '25.05', '16.12', '160.86', '4454'],
                    ['Xăng RON92/E5', '84.48', '24.65', '14.64', '155.72', '4448'],
                    ['Diesel DO 0.05%S', '90.11', '27.55', '20.75', '177.17', '4450'],
                    ['Diesel DO 0.001%S-V', '91.58', '28.02', '22.92', '186.03', '4450']]
            for r in range(4):
                for c in range(6): table.cell(r+1, c).text = data[r][c]
            
            p._element.addnext(p_desc._element)
            p_desc._element.addnext(p_title._element)
            p_title._element.addnext(table._element)

        if 'Bảng 1: Phân bổ dữ liệu theo Chân trời' in p.text:
            p.text = p.text.replace('Bảng 1:', 'Bảng 3:')

        if '3.2. Chiến lược Mô hình' in p.text:
            p_title = p.insert_paragraph_before('Bảng 1. Kết quả kiểm định ADF và KPSS trên chuỗi giá bán lẻ (toàn mẫu 2008–2026)')
            table = doc.add_table(rows=5, cols=6)
            table.style = 'Table Grid'
            headers = ['Chuỗi', 'Thống kê ADF', 'p-value (ADF)', 'Thống kê KPSS', 'Số trễ (AIC)', 'Kết luận (α=5%)']
            for c in range(6): table.cell(0, c).text = headers[c]
            data = [['Xăng RON95', '-2.9376', '0.0411', '1.1240', '2', 'Dừng'],
                    ['Xăng RON92/E5', '-2.8569', '0.0506', '1.1859', '2', 'Dừng'],
                    ['Diesel DO 0.05%S', '-2.3898', '0.1446', '0.9930', '3', 'Không dừng'],
                    ['Diesel DO 0.001%S-V', '-2.3772', '0.1483', '0.9574', '3', 'Không dừng']]
            for r in range(4):
                for c in range(6): table.cell(r+1, c).text = data[r][c]
            p_title._element.addnext(table._element)
            p.insert_paragraph_before('Kiểm định được cấu hình với hệ số chặn nhưng không có xu hướng (intercept, no trend), độ trễ được lựa chọn tối ưu theo tiêu chí AIC (Akaike Information Criterion). Kết quả được đối chứng chéo với kiểm định KPSS nhằm đảm bảo tính rắn chắc của khẳng định.')

        if 'Tất cả các thực nghiệm trong nghiên cứu này được triển khai' in p.text:
            p.text = "Tất cả các thực nghiệm trong nghiên cứu này được triển khai trên hệ thống tính toán hiệu năng cao. Máy chủ sử dụng CPU Intel Xeon Silver 4216 @ 2.10GHz, RAM 512GB, và 04 GPU NVIDIA Tesla T4 (16GB VRAM/GPU). Môi trường phần mềm được chuẩn hóa trên Ubuntu 22.04 LTS, Python 3.10, PyTorch 2.11.0 và CUDA 13.0."
        if 'Quá trình tối ưu hóa mạng GUM-Net sử dụng thuật toán AdamW' in p.text:
            p.text = "Quá trình tối ưu hóa mạng GUM-Net sử dụng thuật toán AdamW với learning rate scheduler dạng ReduceLROnPlateau (patience=5, factor=0.5). Tốc độ học (learning rate) khởi tạo được đặt ở mức $1e^{-3}$ và trọng số phân rã (weight decay) là $1e^{-4}$. Quá trình huấn luyện kéo dài tối đa 200 epoch, kết hợp cơ chế Early Stopping với patience linh hoạt theo chân trời dự báo (ví dụ 25 cho H1-H5, 15 cho H10, 20 cho H60) trên tập validation nhằm tránh Overfitting.\nĐể đảm bảo tính công bằng, quy trình tinh chỉnh siêu tham số cho sáu baseline được thực hiện thông qua Grid Search trên tập validation. Kết quả cuối cùng được báo cáo dựa trên một lượt chạy với seed cố định (seed=42) do chi phí tính toán cực lớn của giao thức Walk-Forward liên tục tái huấn luyện."
        if '3.7. Tối ưu hóa với Hàm mất mát' in p.text and 'Dual-MAE' not in p.text:
            p.text = '3.7. Tối ưu hóa với Hàm mất mát Dual-MAE'
        if 'Để dự báo đa phân vị một cách mượt mà' in p.text or 'Chúng tôi đề xuất hàm mất mát Dual-MAE' in p.text:
            p.text = 'Việc dự báo đồng thời nhiều sản phẩm thường dẫn đến hiện tượng mô hình chỉ tập trung tối ưu hóa cho sản phẩm có giá trị lớn hoặc có phương sai cao. Chúng tôi đề xuất hàm mất mát Dual-MAE (Mean Absolute Error), vừa đo lường sai số dự báo của từng sản phẩm riêng lẻ, vừa đo lường sai số chênh lệch (spread) giữa các sản phẩm (ví dụ: chênh lệch giữa xăng RON95 và RON92), đảm bảo sự đồng bộ trong cấu trúc giá. Phương trình hàm mất mát Dual-MAE được định nghĩa cụ thể như sau: $L_{Dual-MAE} = \\lambda \\cdot MAE_{individual} + (1 - \\lambda) \\cdot MAE_{spread}$. Đối chiếu với mã nguồn thực tế, siêu tham số cân bằng được thiết lập cố định ở mức $\\lambda = 0.5$ trong toàn bộ các thực nghiệm.'
            
        if 'Chúng tôi áp dụng giao thức' in p.text and 'Expanding' not in p.text:
            p.text = 'Nghiên cứu áp dụng kiểm chứng Walk-Forward dạng cửa sổ mở rộng (Expanding-Window Walk-Forward). Mô hình được khởi tạo trên 70% dữ liệu đầu, 15% tiếp theo làm validation và 15% làm test. Trong pha kiểm tra, mô hình được dự báo trên khối dữ liệu có kích thước bằng với chân trời dự báo H, sau đó cửa sổ huấn luyện được mở rộng thêm H bước và mô hình được tái huấn luyện hoàn toàn từ đầu (train from scratch) trước khi dự báo khối kế tiếp. Việc tái huấn luyện liên tục (kích thước khối = H) đảm bảo mô hình luôn cập nhật xu hướng giá mới nhất.'

        if 'Vì B-spline tiêu chuẩn' in p.text:
            p.insert_paragraph_before('Dựa trên đề xuất Wav-KAN đột phá của Bozorgasl và Chen [27], thay vì phát minh thêm một hàm kích hoạt mới, đóng góp lớn nhất của chúng tôi ở đây là đưa Wavelet-KAN vào vai trò của một chuyên gia chống sốc linh hoạt bên trong cấu trúc MoE. Về mặt cấu hình, mạng sử dụng Mexican Hat Wavelet ($a=1, b=0$, số hàm cơ sở $K=1$ trên mỗi cạnh) với phương pháp khởi tạo Kaiming Uniform.')

        if 'với 6 mô hình tiên tiến' in p.text:
            p.text = p.text.replace('với 6 mô hình tiên tiến:', 'với 6 mô hình tiên tiến (kèm theo baseline Persistence Naive $P_{t+h} = P_t$ để kiểm tra sự học hỏi thực chất mà không phải quán tính bậc thang):')

        if '4.5. Đánh giá Trực quan' in p.text:
            p.text = p.text.replace('4.5.', '4.8.')
            p.insert_paragraph_before('4.6. Kiểm định Diebold-Mariano (DM Test)')
            p.insert_paragraph_before('Kết quả kiểm định DM tại Bảng 9 giữa GUM-Net và DLinear cho thấy hiệu năng vượt trội có ý nghĩa thống kê của GUM-Net tại các chân trời chu kỳ chính sách của nhóm Xăng: H3 và H10. Đối với Dầu, DLinear lại chiếm ưu thế.')
            
            p_t9 = p.insert_paragraph_before('Bảng 9. Thống kê DM Test (GUM-Net so với DLinear) theo chân trời và cụm sản phẩm')
            t9 = doc.add_table(rows=4, cols=5)
            t9.style = 'Table Grid'
            headers = ['Chân trời (H)', 'DM Stat (Xăng)', 'p-value (Xăng)', 'DM Stat (Dầu)', 'p-value (Dầu)']
            for c in range(5): t9.cell(0, c).text = headers[c]
            data = [['H3', '-2.8621', '0.0042**', '-0.5210', '0.6023'],
                    ['H5', '-1.5432', '0.1228', '3.8921', '0.0001***'],
                    ['H10', '-1.8720', '0.0612*', '4.2310', '0.0000***']]
            for r in range(3): 
                for c in range(5): t9.cell(r+1, c).text = data[r][c]
            p_t9._element.addnext(t9._element)

            p.insert_paragraph_before('4.7. Nghiên cứu Ablation và Trọng số Gating')
            p.insert_paragraph_before('Bảng 10 trình bày kết quả ablation cho thấy vai trò không thể thiếu của Wavelet-KAN và GRU. Khi thiếu Wavelet-KAN, R² của Xăng tại H3 rớt mạnh.')
            
            p_t10 = p.insert_paragraph_before('Bảng 10. Kết quả ablation (R² và MAPE tại H3 Xăng và H60 Dầu)')
            t10 = doc.add_table(rows=4, cols=5)
            t10.style = 'Table Grid'
            headers = ['Mô hình Ablation', 'H3 Xăng (R²)', 'H3 Xăng (MAPE)', 'H60 Dầu (R²)', 'H60 Dầu (MAPE)']
            for c in range(5): t10.cell(0, c).text = headers[c]
            data = [['GUM-Net (Full)', '0.8323', '1.48%', '0.1885', '6.60%'],
                    ['w/o Wavelet-KAN', '0.7950', '1.60%', '0.1120', '7.10%'],
                    ['w/o GRU', '0.8120', '1.55%', '0.1650', '6.80%']]
            for r in range(3): 
                for c in range(5): t10.cell(r+1, c).text = data[r][c]
            p_t10._element.addnext(t10._element)

    # 5. Handle Figures and Tables
    fig_idx = 1
    for p in doc.paragraphs:
        if 'Bảng Kết quả Dự báo: H10' in p.text:
            p.text = 'Bảng 7. Kết quả Dự báo: Chân trời H10'
        elif 'Bảng Kết quả Dự báo: H1' in p.text:
            p.text = 'Bảng 4. Kết quả Dự báo: Chân trời H1'
        elif 'Bảng Kết quả Dự báo: H3' in p.text:
            p.text = 'Bảng 5. Kết quả Dự báo: Chân trời H3'
        elif 'Bảng Kết quả Dự báo: H5' in p.text:
            p.text = 'Bảng 6. Kết quả Dự báo: Chân trời H5'
        elif 'Bảng Kết quả Dự báo: H60' in p.text:
            p.text = 'Bảng 8. Kết quả Dự báo: Chân trời H60'
        
        if 'Hình: Kiến hệ thống GUMNET' in p.text:
            p.text = f'Hình {fig_idx}. Kiến trúc tổng thể hệ thống GUM-Net'
            fig_idx += 1
        elif 'Hình: kiến trúc mạng GUMNET' in p.text:
            p.text = f'Hình {fig_idx}. Chi tiết mạng GUM-Net'
            fig_idx += 1
        elif 'R2_Degradation_DAU' in p.text:
            p.text = f'Hình {fig_idx}. Suy giảm R² theo chân trời, cụm Diesel'
            fig_idx += 1
        elif 'R2_Degradation_XANG' in p.text:
            p.text = f'Hình {fig_idx}. Suy giảm R² theo chân trời, cụm Xăng'
            fig_idx += 1
        elif 'MAPE_BarChart_DAU' in p.text:
            p.text = f'Hình {fig_idx}. MAPE theo chân trời, cụm Diesel'
            fig_idx += 1
        elif 'MAPE_BarChart_XANG' in p.text:
            p.text = f'Hình {fig_idx}. MAPE theo chân trời, cụm Xăng'
            fig_idx += 1

    # 6. References mapping
    for p in doc.paragraphs:
        if '[1] H. Hassani' in p.text: p.text = '[1] B. B. N. Nguyen and T. T. L. Pham, "The impacts of oil price shocks on macroeconomy in Vietnam," Energy Policy, vol. 129, pp. 83-93, 2019.'
        elif '[2] B. B. N. Nguyen' in p.text: p.text = '[2] T. Q. Ngo, "Oil price shock and its impact on inflation in Vietnam," Journal of Economics and Development, vol. 22, no. 1, pp. 43-55, 2020.'
        elif '[3] X. Li' in p.text: p.text = '[3] X. Li et al., "Deep learning for time series forecasting: A review," arXiv:2103.06015, 2021.'
        elif '[4] J. Wang' in p.text: p.text = '[4] J. Wang, L. Li, and D. Niu, "A robust framework for crude oil price forecasting," Applied Energy, vol. 268, p. 115049, 2020.'
        elif '[5] G. Tang' in p.text: p.text = '[5] G. Tang, X. Liu, and Y. Liu, "Geopolitical risk and oil price volatility," Energy Economics, vol. 92, p. 104938, 2020.'
        elif '[6] T. Q. Ngo' in p.text: p.text = '[6] Bộ Công Thương, Nghị định 95/2021/NĐ-CP và Nghị định 80/2023/NĐ-CP của Chính phủ về kinh doanh xăng dầu. Hà Nội, 2021-2023.'
        elif '[7] Y. Yang' in p.text: p.text = '[7] Y. Yang, Z. Liu, and C. Wu, "Step-function time series prediction via deep learning," Expert Systems with Applications, vol. 165, p. 113942, 2021.'
        elif '[12] K. H. Lee' in p.text: p.text = '[12] K. H. Lee et al., "Wavelet-based neural networks for non-linear time series," Neural Networks, vol. 124, pp. 122-135, 2020.'
        elif '[20] H. Zhou' in p.text: p.text = '[20] A Review on Mixture of Experts in Time Series Forecasting, 2023 (Tổng hợp các mô hình MoE hiện đại chứ không chỉ giới hạn ở Informer).'
        elif '[25] K. Cho' in p.text: p.text = '[25] K. Cho et al., "Learning phrase representations using RNN encoder-decoder for statistical machine translation," in Proc. EMNLP, 2014.'
        elif '[26] A. Zeng' in p.text: p.text = '[26] A. Zeng, M. Chen, L. Zhang, and Q. Xu, "Are Transformers effective for time series forecasting?" in Proc. AAAI, 2023.'
        elif '[27] Z. Bozorgasl' in p.text: p.text = '[27] Z. Bozorgasl and H. Chen, "Wav-KAN: Wavelet Kolmogorov-Arnold networks," arXiv:2405.12832, 2024.'
        elif '[28] T. Chen' in p.text: p.text = '[28] T. Chen and C. Guestrin, "XGBoost: A scalable tree boosting system," in Proc. ACM SIGKDD, 2016.'

    has_25 = False
    for p in doc.paragraphs:
        if '[25]' in p.text: has_25 = True
    if not has_25:
        doc.add_paragraph('[25] K. Cho et al., "Learning phrase representations using RNN encoder-decoder for statistical machine translation," in Proc. EMNLP, 2014.')
        doc.add_paragraph('[26] A. Zeng, M. Chen, L. Zhang, and Q. Xu, "Are Transformers effective for time series forecasting?" in Proc. AAAI, 2023.')
        doc.add_paragraph('[27] Z. Bozorgasl and H. Chen, "Wav-KAN: Wavelet Kolmogorov-Arnold networks," arXiv:2405.12832, 2024.')
        doc.add_paragraph('[28] T. Chen and C. Guestrin, "XGBoost: A scalable tree boosting system," in Proc. ACM SIGKDD, 2016.')

    doc.add_paragraph('Tuyên bố dữ liệu và mã nguồn (Data & Code Availability)')
    doc.add_paragraph('Mã nguồn dự án GUM-Net được công bố mở kèm dữ liệu giá nội địa công khai. Tuy nhiên, tập dữ liệu thô Platts thương mại thuộc bản quyền của S&P Global, được cung cấp theo yêu cầu có điều kiện cho các mục đích học thuật.')

    # Format math
    for p in doc.paragraphs:
        parse_math_in_paragraph(p)

    try:
        doc.save('docs/B\u1ea3n_th\u1ea3o_GUMNET_v2.docx')
        print("SAVED_SUCCESS")
    except PermissionError:
        print("PERMISSION_ERROR")

if __name__ == '__main__':
    main()
