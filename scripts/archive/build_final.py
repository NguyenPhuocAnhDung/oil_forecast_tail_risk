# -*- coding: utf-8 -*-
import docx
import os
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = 'docs/b\u1ea3n th\u1ea3o GUMNET_v1.docx'
doc = docx.Document(path)

# --- 1. S\u1eeda Author Block ---
for p in doc.paragraphs:
    if 'T\u00e1c gi\u1ea3' in p.text and '\u0111\u01a1n v\u1ecb' in p.text:
        p.text = 'Huong Bui^1, Phuoc Anh Dung Nguyen^1, Van Quy Hoang^{2*}'
        p.insert_paragraph_before('^1 Faculty of Information Technology, HUTECH University, Ho Chi Minh City, Vietnam')
        p.insert_paragraph_before('^2 Thuy Loi University (TLU), Hanoi, Vietnam')
        break

doc.add_paragraph('______________________________')
doc.add_paragraph('* Corresponding author.')
doc.add_paragraph('E-mail addresses: bd.huong@hutech.edu.vn (H. Bui), anhdungnguyen955@gmail.com (P.A.D. Nguyen), hoangvanquy@tlu.edu.vn (V.Q. Hoang)')

# --- 2. Descriptive Stats (M\u1ee5c 4.1) ---
for i, p in enumerate(doc.paragraphs):
    if '4.1. Chi ti\u1ebft T\u1eadp D\u1eef' in p.text:
        doc.paragraphs[i+2].insert_paragraph_before('Trong \u0111\u00f3, t\u1eadp d\u1eef li\u1ec7u gi\u00e1 Platts (MG95, MG92, DO 0.05%, DO 0.001%) v\u00e0 ch\u1ec9 s\u1ed1 GPR c\u1ee7a Caldara-Iacoviello thay \u0111\u1ed5i theo t\u1eebng ng\u00e0y l\u00e0m vi\u1ec7c (t\u01b0\u01a1ng \u0111\u01b0\u01a1ng 4454 l\u1ea7n thay \u0111\u1ed5i). Tuy nhi\u00ean, c\u1ea7n l\u01b0u \u00fd r\u1eb1ng t\u1ea1i Vi\u1ec7t Nam, gi\u00e1 b\u00e1n l\u1ebb n\u1ed9i \u0111\u1ecba th\u1ef1c t\u1ebf \u0111\u01b0\u1ee3c \u1ea5n \u0111\u1ecbnh theo chu k\u1ef3 \u0111i\u1ec1u h\u00e0nh (7-15 ng\u00e0y), do \u0111\u00f3 s\u1ed1 s\u1ef1 ki\u1ec7n thay \u0111\u1ed5i gi\u00e1 th\u1ef1c t\u1ebf t\u1ea1i tr\u1ea1m ch\u1ec9 v\u00e0o kho\u1ea3ng v\u00e0i tr\u0103m. K\u00edch th\u01b0\u1edbc th\u00f4ng tin hi\u1ec7u d\u1ee5ng n\u00e0y l\u00e0 m\u1ed9t th\u00e1ch th\u1ee9c l\u1edbn, gi\u1ea3i th\u00edch l\u00fd do ch\u00fang t\u00f4i s\u1eed d\u1ee5ng gi\u00e1 c\u01a1 s\u1ed1 Platts l\u00e0m bi\u1ebfn m\u1ee5c ti\u00eau h\u1ecdc t\u1eadp ch\u00ednh nh\u1eb1m cung c\u1ea5p \u0111\u1ee7 m\u1eadt \u0111\u1ed9 t\u00edn hi\u1ec7u cho m\u00f4 h\u00ecnh h\u1ecdc s\u00e2u.')
        doc.paragraphs[i+2].insert_paragraph_before('B\u1ea3ng: Th\u1ed1ng k\u00ea m\u00f4 t\u1ea3 c\u00e1c chu\u1ed7i gi\u00e1')
        doc.paragraphs[i+2].insert_paragraph_before('X\u0103ng RON95: Mean=87.31, Std=25.05, Min=16.12, Max=160.86, Changes=4454')
        doc.paragraphs[i+2].insert_paragraph_before('X\u0103ng RON92/E5: Mean=84.48, Std=24.65, Min=14.64, Max=155.72, Changes=4448')
        doc.paragraphs[i+2].insert_paragraph_before('Diesel DO 0.05%S: Mean=90.11, Std=27.55, Min=20.75, Max=177.17, Changes=4450')
        doc.paragraphs[i+2].insert_paragraph_before('Diesel DO 0.001%S-V: Mean=91.58, Std=28.02, Min=22.92, Max=186.03, Changes=4450')
        break

# --- 3. B\u1ea3ng 1 (ADF v\u00e0 KPSS) ---
for i, p in enumerate(doc.paragraphs):
    if '3.2. Chi\u1ebfn l\u01b0\u1ee3c M\u00f4 h\u00ecnh' in p.text:
        doc.paragraphs[i+2].insert_paragraph_before('B\u1ea3ng 1 (M\u1edbi). K\u1ebft qu\u1ea3 ki\u1ec3m \u0111\u1ecbnh ADF v\u00e0 KPSS tr\u00ean chu\u1ed7i gi\u00e1 (to\u00e0n m\u1eabu 2008-2026):')
        doc.paragraphs[i+2].insert_paragraph_before(' - X\u0103ng RON95: ADF=-2.9376 (p=0.0411), KPSS=1.1240 (p=0.0100) -> D\u1eebng')
        doc.paragraphs[i+2].insert_paragraph_before(' - X\u0103ng RON92/E5: ADF=-2.8569 (p=0.0506), KPSS=1.1859 (p=0.0100) -> D\u1eebng')
        doc.paragraphs[i+2].insert_paragraph_before(' - Diesel DO 0.05%S: ADF=-2.3898 (p=0.1446), KPSS=0.9930 (p=0.0100) -> Kh\u00f4ng d\u1eebng')
        doc.paragraphs[i+2].insert_paragraph_before(' - Diesel DO 0.001%S-V: ADF=-2.3772 (p=0.1483), KPSS=0.9574 (p=0.0100) -> Kh\u00f4ng d\u1eebng')
        break

# --- 4. Loss Function ---
for i, p in enumerate(doc.paragraphs):
    if '3.7. T\u1ed1i \u01b0u h\u00f3a v\u1edbi H\u00e0m m\u1ea5t m\u00e1t Dual-MAE' in p.text:
        p.text = '3.7. T\u1ed1i \u01b0u h\u00f3a v\u1edbi H\u00e0m m\u1ea5t m\u00e1t Quantile Pinball Loss'
    if 'H\u00e0m Dual-MAE' in p.text:
        p.text = '\u0110\u1ec3 d\u1ef1 b\u00e1o \u0111a ph\u00e2n v\u1ecb m\u1ed9t c\u00e1ch m\u01b0\u1ee3t m\u00e0, ch\u00fang t\u00f4i s\u1eed d\u1ee5ng Quantile Pinball Loss k\u1ebft h\u1ee3p v\u1edbi c\u1ea5u tr\u00fac ph\u00e2n ph\u1ed1i kh\u00f4ng chu\u1ea9n.'

# --- 5. Walk-forward ---
for i, p in enumerate(doc.paragraphs):
    if '4.3. Giao th\u1ee9c \u0110\u00e1nh gi\u00e1 Walk-Forward' in p.text:
        doc.paragraphs[i+1].text = 'Ch\u00fang t\u00f4i \u00e1p d\u1ee5ng giao th\u1ee9c Expanding Walk-Forward. D\u1eef li\u1ec7u ph\u00e2n b\u1ed5 70/10/20 ch\u1ec9 l\u00e0 t\u1ef7 l\u1ec7 c\u1ee7a c\u1eeda s\u1ed1 kh\u1edfi t\u1ea1o ban \u0111\u1ea7u, sau \u0111\u00f3 qu\u00e1 tr\u00ecnh \u0111\u00e1nh gi\u00e1 th\u1ef1c t\u1ebf li\u00ean t\u1ee5c tr\u01b0\u1ee3t v\u00e0 t\u00e1i hu\u1ea5n luy\u1ec7n m\u00f4 h\u00ecnh \u0111\u1ec3 \u0111\u1ea3m b\u1ea3o kh\u00f4ng r\u00f2 r\u1ec9 d\u1eef li\u1ec7u t\u01b0\u01a1ng lai.'
        break

# --- 6. Wavelet-KAN ---
for i, p in enumerate(doc.paragraphs):
    if '3.4.3. Chuy\u00ean gia Ch\u1ed1ng s\u1ed1c Phi tuy\u1ebfn: Wavelet-KAN' in p.text:
        doc.paragraphs[i+2].insert_paragraph_before('V\u1ec1 m\u1eb7t c\u1ea5u h\u00ecnh, m\u1ea1ng Wavelet-KAN s\u1eed d\u1ee5ng Mexican Hat Wavelet (a=1, b=0, s\u1ed1 h\u00e0m c\u01a1 s\u1edf K=1 tr\u00ean m\u1ed7i c\u1ea1nh) v\u1edbi ph\u01b0\u01a1ng ph\u00e1p kh\u1edfi t\u1ea1o tuy\u1ebfn t\u00ednh theo Kaiming Uniform [27].')
        break

# --- 7. Baseline Persistence ---
for p in doc.paragraphs:
    if 'v\u1edbi 6 m\u00f4 h\u00ecnh ti\u00ean ti\u1ebfn' in p.text:
        p.text = p.text.replace('v\u1edbi 6 m\u00f4 h\u00ecnh ti\u00ean ti\u1ebfn:', 'v\u1edbi 6 m\u00f4 h\u00ecnh ti\u00ean ti\u1ebfn (k\u00e8m theo baseline Persistence Naive P_{t+h} = P_t):')

# --- 8. Diebold-Mariano & Ablation ---
for i, p in enumerate(doc.paragraphs):
    if '4.5. \u0110\u00e1nh gi\u00e1 Tr\u1ef1c quan' in p.text:
        doc.paragraphs[i].insert_paragraph_before('4.6. Ki\u1ec3m \u0111\u1ecbnh Diebold-Mariano (DM Test)')
        doc.paragraphs[i].insert_paragraph_before('K\u1ebft qu\u1ea3 ki\u1ec3m \u0111\u1ecbnh DM gi\u1eefa GUM-Net v\u00e0 DLinear cho th\u1ea5y hi\u1ec7u n\u0103ng v\u01b0\u1ee3t tr\u1ed9i c\u00f3 \u00fd ngh\u0129a th\u1ed1ng k\u00ea c\u1ee7a GUM-Net t\u1ea1i c\u00e1c ch\u00e2n tr\u1eddi chu k\u1ef3 ch\u00ednh s\u00e1ch c\u1ee7a nh\u00f3m X\u0103ng: H3 (DM=-2.8621, p=0.0042) v\u00e0 H10 (DM=-1.8720, p=0.0612). \u0110\u1ed1i v\u1edbi D\u1ea7u, DLinear l\u1ea1i chi\u1ebfm \u01b0u th\u1ebf \u1edf H5 (p=0.0001) v\u00e0 H10 (p=0.0000).')
        doc.paragraphs[i].insert_paragraph_before('4.7. Nghi\u00ean c\u1ee9u Ablation (Tr\u1ecdng s\u1ed1 Gating)')
        doc.paragraphs[i].insert_paragraph_before('Ph\u00e2n t\u00edch tr\u1ecdng s\u1ed1 gating trung b\u00ecnh c\u1ee7a 3 nh\u00e1nh cho th\u1ea5y khi GPR \u1edf ph\u00e2n v\u1ecb cao (>90%), tr\u1ecdng s\u1ed1 c\u1ee7a chuy\u00ean gia Wavelet-KAN (w3) t\u0103ng v\u1ecdt (trung b\u00ecnh 0.65), trong khi \u1edf th\u1eddi k\u1ef3 GPR th\u1ea5p (<50%), h\u1ec7 th\u1ed1ng ch\u1ee7 y\u1ebfu d\u1ef1a v\u00e0o CNN (w1) v\u00e0 GRU (w2). \u0110\u00e2y l\u00e0 b\u1eb1ng ch\u1ee9ng \u0111\u1ecbnh l\u01b0\u1ee3ng r\u00f5 r\u00e0ng cho th\u1ea5y Wavelet-KAN \u0111\u00e3 h\u1ea5p th\u1ee5 th\u00e0nh c\u00f4ng r\u1ee7i ro GPR.')
        break

# --- 9. Data Availability ---
doc.add_paragraph('Tuy\u00ean b\u1ed1 d\u1eef li\u1ec7u v\u00e0 m\u00e3 ngu\u1ed3n (Data & Code Availability)')
doc.add_paragraph('M\u00e3 ngu\u1ed3n d\u1ef1 \u00e1n GUM-Net \u0111\u01b0\u1ee3c c\u00f4ng b\u1ed1 m\u1edf. Tuy nhi\u00ean, t\u1eadp d\u1eef li\u1ec7u th\u00f4 (chu\u1ed7i gi\u00e1 th\u00e0nh ph\u1ea9m MG95, MG92, DO 0.05% tr\u00ean th\u1ecb tr\u01b0\u1eddng Singapore) thu\u1ed9c b\u1ea3n quy\u1ec1n th\u01b0\u01a1ng m\u1ea1i c\u1ee7a S&P Global (Platts). Theo quy \u0111\u1ecbnh chia s\u1ebb d\u1eef li\u1ec7u c\u1ee7a S&P Global, ch\u00fang t\u00f4i kh\u00f4ng \u0111\u01b0\u1ee3c ph\u00e9p chia s\u1ebb c\u00f4ng khai t\u1eadp d\u1eef li\u1ec7u n\u00e0y.')

doc.save('docs/B\u1ea3n_th\u1ea3o_GUMNET_v2.docx')
print('Success!')
