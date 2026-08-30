import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re
import pandas as pd

# Define helper to convert string or math token into Word XML
def create_inline_omml(xml_body):
    return f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{xml_body}</m:oMath>'

# Math snippets dictionary
M = {
    'P_tc': '<m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t,c</m:t></m:r></m:sub></m:sSub>',
    'c_set': '<m:r><m:t>c ∈ {MG95, DO 0.001%}</m:t></m:r>',
    'L_30': '<m:r><m:t>L = 30</m:t></m:r>',
    'h_set': '<m:r><m:t>h ∈ {1, 3, 5, 7, 10, 20, 60}</m:t></m:r>',
    'r_thc': '<m:sSub><m:e><m:r><m:t>r</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t> = ln(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t>) − ln(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t,c</m:t></m:r></m:sub></m:sSub><m:r><m:t>)</m:t></m:r>',
    'P_hat_thc': '<m:sSub><m:e><m:r><m:t>P̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t,c</m:t></m:r></m:sub></m:sSub><m:r><m:t> · exp(</m:t></m:r><m:sSub><m:e><m:r><m:t>r̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h,c</m:t></m:r></m:sub></m:sSub><m:r><m:t>)</m:t></m:r>',
    'X_price': '<m:sSub><m:e><m:r><m:t>X</m:t></m:r></m:e><m:sub><m:r><m:t>price</m:t></m:r></m:sub></m:sSub>',
    'X_macro': '<m:sSub><m:e><m:r><m:t>X</m:t></m:r></m:e><m:sub><m:r><m:t>macro</m:t></m:r></m:sub></m:sSub>',
    'X_shock': '<m:sSub><m:e><m:r><m:t>X</m:t></m:r></m:e><m:sub><m:r><m:t>shock</m:t></m:r></m:sub></m:sSub>',
    'k_set': '<m:r><m:t>k ∈ {3, 7, 15}</m:t></m:r>',
    'router_in': '<m:r><m:t>[</m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>CNN</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>GRU</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>KAN</m:t></m:r></m:sub></m:sSub><m:r><m:t> ‖ Emb(h) ‖ </m:t></m:r><m:sSub><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>ctx</m:t></m:r></m:sub></m:sSub><m:r><m:t>]</m:t></m:r>',
    'w_h': '<m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> = Softmax(MLP(·))</m:t></m:r>',
    'f_fused': '<m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>fused</m:t></m:r></m:sub></m:sSub><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>∑</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>h,i</m:t></m:r></m:sub></m:sSub><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>',
    'x_ctx': '<m:sSub><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>ctx</m:t></m:r></m:sub></m:sSub>',
    'q_set': '<m:r><m:t>q ∈ {0.1, 0.5, 0.9}</m:t></m:r>',
    'gamma_h_set': '<m:sSub><m:e><m:r><m:t>γ</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> ∈ (0, 1)</m:t></m:r>',
    'r_hat_q': '<m:sSubSup><m:e><m:r><m:t>r̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub><m:sup><m:r><m:t>(q)</m:t></m:r></m:sup></m:sSubSup><m:r><m:t> = </m:t></m:r><m:sSub><m:e><m:r><m:t>Head</m:t></m:r></m:e><m:sub><m:r><m:t>q</m:t></m:r></m:sub></m:sSub><m:r><m:t>(</m:t></m:r><m:sSub><m:e><m:r><m:t>f</m:t></m:r></m:e><m:sub><m:r><m:t>fused</m:t></m:r></m:sub></m:sSub><m:r><m:t>) + </m:t></m:r><m:sSub><m:e><m:r><m:t>γ</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub><m:r><m:t> · </m:t></m:r><m:sSubSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub><m:sup><m:r><m:t>target</m:t></m:r></m:sup></m:sSubSup>',
    'L_tot': '<m:r><m:t>ℒ = </m:t></m:r><m:sSub><m:e><m:r><m:t>ℒ</m:t></m:r></m:e><m:sub><m:r><m:t>pinball</m:t></m:r></m:sub></m:sSub><m:r><m:t> + α</m:t></m:r><m:sSub><m:e><m:r><m:t>ℒ</m:t></m:r></m:e><m:sub><m:r><m:t>balance</m:t></m:r></m:sub></m:sSub>',
    'alpha_val': '<m:r><m:t>α = 0.01</m:t></m:r>',
    'L_bal': '<m:sSub><m:e><m:r><m:t>ℒ</m:t></m:r></m:e><m:sub><m:r><m:t>balance</m:t></m:r></m:sub></m:sSub>',
    'gamma_h': '<m:sSub><m:e><m:r><m:t>γ</m:t></m:r></m:e><m:sub><m:r><m:t>h</m:t></m:r></m:sub></m:sSub>',
    't_minus_1': '<m:r><m:t>t − 1</m:t></m:r>',
    'N_val': '<m:r><m:t>N = 4.512</m:t></m:r>',
    'R2': '<m:sSup><m:e><m:r><m:t>R</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>',
    'DA_formula': '<m:r><m:t>DA = </m:t></m:r><m:f><m:num><m:r><m:t>1</m:t></m:r></m:num><m:den><m:r><m:t>N</m:t></m:r></m:den></m:f><m:sSubSup><m:e><m:r><m:t>∑</m:t></m:r></m:e><m:sub><m:r><m:t>t=1</m:t></m:r></m:sub><m:sup><m:r><m:t>N</m:t></m:r></m:sup></m:sSubSup><m:r><m:t> 𝕀[sign(</m:t></m:r><m:sSub><m:e><m:r><m:t>P̂</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub></m:sSub><m:r><m:t> − </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub></m:sSub><m:r><m:t>) = sign(</m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t+h</m:t></m:r></m:sub></m:sSub><m:r><m:t> − </m:t></m:r><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub></m:sSub><m:r><m:t>)]</m:t></m:r>',
    'q_interval': '<m:r><m:t>[</m:t></m:r><m:sSub><m:e><m:r><m:t>q</m:t></m:r></m:e><m:sub><m:r><m:t>0.1</m:t></m:r></m:sub></m:sSub><m:r><m:t>, </m:t></m:r><m:sSub><m:e><m:r><m:t>q</m:t></m:r></m:e><m:sub><m:r><m:t>0.9</m:t></m:r></m:sub></m:sSub><m:r><m:t>]</m:t></m:r>',
    'p_val': '<m:r><m:t>p &lt; 0.001</m:t></m:r>',
    'mexican_hat': '<m:r><m:t>ψ(z) = (1 − </m:t></m:r><m:sSup><m:e><m:r><m:t>z</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup><m:r><m:t>) · </m:t></m:r><m:sSup><m:e><m:r><m:t>e</m:t></m:r></m:e><m:sup><m:r><m:t>−0.5z²</m:t></m:r></m:sup></m:sSup>'
}

def add_styled_paragraph(p, segments, line_spacing=1.15, space_after=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p.text = ""
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    
    for seg in segments:
        if isinstance(seg, str):
            # Parse text with Table/Figure/Reference links
            parts = re.split(r'(Bảng \d+|Hình \d+|\[\d+\])', seg)
            for part in parts:
                if not part:
                    continue
                m_tbl = re.match(r'Bảng (\d+)', part)
                m_fig = re.match(r'Hình (\d+)', part)
                m_ref = re.match(r'\[(\d+)\]', part)
                if m_tbl:
                    t_id = f"tbl_{m_tbl.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{t_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                elif m_fig:
                    f_id = f"fig_{m_fig.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{f_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                elif m_ref:
                    r_id = f"ref_{m_ref.group(1)}"
                    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                                   f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                                   f'<w:t>{part}</w:t></w:r></w:hyperlink>')
                    p._p.append(hl)
                else:
                    run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9.5)
        elif isinstance(seg, tuple) and seg[0] == 'm':
            m_xml = M[seg[1]]
            p._p.append(parse_xml(create_inline_omml(m_xml)))
        elif isinstance(seg, tuple) and seg[0] == 'ref':
            r_id = f"ref_{seg[1]}"
            hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                           f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                           f'<w:t>[{seg[1]}]</w:t></w:r></w:hyperlink>')
            p._p.append(hl)
        elif isinstance(seg, tuple) and seg[0] == 'refs':
            for idx, r_num in enumerate(seg[1]):
                r_id = f"ref_{r_num}"
                hl = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{r_id}" w:history="1">'
                               f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="19"/><w:color w:val="1A56DB"/></w:rPr>'
                               f'<w:t>[{r_num}]</w:t></w:r></w:hyperlink>')
                p._p.append(hl)
                if idx < len(seg[1]) - 1:
                    r_c = p.add_run(", ")
                    r_c.font.name = 'Times New Roman'
                    r_c.font.size = Pt(9.5)

def build_perfect_conference_doc(out_fname):
    doc = docx.Document('GUMNetHet_FAIR_ban_thao_template.docx')
    
    # 1. P[0] Title
    p0 = doc.paragraphs[0]
    p0.text = "ROBUST PROBABILISTIC ENERGY FORECASTING UNDER GEOPOLITICAL SHOCKS: AN ADAPTIVE MIXTURE OF LOCAL-GLOBAL EXPERTS"
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p0.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.bold = True
        
    # 2. P[1] Authors
    p1 = doc.paragraphs[1]
    p1.text = "Nguyễn Phước Anh Dũng¹, Bùi Danh Hường¹*, Hoàng Văn Quý²"
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p1.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.bold = True
        
    # 3. P[2] Affiliations
    p2 = doc.paragraphs[2]
    p2.text = (
        "¹Khoa Công nghệ Thông tin, Trường Đại học Công nghệ TP.HCM (HUTECH), TP. Hồ Chí Minh, Việt Nam\n"
        "²Khoa Công nghệ Thông tin, Trường Đại học Thủy lợi (TLU), Hà Nội, Việt Nam\n"
        "*Tác giả liên hệ: Bùi Danh Hường (bd.huong@hutech.edu.vn)"
    )
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p2.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(8.5)
        r.italic = True

    # 4. P[5] Abstract
    add_styled_paragraph(doc.paragraphs[5], [
        "Tóm tắt—Dự báo giá xăng dầu thành phẩm trở nên khó khăn khi chuỗi giá chịu đồng thời biến động ngắn hạn, dịch chuyển chế độ vĩ mô "
        "và các cú sốc địa chính trị có tính phi tuyến, đuôi dày. Bài báo đề xuất GUMNetHet (Heterogeneous Gated Unified Mixture Network), "
        "một mạng kết hợp chuyên gia không đồng nhất cho dự báo xác suất đa chu kỳ. Mô hình phân vùng đặc trưng thành ba nhóm và giao cho "
        "ba chuyên gia chuyên biệt: CNN-1D đa tỷ lệ cho động lượng giá, GRU-Attention cho chế độ vĩ mô, và Wavelet-KAN cho quan hệ phi tuyến "
        "nhạy cú sốc. Bộ định tuyến nhận biết horizon kết hợp biểu diễn chuyên gia, nhúng horizon và thống kê ngữ cảnh để phân bổ trọng số động; "
        "đầu ra đa phân vị ", ('m', 'q_set'), " được tối ưu bằng pinball loss và điều chuẩn cân bằng tải. Trên dữ liệu đa nguồn 11/2008–04/2026 (",
        ('m', 'N_val'), " quan sát) với expanding walk-forward, GUMNetHet đạt MAE thấp nhất trong nhóm baseline được báo cáo ở cả bảy horizon H1–H60 cho "
        "MG95 và DO 0.001%. Tại H60, MAE giảm 30,1% cho xăng và 22,9% cho dầu so với baseline tốt nhất tương ứng. Directional accuracy đạt "
        "90,95–95,56% ở H1–H7 của xăng và 76,65–84,92% ở H1–H7 của dầu, nhưng giảm mạnh tại H10/H60, cho thấy cần phân biệt giữa độ chính xác "
        "mức giá và độ chính xác hướng ở horizon dài. Khoảng phân vị 80% đạt PICP=82,4% với PINAW=0,142. Với 0,34M tham số và độ trễ 1,42 ms/mẫu, "
        "GUMNetHet cho thấy sự cân bằng tốt giữa độ chính xác, bất định và chi phí tính toán."
    ], line_spacing=1.10, space_after=4)

    # 5. P[8], P[9], P[10] Introduction
    add_styled_paragraph(doc.paragraphs[8], [
        "Thị trường xăng dầu Việt Nam có đặc thù cấu trúc nguồn cung mang tính chiến lược: khoảng 70% tổng sản lượng "
        "tiêu thụ nội địa được cung ứng bởi hai nhà máy lọc hóa dầu trong nước (Dung Quất và Nghi Sơn), trong khi 30% "
        "nhu cầu còn lại bắt buộc phải nhập khẩu trực tiếp từ các thị trường quốc tế. Trong đó, thị trường Singapore là "
        "địa bàn nhập khẩu trọng yếu nhất, với giá giao dịch thành phẩm Mean of Platts Singapore (MOPS)—tiêu biểu là Mogas 95 "
        "(MG95) và Gasoil 0.001%S (DO 0.001%)—đóng vai trò là hệ quy chiếu định giá cơ sở cho mọi hợp đồng thương mại, "
        "tính toán giá vốn và quản lý chuỗi cung ứng. Tuy nhiên, giá Platts biến động cực kỳ phức tạp do phản ứng đồng thời "
        "với các cú sốc cung–cầu toàn cầu, biến động tỷ giá và rủi ro địa chính trị (GPR) ", ('refs', [1, 2, 3]),
        ". Các giai đoạn sụp đổ giá dầu 2014–2016, chiến tranh giá OPEC+ năm 2020, xung đột Nga–Ukraine 2022 và căng thẳng Biển Đỏ 2023–2024 cho thấy chuỗi giá "
        "thành phẩm thường xuyên xuất hiện bước nhảy phi tuyến, dịch chuyển chế độ và phân cụm biến động mạnh. Do đó, nhu cầu "
        "dự báo chính xác giá xăng dầu Platts trong ngắn hạn (H1–H7) và trung hạn (H10–H60) là đòi hỏi cấp thiết phục vụ trực "
        "tiếp cho việc ra quyết định kinh doanh, tối ưu hóa kế hoạch mua hàng, quản trị tồn kho và phòng hộ rủi ro (hedging) "
        "của các doanh nghiệp đầu mối xăng dầu lớn như Petrolimex và PVOIL."
    ])

    add_styled_paragraph(doc.paragraphs[9], [
        "Các kiến trúc chuỗi thời gian hiện đại như PatchTST [8], iTransformer [9], TimesNet [10], DLinear [11], Mamba [15] và "
        "Chronos [16] đã cải thiện đáng kể hiệu năng trên các benchmark tổng quát, nhưng phần lớn đều xử lý toàn bộ các biến đầu vào "
        "trong một không gian biểu diễn tương đối đồng nhất. Với dữ liệu năng lượng, ba nhóm tín hiệu có bản chất cơ bản khác nhau: "
        "động lượng giá tần số cao, trạng thái vĩ mô biến đổi chậm và phản ứng phi tuyến nhạy cú sốc biên độ lớn đòi hỏi các inductive "
        "bias chuyên biệt. Các mạng kết hợp chuyên gia (MoE) truyền thống ", ('refs', [20, 21, 22]),
        " thường đưa cùng một tập đặc trưng tới mọi expert, làm suy giảm mức độ chuyên môn hóa."
    ])

    add_styled_paragraph(doc.paragraphs[10], [
        "GUMNetHet giải quyết triệt để điểm nghẽn này bằng kỹ thuật phân vùng đặc trưng (feature partitioning): nhóm giá và benchmark "
        "được xử lý bởi CNN-1D đa tỷ lệ; nhóm vĩ mô và chỉ số GPR bởi GRU-Attention; nhóm tỷ lệ crack-spread và độ biến động bởi Wavelet-KAN. "
        "Ba biểu diễn được hợp nhất linh hoạt thông qua một bộ định tuyến (gating router) phụ thuộc vào horizon dự báo và ngữ cảnh thị trường. "
        "Đóng góp chính của bài báo gồm: (i) Kiến trúc MoE dị thể với phân vùng đặc trưng theo bản chất kinh tế và miền tần số; (ii) Bộ định tuyến "
        "nhận biết horizon (horizon-aware routing); (iii) Đầu ra đa phân vị (multi-quantile head) kết hợp residual scaling để kiểm soát độ trôi phương sai; "
        "và (iv) Đánh giá thực nghiệm mở rộng walk-forward trên ", ('m', 'N_val'), " ngày giao dịch với bảng kết quả đầy đủ các chỉ số MAE, RMSE, MAPE, ",
        ('m', 'R2'), " và DA%."
    ])

    # 6. P[12], P[13] Related Work
    add_styled_paragraph(doc.paragraphs[12], [
        "Kilian [2] và Baumeister & Kilian [3] đặt nền tảng cho việc phân tách cú sốc cung–cầu trong thị trường dầu, trong khi chỉ số GPR của Caldara & Iacoviello [1] "
        "cung cấp thước đo định lượng cho rủi ro địa chính trị. Ở phía mô hình, PatchTST [8], iTransformer [9] và TimesNet [10] đại diện cho các thiết kế "
        "Transformer/biến thiên thời gian hiện đại; DLinear [11] cho thấy mô hình tuyến tính có phân rã vẫn là baseline mạnh. TFT [12], N-BEATS [13] và N-HiTS [14] "
        "mở rộng dự báo đa horizon; Mamba [15] đưa state-space chọn lọc vào sequence modeling; Chronos [16], TimesFM [17], MOIRAI [18] và TTM [19] mở rộng sang foundation time-series models."
    ])

    add_styled_paragraph(doc.paragraphs[13], [
        "Về MoE, Jacobs et al. [20] giới thiệu adaptive mixtures of local experts; Shazeer et al. [21] và Switch Transformer [22] phát triển sparse gating ở quy mô lớn. "
        "KAN [23] và Wav-KAN [24] thay thế activation cố định bằng hàm cơ sở có thể học, trong đó wavelet phù hợp với tín hiệu cục bộ đa độ phân giải [25]. "
        "Time-MoE [26] và TimeMixer++ [27] tiếp tục cho thấy giá trị của expert routing trong chuỗi thời gian. Khác với các thiết kế expert đồng nhất, GUMNetHet gán các loại "
        "kiến trúc khác nhau cho các nhóm biến khác nhau và điều kiện hóa router trên horizon."
    ])

    # 7. P[16] Section 3.1: Problem & Partitioning (WITH FULL OMML)
    add_styled_paragraph(doc.paragraphs[16], [
        "Với giá mục tiêu ", ('m', 'P_tc'), " của sản phẩm ", ('m', 'c_set'), ", lookback ", ('m', 'L_30'), " và horizon ", ('m', 'h_set'),
        ", mục tiêu học là lợi suất log tích lũy ", ('m', 'r_thc'), ". Dự báo mức giá được khôi phục bởi ", ('m', 'P_hat_thc'),
        ". Đầu vào được chia thành ", ('m', 'X_price'), " (Platts liên sản phẩm, WTI, Brent), ", ('m', 'X_macro'), " (DXY, GPR, sản lượng và biến làm mịn), và ",
        ('m', 'X_shock'), " (crack-spread, realized volatility, biến lịch)."
    ])

    # 8. P[18] Section 3.2: 3 Experts & Router (WITH FULL OMML)
    add_styled_paragraph(doc.paragraphs[18], [
        "Expert giá dùng ba nhánh Conv1D với kernel ", ('m', 'k_set'), " và attention pooling; expert vĩ mô dùng GRU hai lớp với dropout 0,1; "
        "expert shock-sensitive dùng Wav-KAN với hàm cơ sở Mexican Hat ", ('m', 'mexican_hat'), ". Router nhận ", ('m', 'router_in'),
        " và sinh ", ('m', 'w_h'), "; biểu diễn hợp nhất là ", ('m', 'f_fused'), ". ", ('m', 'x_ctx'), " gồm thống kê mức và biến động gần nhất, "
        "nhờ đó trọng số expert thay đổi theo cả horizon lẫn chế độ thị trường."
    ])

    # 9. P[22] Section 3.3: Quantile Head & Loss (WITH FULL OMML)
    add_styled_paragraph(doc.paragraphs[22], [
        "Head sinh ba phân vị ", ('m', 'q_set'), ". Residual scaling học hệ số neo ", ('m', 'gamma_h_set'), " để giới hạn độ trôi ở horizon dài, với dạng ",
        ('m', 'r_hat_q'), ". Hàm mục tiêu là ", ('m', 'L_tot'), ", ", ('m', 'alpha_val'), "; ", ('m', 'L_bal'), " điều chuẩn phân bố trọng số router "
        "để giảm nguy cơ gate collapse. Trong bản revised, ", ('m', 'gamma_h'), " học được nằm trong khoảng 0,18–0,42, cho thấy cơ chế neo không suy thoái thành persistence baseline."
    ])

    # 10. P[25], P[26] Section 4.1: Data & Walk-forward
    add_styled_paragraph(doc.paragraphs[25], [
        "Dữ liệu bao phủ 03/11/2008–30/04/2026 với ", ('m', 'N_val'), " quan sát ngày giao dịch. Hai target được báo cáo là MG95 và DO 0.001% theo Platts; "
        "các biến ngoại sinh gồm Platts liên sản phẩm, WTI, Brent, GPR [1], DXY, sản lượng dầu, crack-spread, realized volatility và biến lịch. "
        "Bài báo không dự báo giá bán lẻ điều hành trong nước; các kết luận được giới hạn ở giá thành phẩm Platts."
    ])

    add_styled_paragraph(doc.paragraphs[26], [
        "Để tránh look-ahead bias, các biến ngày được dùng ở ", ('m', 't_minus_1'), "; GPR được áp lag 30 ngày lịch; sản lượng dầu áp lag 7 ngày; "
        "các biến rolling được tính sau khi lag. Mọi scaler chỉ fit trên train tại từng bước walk-forward. Lookback ", ('m', 'L_30'), "; train/validation=85/15 ở mỗi lần mở rộng. "
        "Cửa sổ test tăng từ 100 ngày ở H1–H5 lên 600 ngày ở H60."
    ])

    # 11. P[28] Section 4.2: Diagnostics
    add_styled_paragraph(doc.paragraphs[28], [
        "Kiểm định ADF và KPSS trong Bảng 1 cho bằng chứng về tính dừng ở mức giá; sau chuyển log-return, kiểm định ADF bác bỏ giả thuyết nghiệm đơn vị (",
        ('m', 'p_val'), ") trên tất cả các chuỗi. Độ nhọn (Kurtosis) của lợi suất rất cao (đạt 213,35 ở WTI và 17,51 ở MG95), khẳng định tính chất đuôi dày và các cú sốc cực đoan trong chuỗi giá năng lượng."
    ])

    # 12. P[32] Section 4.3: Baselines & Metrics (WITH FULL OMML)
    add_styled_paragraph(doc.paragraphs[32], [
        "Các baseline đối chuẩn được báo cáo chi tiết gồm 6 mô hình đại diện tiêu biểu: PatchTST [8], iTransformer [9], TimesNet [10], "
        "DLinear [11], Chronos [16] và BiMamba (ở bảng xăng của nguồn thực nghiệm). Các mô hình dùng chung lookback ", ('m', 'L_30'), ", quy tắc trễ dữ liệu và giao thức "
        "walk-forward expanding. Các chỉ số đánh giá điểm gồm MAE, RMSE, MAPE (%) và Hệ số xác định ", ('m', 'R2'), ". Độ chính xác xu hướng được tính theo ",
        ('m', 'DA_formula'), ". Đánh giá phân phối xác suất sử dụng hệ số bao phủ PICP và độ rộng dải chuẩn hóa PINAW."
    ])

    # 13. P[36] Section 5.1: Results
    add_styled_paragraph(doc.paragraphs[36], [
        "Hình 2 cùng Bảng 2 và Bảng 3 cho thấy GUMNetHet duy trì MAE thấp hơn nhóm baseline được báo cáo ở cả hai sản phẩm, đặc biệt ở H20–H60. Ở H60, MAE của xăng là 4,847 "
        "so với 6,933 của baseline tốt nhất (giảm 30,1%); với dầu là 7,066 so với 9,167 (giảm 22,9%). Tuy nhiên, ", ('m', 'R2'), " tại H60 giảm còn 0,155 (xăng) và −0,007 (dầu), "
        "vì vậy kết quả dài hạn nên được hiểu là ổn định sai số mức giá tốt hơn baseline, không phải dự báo quỹ đạo dài hạn hoàn hảo."
    ])

    # 14. P[42] Section 5.2: DA%
    add_styled_paragraph(doc.paragraphs[42], [
        "Độ chính xác hướng trong Hình 3 cho thấy một hành vi khác với sai số mức. Với xăng (Bảng 2), GUMNetHet đạt 91,46%, 91,37%, 90,95% và 95,56% ở H1, H3, H5, H7; "
        "với dầu (Bảng 3) tương ứng là 84,92%, 76,65%, 76,88% và 83,28%. Ở H20, DA vẫn cao (91,65% xăng; 71,11% dầu). Ngược lại, H10 và H60 giảm dưới 50%, "
        "lần lượt 42,24%/27,95% cho xăng và 32,29%/19,10% cho dầu. Do đó, GUMNetHet phù hợp hơn như mô hình dự báo mức và biên bất định ở horizon dài, thay vì công cụ sinh tín hiệu hướng."
    ])

    # 15. P[46], P[48] Section 5.3: Ablation & Router
    add_styled_paragraph(doc.paragraphs[46], [
        "Khoảng ", ('m', 'q_interval'), " đạt PICP=82,4% so với danh định 80% và PINAW=0,142. Hình 4 cho thấy biên bất định mở rộng khi biến động tăng. "
        "Ablation trong Bảng 4 cho thấy thay Wav-KAN bằng MLP gây suy giảm lớn nhất trong các biến thể expert; router đồng nhất cũng làm MAE tăng "
        "đáng kể, củng cố vai trò của chuyên môn hóa và định tuyến thích ứng."
    ])

    add_styled_paragraph(doc.paragraphs[48], [
        "Bản revised còn ghi nhận rằng loại bỏ residual scaling làm MAE tăng khoảng 8,5%/6,3% ở H20 và 14,1%/11,8% ở H60 cho xăng/dầu. Phân tích router trong "
        "Hình 4 cho thấy ở GPR thấp và horizon ngắn, CNN có trọng số trung bình khoảng 0,48; khi GPR vượt phân vị 90%, trọng số Wav-KAN tăng từ khoảng 0,29 lên 0,61 ở horizon trung–dài và CNN giảm xuống khoảng 0,21. Các kết quả này cho thấy router thực sự thay đổi chế độ thay vì chỉ trung bình hóa đầu ra."
    ])

    # 16. P[53] Section 5.4: Computational Cost
    add_styled_paragraph(doc.paragraphs[53], [
        "Kết quả trong Bảng 5 cho thấy GUMNetHet không nhẹ bằng DLinear, nhưng nhỏ hơn đáng kể các Transformer đại diện và vẫn đạt độ trễ 1,42 ms/mẫu. Vì vậy lợi thế của mô hình là hiệu quả tương đối cao trong nhóm phi tuyến mạnh, phù hợp cho hệ thống giám sát và cập nhật dự báo cận thời gian thực."
    ])

    # 17. P[58] Section 7: Conclusion
    add_styled_paragraph(doc.paragraphs[58], [
        "Bài báo trình bày GUMNetHet, một MoE dị thể cho dự báo xác suất đa horizon giá xăng dầu thành phẩm dưới cú sốc địa chính trị. Kiến trúc kết hợp CNN đa tỷ lệ, GRU-Attention và Wavelet-KAN thông qua router nhận biết horizon, cùng đầu ra đa phân vị và residual scaling. Trên expanding walk-forward, GUMNetHet đạt MAE thấp nhất trong nhóm baseline được báo cáo ở H1–H60 cho cả MG95 và DO 0.001%; lợi thế rõ nhất xuất hiện ở H60. Các bảng kết quả đầy đủ cho thấy ưu thế này đi kèm ",
        ('m', 'R2'), " suy giảm ở horizon dài và DA thấp tại H10/H60, qua đó xác định rõ phạm vi sử dụng của mô hình. Ablation, fan chart và phân tích router hỗ trợ giả thuyết rằng phân vùng đặc trưng và định tuyến theo chế độ là nguồn chính của hiệu năng. Công việc tiếp theo tập trung vào multi-seed robustness, probabilistic baselines, kiểm định thống kê cho multi-step forecasts và tích hợp news embedding địa chính trị theo thời gian thực."
    ])

    # Format all Table captions with Bookmarks
    def make_table_caption(p, text, tbl_id):
        p.style = 'Caption'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.text = ""
        bm_id = f"300{tbl_id}"
        bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="tbl_{tbl_id}"/>')
        bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>')
        p._p.append(bm_start)
        p._p.append(bm_end)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8.5)
        run.bold = False

    def make_figure_caption(p, text, fig_id):
        p.style = 'Caption'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        p.text = ""
        bm_id = f"400{fig_id}"
        bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="fig_{fig_id}"/>')
        bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>')
        p._p.append(bm_start)
        p._p.append(bm_end)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8.0)
        run.bold = False

    make_table_caption(doc.paragraphs[29], "Bảng 1. Chẩn đoán thống kê rút gọn của các chuỗi chính.", "1")
    make_table_caption(doc.paragraphs[39], "Bảng 2. Kết quả chi tiết trên MG95 (Seed=42). MAE/RMSE tính theo USD/thùng; giá trị DA là phần trăm.", "2")
    make_table_caption(doc.paragraphs[40], "Bảng 3. Kết quả chi tiết trên DO 0.001% (Seed=42). Bảng nguồn thực nghiệm không báo cáo BiMamba cho mục tiêu dầu.", "3")
    make_table_caption(doc.paragraphs[47], "Bảng 4. Ablation rút gọn của GUMNetHet (Seed=42).", "4")
    make_table_caption(doc.paragraphs[52], "Bảng 5. Chi phí tính toán trên GPU Tesla T4.", "5")

    make_figure_caption(doc.paragraphs[20], "Hình 1. Kiến trúc GUMNetHet: phân vùng đặc trưng, ba expert dị thể, horizon-aware router và multi-quantile head.", "1")
    make_figure_caption(doc.paragraphs[38], "Hình 2. Đường cong MAE và R² qua bảy horizon cho MG95 và DO 0.001% (Seed=42).", "2")
    make_figure_caption(doc.paragraphs[44], "Hình 3. Directional accuracy (DA%) của GUMNetHet và các baseline qua H1–H60 cho xăng và dầu.", "3")
    make_figure_caption(doc.paragraphs[50], "Hình 4. Trên: fan chart đa phân vị dưới biến động mạnh. Dưới: trọng số router trong chế độ GPR thấp và GPR cao.", "4")

    # Format all tables: 100% pure white, standard borders
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = table._tbl.tblPr
        for old_b in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old_b)
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                            f'<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                            f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
                            f'<w:insideV w:val="none"/>'
                            f'<w:left w:val="none"/>'
                            f'<w:right w:val="none"/>'
                            f'</w:tblBorders>')
        tblPr.append(borders)
        
        for row_idx, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            if row_idx == 0:
                trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                for old_shd in tcPr.findall(qn('w:shd')):
                    tcPr.remove(old_shd)
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="FFFFFF"/>')
                tcPr.append(shd)
                mar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:left w:w="100" w:type="dxa"/><w:right w:w="100" w:type="dxa"/></w:tcMar>')
                tcPr.append(mar)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.05
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(8.5)

    doc.save(out_fname)
    print(f"Generated {out_fname} with 100% OMML formatted formulas!")

if __name__ == '__main__':
    build_perfect_conference_doc('GUMNETHet_FAIRv2_template.docx')
    build_perfect_conference_doc('GUMNETHET_FAIR_v2_TIENG_VIET.docx')
    build_perfect_conference_doc('GUMNETHET_FAIR_v2.docx')
